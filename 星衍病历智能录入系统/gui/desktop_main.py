"""
MedVoiceApp —— 原生 PyQt5 桌面主窗口（从 main.py 拆分）
"""
import sys
import os
import re
import json
import time
import threading
import difflib

from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QPushButton, QTextEdit, QComboBox, QLabel, QSplitter,
    QListWidget, QListWidgetItem, QStatusBar, QToolBar,
    QMessageBox, QCheckBox, QGroupBox, QFileDialog,
    QDialog, QLineEdit, QTableWidget, QTableWidgetItem,
    QTabWidget, QHeaderView, QTextBrowser, QToolButton,
    QAction, QMenu, QScrollArea, QFrame, QInputDialog,
)
from PyQt5.QtCore import Qt, QTimer, pyqtSignal, QThread
from PyQt5.QtGui import QFont, QColor, QPalette, QTextCursor, QTextCharFormat, QTextDocument

try:
    from PyQt5.QtWebEngineWidgets import QWebEngineView  # noqa: F401
    _HAS_WEBENGINE = True
except ImportError:
    _HAS_WEBENGINE = False

from corrector import Corrector
from asr_engine import ASREngine
from template_engine import TemplateEngine
from rule_engine import RuleEngine
from section_parser import SectionParser, SmartDictation
from medical_classifier import MedicalClassifier
from crash_logger import CrashLogger
from database import Database
from login_dialog import LoginDialog, UserManagerDialog
from record_manager_dialog import RecordManagerDialog
from voice_command import VoiceCommandParser
from asr_engine import get_microphone_list
from diagnosis_assistant import DiagnosisAssistant
from knowledge_qa import KnowledgeQA
from license_manager import LicenseManager
from activation_dialog import ActivationDialog, TrialInfoBar
from phrase_library import PhraseLibrary
from phrase_dialog import PhraseDialog
from correction_feedback import CorrectionFeedback
from ux_components import Toast, RecordingIndicator, OnboardingGuide, FieldHighlighter
from audio_widgets import WaveformWidget, AudioPlayerWidget, AsrPreviewPanel
from diff_review_dialog import DiffReviewDialog
from correction_memory import get_memory
from topk_engine import get_topk_engine
from recording_handler import RecordingHandler
from threads import ListenThread, CorrectThread, DiagnosisThread, create_listen_thread, stop_listen_thread

# 对话框（从 gui 包导入）
from gui.dialogs import RuleManagerDialog, TemplateManagerDialog, SectionDialog, FieldWordsPanel


class MedVoiceApp(QMainWindow):
    # ─── 常量 ─────────────────────────────────────────
    WIN_X, WIN_Y, WIN_W, WIN_H = 100, 100, 1200, 800
    DURATION_TIMER_MS = 500
    MAX_RECORDS_DISPLAY = 200
    TERM_BUDGET = 300
    POSTPROCESS_BUDGET = 120
    TOPK_FAST = 12

    file_transcribed = pyqtSignal(str, str)
    _maintenance_message = pyqtSignal(str)  # 后台维护线程的安全消息通道

    def __init__(self, db=None, current_user=None):
        super().__init__()
        self.db = db
        self.current_user = current_user or {}
        self.current_record_id = None
        uname = self.current_user.get("username", "")
        title = "星衍AI智能病历录入系统 v2.0"
        if uname:
            role_txt = "管理员" if self.current_user.get("role") == "admin" else "医生"
            title += "  |  当前用户：%s（%s）" % (uname, role_txt)
        self.setWindowTitle(title)
        self.setGeometry(self.WIN_X, self.WIN_Y, self.WIN_W, self.WIN_H)

        # 核心引擎
        self.rule_engine = RuleEngine()
        self.corrector = Corrector(rule_engine=self.rule_engine)
        self.template_engine = TemplateEngine()
        self.parser = SectionParser()
        self.smart_dictation = SmartDictation(self.parser)
        self.classifier = MedicalClassifier()

        # AI 辅助诊断
        self.diagnosis_assistant = DiagnosisAssistant()
        self.qa_engine = KnowledgeQA()
        self.diagnosis_thread = None

        # 常用语句库
        self.phrase_lib = PhraseLibrary()
        self._phrase_dialog = None

        # 纠错反馈收集
        self.feedback = CorrectionFeedback()
        self.memory = None
        self.topk_engine = None

        # 崩溃日志
        self.crash_logger = CrashLogger()
        self.crash_logger.log_event("应用启动")

        # 加载语音模型
        model_path = os.path.join(os.path.dirname(__file__), "model")
        self.asr = ASREngine(model_path=model_path)
        self.listen_thread = None

        # 状态
        self.current_dept = "通用"
        self.asr.set_hotwords("通用")
        self.is_listening = False
        self.partial_text = ""
        self._auto_stop_timer = None
        self.voice_command = VoiceCommandParser()
        self._record_start_ts = None
        self._duration_timer = QTimer(self)
        self._duration_timer.setInterval(self.DURATION_TIMER_MS)
        self._duration_timer.timeout.connect(self._update_record_duration)

        self._init_ui()
        self._apply_dark_theme()

        # 默认加载通用词库
        self.corrector.set_department("通用")
        self._load_departments()

        # 光标移动时自动检测当前字段
        self.text_edit.cursorPositionChanged.connect(self._on_cursor_moved)

        # 启动时：自动备份数据库 + 从历史病历预热用户热词
        self._startup_maintenance_thread = threading.Thread(
            target=self._startup_maintenance,
            daemon=True,
            name="StartupMaintenance"
        )
        self._startup_maintenance_thread.start()

        # 首次使用新手引导
        self._guide = OnboardingGuide.show_if_needed(self)

        # 音频文件转写完成信号
        self.file_transcribed.connect(self._on_file_transcribed)
        self._maintenance_message.connect(lambda msg: self.status_bar.showMessage(msg))

        # 转写线程并发控制（最多 2 个同时转写）
        self._transcribe_semaphore = threading.Semaphore(2)

        # 模板搜索过滤
        self.template_combo.lineEdit().textEdited.connect(self._filter_templates)

    # ─── 以下方法保持原有逻辑，仅调整内部引用 ───

    def _startup_maintenance(self):
        if not self.db:
            return
        try:
            path = self.db.auto_backup()
            self.crash_logger.log_event("数据库自动备份: %s" % path)
        except Exception as e:
            self._maintenance_message.emit(f"⚠ 自动备份失败: {e}")
        try:
            self._refresh_user_hotwords(silent=True)
        except Exception as e:
            self._maintenance_message.emit(f"⚠ 用户热词预热失败: {e}")

    def _refresh_user_hotwords(self, silent=False):
        if not self.db:
            return
        user_id = None if self.current_user.get("role") == "admin" else self.current_user.get("id")
        records = self.db.list_records(user_id=user_id)
        if not records:
            if not silent:
                self.status_bar.showMessage("暂无历史病历，无法提取个人热词")
            return
        from collections import Counter
        counter = Counter()
        for r in records[:self.MAX_RECORDS_DISPLAY]:
            content = r.get("content", "") or ""
            for w in self.corrector.active_words:
                if len(w) >= 2 and w in content:
                    counter[w] += 1
        common = [w for w, c in counter.most_common(self.MAX_RECORDS_DISPLAY) if c >= 2]
        if common:
            self.asr.update_user_hotwords(common)
            self.asr.set_hotwords(self.current_dept)
            self._refresh_topk_hotwords(silent=True)
            if not silent:
                self.status_bar.showMessage("已从 %d 份病历提取 %d 个个人高频词并加入热词" % (len(records), len(common)))
        elif not silent:
            self.status_bar.showMessage("未提取到足够的高频词")

    def _refresh_topk_hotwords(self, silent=False):
        try:
            topk = self._get_topk_engine()
            if topk:
                topk.refresh_asr_hotwords(
                    self.asr,
                    dept=self.current_dept,
                    doctor_id=self.current_user.get("id") if isinstance(self.current_user, dict) else None,
                    term_budget=self.TERM_BUDGET,
                    postprocess_budget=self.POSTPROCESS_BUDGET,
                )
                self.asr.set_hotwords(self.current_dept)
                if not silent:
                    self.status_bar.showMessage("已基于记忆库刷新 Top-K 术语热词")
        except Exception as e:
            print(f"[Main] 刷新 Top-K 热词失败: {e}")

    def _on_cursor_moved(self):
        cursor = self.text_edit.textCursor()
        pos = cursor.position()

        if not hasattr(self, '_last_cursor_line'):
            self._last_cursor_line = -1
            self._last_detected_field = None

        block = cursor.block()
        current_line = block.blockNumber()

        if current_line == self._last_cursor_line:
            field = self._last_detected_field
        else:
            text = self.text_edit.toPlainText()
            field = self._detect_field_at_position(text, pos)
            self._last_cursor_line = current_line
            self._last_detected_field = field

        if field and field != self.field_panel._current_field:
            self.field_panel.set_current_field(field)

    def _detect_field_at_position(self, text, pos):
        best_field = None
        best_pos = -1
        for keyword, standard_field in self.parser.keyword_to_field.items():
            pattern = re.compile(re.escape(keyword) + r'[：: \t]*')
            for m in pattern.finditer(text, 0, pos):
                if m.end() > best_pos:
                    best_pos = m.end()
                    best_field = standard_field
        return best_field

    def _insert_term_at_cursor(self, term):
        cursor = self.text_edit.textCursor()
        cursor.insertText(term)
        self.text_edit.setTextCursor(cursor)
        self.text_edit.setFocus()

    # ══════════════════════════════════════════════════════
    # UI 初始化
    # ══════════════════════════════════════════════════════

    def _init_ui(self):
        toolbar = QToolBar()
        self.addToolBar(toolbar)

        # 科室选择
        toolbar.addWidget(QLabel("  科室："))
        self.dept_combo = QComboBox()
        self.dept_combo.setMinimumWidth(120)
        self.dept_combo.currentTextChanged.connect(self._on_dept_changed)
        toolbar.addWidget(self.dept_combo)

        toolbar.addSeparator()

        # 模板选择
        toolbar.addWidget(QLabel("  模板："))
        self.template_combo = QComboBox()
        self.template_combo.setMinimumWidth(160)
        self.template_combo.setEditable(True)
        self.template_combo.setInsertPolicy(QComboBox.NoInsert)
        self.template_combo.lineEdit().setPlaceholderText("搜索模板...")
        self.template_combo.currentTextChanged.connect(self._on_template_changed)
        toolbar.addWidget(self.template_combo)

        toolbar.addSeparator()

        # 录音按钮
        self.record_btn = QPushButton("\U0001f3a4 开始录音")
        self.record_btn.setCheckable(True)
        self.record_btn.clicked.connect(self._toggle_recording)
        self.record_btn.setStyleSheet("""
            QPushButton {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #00d4ff, stop:1 #0066ff);
                color: #0a0e27;
                font-weight: bold;
                padding: 8px 20px;
                border-radius: 20px;
                font-size: 14px;
            }
            QPushButton:checked {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #ff4444, stop:1 #cc0000);
                color: white;
            }
            QPushButton:hover {
                padding: 8px 25px;
            }
        """)
        toolbar.addWidget(self.record_btn)

        # 录音状态指示器
        self.recording_indicator = RecordingIndicator()
        toolbar.addWidget(self.recording_indicator)

        # 录音模式选择
        mode_label = QLabel("  模式：")
        toolbar.addWidget(mode_label)
        self.record_mode_combo = QComboBox()
        self.record_mode_combo.addItems(["手动停止", "连续录音（60s）", "连续录音（120s）"])
        self.record_mode_combo.setToolTip("手动停止：点击按钮开始/结束\n连续录音：到达时长自动停止")
        self.record_mode_combo.setMinimumWidth(140)
        toolbar.addWidget(self.record_mode_combo)

        # 麦克风设备选择
        mic_label = QLabel("  麦克风：")
        toolbar.addWidget(mic_label)
        self.mic_combo = QComboBox()
        self.mic_combo.setMinimumWidth(150)
        self.mic_combo.setToolTip("选择录音输入设备")
        self._load_microphones()
        self.mic_combo.currentIndexChanged.connect(self._on_mic_changed)
        toolbar.addWidget(self.mic_combo)

        toolbar.addSeparator()

        # 功能按钮
        correct_btn = QPushButton("✨ 纠错")
        correct_btn.clicked.connect(self._run_correction)
        toolbar.addWidget(correct_btn)

        clear_btn = QPushButton("\U0001f5d1 清除")
        clear_btn.clicked.connect(self._clear_text)
        toolbar.addWidget(clear_btn)

        save_btn = QPushButton("\U0001f4be 导出")
        save_btn.clicked.connect(self._save_text)
        toolbar.addWidget(save_btn)

        copy_btn = QPushButton("\U0001f4cb 复制全文")
        copy_btn.setToolTip("复制病历全文到剪贴板（Ctrl+Shift+C）")
        copy_btn.clicked.connect(self._copy_all_text)
        toolbar.addWidget(copy_btn)

        toolbar.addSeparator()

        phrase_btn = QPushButton("\U0001f4ac 常用语")
        phrase_btn.setToolTip("打开常用语句库（F3）")
        phrase_btn.clicked.connect(self._open_phrase_library)
        toolbar.addWidget(phrase_btn)

        autofill_btn = QPushButton("\U0001f4cb 首页→病程")
        autofill_btn.setToolTip("将当前入院记录的内容自动填入首次病程记录")
        autofill_btn.clicked.connect(self._autofill_progress_note)
        toolbar.addWidget(autofill_btn)

        apply_btn = QPushButton("⚡ 一键套用")
        apply_btn.setToolTip("语音输入核心信息，自动替换模板中的占位符 X")
        apply_btn.clicked.connect(self._smart_apply_template)
        toolbar.addWidget(apply_btn)

        # "更多"溢出菜单
        more_btn = QToolButton()
        more_btn.setText("更多 ⌄")
        more_btn.setPopupMode(QToolButton.InstantPopup)
        more_menu = QMenu(more_btn)
        more_menu.addAction("\U0001f4dd 模板管理", self._open_template_manager)
        more_menu.addAction("\U0001f4cf 规则管理", self._open_rule_manager)
        more_menu.addAction("\U0001f4cb 结构化解析", self._open_struct_view)
        more_menu.addSeparator()
        more_menu.addAction("\U0001f9e0 重训语言模型", self._retrain_lm)
        more_menu.addAction("\U0001f4cb 崩溃日志", self._view_crash_log)
        more_btn.setMenu(more_menu)
        more_btn.setStyleSheet("""
            QToolButton { padding: 6px 12px; border-radius: 6px; }
            QToolButton:hover { background: rgba(0,212,255,0.1); }
        """)
        toolbar.addWidget(more_btn)

        toolbar.addSeparator()

        save_record_btn = QPushButton("\U0001f4be 保存病历")
        save_record_btn.setToolTip("保存到病历库（Ctrl+S）")
        save_record_btn.clicked.connect(self._save_record)
        toolbar.addWidget(save_record_btn)

        record_lib_btn = QPushButton("\U0001f4da 病历库")
        record_lib_btn.clicked.connect(self._open_record_manager)
        toolbar.addWidget(record_lib_btn)

        backup_btn = QPushButton("\U0001f6e1 备份")
        backup_btn.setToolTip("备份/恢复病历数据库")
        backup_btn.clicked.connect(self._open_backup_menu)
        toolbar.addWidget(backup_btn)

        if self.current_user.get("role") == "admin":
            user_mgr_btn = QPushButton("\U0001f465 用户管理")
            user_mgr_btn.clicked.connect(self._open_user_manager)
            toolbar.addWidget(user_mgr_btn)

        # 快捷键
        shortcuts = [
            ("Ctrl+S", self._save_record),
            ("Ctrl+Shift+C", self._copy_all_text),
            ("F2", self._toggle_recording),
            ("Ctrl+R", self._toggle_recording),
            ("Ctrl+E", self._save_text),
            ("F11", self._toggle_focus_mode),
            ("F9", self._toggle_left_panel),
            ("F3", self._open_phrase_library),
            ("F4", self._run_correction),
        ]
        for key, callback in shortcuts:
            action = QAction(self)
            action.setShortcut(key)
            action.triggered.connect(callback)
            self.addAction(action)

        # 中央部件
        central = QWidget()
        self.setCentralWidget(central)
        main_layout = QHBoxLayout(central)
        main_layout.setContentsMargins(10, 10, 10, 10)
        main_layout.setSpacing(10)

        splitter = QSplitter(Qt.Horizontal)
        self.splitter = splitter

        # 左侧面板 - 纠错日志
        left_panel = QWidget()
        self.left_panel = left_panel
        left_layout = QVBoxLayout(left_panel)
        left_layout.setContentsMargins(0, 0, 0, 0)

        left_layout.addWidget(QLabel("\U0001f4cb 纠错日志"))

        filter_bar = QWidget()
        filter_layout = QHBoxLayout(filter_bar)
        filter_layout.setContentsMargins(0, 0, 0, 5)
        filter_layout.setSpacing(5)

        self.filter_typo = QCheckBox("\U0001f4e2 错别字")
        self.filter_logic = QCheckBox("\U0001f9e0 逻辑错误")
        self.filter_missing = QCheckBox("⚠️ 缺项提醒")
        self.filter_typo.setChecked(True)
        self.filter_logic.setChecked(True)
        self.filter_missing.setChecked(True)

        for cb in [self.filter_typo, self.filter_logic, self.filter_missing]:
            cb.setStyleSheet("""
                QCheckBox { color: #b8c5d6; font-size: 11px; spacing: 3px; }
                QCheckBox::indicator {
                    width: 14px; height: 14px;
                    border: 1px solid rgba(0,212,255,0.3);
                    border-radius: 3px;
                    background: rgba(0,212,255,0.05);
                }
                QCheckBox::indicator:checked {
                    background: rgba(0,212,255,0.3);
                    border-color: #00d4ff;
                }
            """)
            cb.stateChanged.connect(self._apply_filter)
            filter_layout.addWidget(cb)

        filter_layout.addStretch()
        self.stats_label = QLabel("")
        self.stats_label.setStyleSheet("color: #6b8a9a; font-size: 10px;")
        filter_layout.addWidget(self.stats_label)

        left_layout.addWidget(filter_bar)

        self.log_list = QListWidget()
        self.log_list.setStyleSheet("""
            QListWidget {
                background: rgba(255,255,255,0.03);
                border: 1px solid rgba(0,212,255,0.1);
                border-radius: 8px;
                padding: 5px;
                font-size: 12px;
            }
            QListWidget::item { padding: 6px; border-bottom: 1px solid rgba(0,212,255,0.05); }
        """)
        self.log_list.itemClicked.connect(self._on_log_item_clicked)
        left_layout.addWidget(self.log_list)

        action_bar = QWidget()
        action_layout = QHBoxLayout(action_bar)
        action_layout.setContentsMargins(0, 5, 0, 0)
        action_layout.setSpacing(6)

        accept_btn = QPushButton("✓ 接受")
        accept_btn.clicked.connect(self._accept_correction)
        accept_btn.setEnabled(False)
        accept_btn.setStyleSheet("""
            QPushButton {
                background: rgba(81, 207, 102, 0.15); color: #51cf66;
                padding: 5px 14px; border-radius: 12px;
                border: 1px solid rgba(81, 207, 102, 0.3); font-size: 11px;
            }
            QPushButton:hover { background: rgba(81, 207, 102, 0.25); }
            QPushButton:disabled { color: #444; border-color: #333; }
        """)

        accept_all_btn = QPushButton("✓✓ 全部接受")
        accept_all_btn.clicked.connect(self._accept_all_corrections)
        accept_all_btn.setStyleSheet("""
            QPushButton {
                background: rgba(81, 207, 102, 0.1); color: #51cf66;
                padding: 5px 10px; border-radius: 12px;
                border: 1px solid rgba(81, 207, 102, 0.2); font-size: 11px;
            }
            QPushButton:hover { background: rgba(81, 207, 102, 0.2); }
        """)

        reject_btn = QPushButton("✗ 拒绝")
        reject_btn.clicked.connect(self._reject_correction)
        reject_btn.setEnabled(False)
        reject_btn.setStyleSheet("""
            QPushButton {
                background: rgba(255, 107, 107, 0.15); color: #ff6b6b;
                padding: 5px 14px; border-radius: 12px;
                border: 1px solid rgba(255, 107, 107, 0.3); font-size: 11px;
            }
            QPushButton:hover { background: rgba(255, 107, 107, 0.25); }
            QPushButton:disabled { color: #444; border-color: #333; }
        """)

        self.log_hint = QLabel("点击日志条目查看详情")
        self.log_hint.setStyleSheet("color: #6b8a9a; font-size: 10px;")

        action_layout.addWidget(accept_btn)
        action_layout.addWidget(accept_all_btn)
        action_layout.addWidget(reject_btn)
        action_layout.addStretch()
        action_layout.addWidget(self.log_hint)

        left_layout.addWidget(action_bar)
        self._accept_btn = accept_btn
        self._reject_btn = reject_btn

        self._all_logs = []
        splitter.addWidget(left_panel)

        # 右侧 - 文本编辑区
        right_panel = QWidget()
        right_layout = QVBoxLayout(right_panel)
        right_layout.setContentsMargins(0, 0, 0, 0)

        # 实时识别显示
        self.partial_label = QLabel("等待输入...")
        self.partial_label.setStyleSheet("""
            color: #00d4ff; font-size: 13px; padding: 5px;
            background: rgba(0,212,255,0.05); border-radius: 5px;
        """)
        self.partial_label.setWordWrap(True)
        self.partial_label.setMaximumHeight(120)
        right_layout.addWidget(self.partial_label)

        # 悬浮识别预览面板
        self.asr_preview = AsrPreviewPanel(self)
        self.asr_preview.accepted.connect(self._on_preview_accept)
        self.asr_preview.rejected.connect(self._on_preview_reject)
        self.asr_preview.retried.connect(self._on_preview_retry)
        self._pending_asr_text = ''

        # 录音事件处理器
        self.recorder = RecordingHandler(self)

        # 实时录音波形图
        self.waveform = WaveformWidget()
        self.waveform.setVisible(False)
        right_layout.addWidget(self.waveform)

        self._wave_timer = QTimer(self)
        self._wave_timer.setInterval(33)
        self._wave_timer.timeout.connect(self._poll_audio_level)

        # 字段常用词面板
        self.field_panel = FieldWordsPanel()
        self.field_panel.setMaximumHeight(340)
        self.field_panel.term_clicked.connect(self._insert_term_at_cursor)
        right_layout.addWidget(self.field_panel)

        # 文本编辑区
        self.text_edit = QTextEdit()
        self.text_edit.setPlaceholderText(
            "选择模板开始，或直接输入病历内容。\n"
            "点击「开始录音」用语音输入，系统会自动纠错。"
        )
        self.text_edit.setStyleSheet("""
            QTextEdit {
                background: rgba(255,255,255,0.03);
                border: 1px solid rgba(0,212,255,0.15);
                border-radius: 10px; padding: 15px;
                font-size: 14px; line-height: 1.8; color: #e0e0e0;
            }
            QTextEdit:focus { border: 1px solid rgba(0,212,255,0.4); }
        """)
        right_layout.addWidget(self.text_edit)

        # 音文对照播放器
        self.audio_player_widget = AudioPlayerWidget()
        right_layout.addWidget(self.audio_player_widget)

        # 字段名语法高亮
        self.field_highlighter = FieldHighlighter(self.text_edit.document())

        # 支持拖拽音频文件转写
        self.setAcceptDrops(True)

        # AI 辅助诊断面板
        ai_group = QGroupBox("\U0001f52c AI 辅助诊断")
        self.ai_group = ai_group
        ai_group.setStyleSheet("""
            QGroupBox {
                color: #00d4ff; font-size: 13px; font-weight: bold;
                border: 1px solid rgba(0,212,255,0.2);
                border-radius: 10px; margin-top: 10px; padding-top: 8px;
            }
            QGroupBox::title { subcontrol-origin: margin; left: 12px; padding: 0 5px; }
        """)
        ai_layout = QVBoxLayout(ai_group)
        ai_layout.setContentsMargins(10, 8, 10, 10)

        ai_header = QHBoxLayout()
        self.ai_analyze_btn = QPushButton("\U0001f52c 分析当前病历")
        self.ai_analyze_btn.clicked.connect(self._run_diagnosis)
        self.ai_analyze_btn.setStyleSheet("""
            QPushButton {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #00d4ff, stop:1 #0088cc);
                color: #0a0e27; padding: 6px 16px;
                border-radius: 12px; font-size: 12px; font-weight: bold;
            }
            QPushButton:hover { background: #00d4ff; }
            QPushButton:disabled { background: #333; color: #666; }
        """)
        qa_btn = QPushButton("\U0001f4a1 知识问答")
        qa_btn.clicked.connect(self._show_qa_dialog)
        qa_btn.setStyleSheet("""
            QPushButton {
                background: qlineargradient(x1:0,y1:0,x2:1,y2:0, stop:0 #ff9a5b, stop:1 #f64236);
                color: #fff; padding: 6px 16px;
                border-radius: 12px; font-size: 12px; font-weight: bold;
            }
            QPushButton:hover { background: #ffaf7b; }
        """)
        self.ai_status_label = QLabel("")
        self.ai_status_label.setStyleSheet("color: #6b8a9a; font-size: 10px;")
        ai_header.addWidget(self.ai_analyze_btn)
        ai_header.addWidget(qa_btn)
        ai_header.addWidget(self.ai_status_label)
        ai_header.addStretch()

        self.ai_collapse_btn = QPushButton("▲ 收起")
        self.ai_collapse_btn.setFixedWidth(64)
        self.ai_collapse_btn.clicked.connect(self._toggle_ai_panel)
        self.ai_collapse_btn.setStyleSheet("""
            QPushButton {
                background: transparent; color: #6b8a9a;
                border: 1px solid rgba(0,212,255,0.2);
                border-radius: 8px; padding: 3px 8px; font-size: 11px;
            }
            QPushButton:hover { color: #00d4ff; border-color: rgba(0,212,255,0.5); }
        """)
        ai_header.addWidget(self.ai_collapse_btn)
        ai_layout.addLayout(ai_header)

        self.ai_result = QTextBrowser()
        self.ai_result.setOpenExternalLinks(False)
        self.ai_result.setStyleSheet("""
            QTextBrowser {
                background: rgba(255,255,255,0.03);
                border: 1px solid rgba(0,212,255,0.15);
                border-radius: 8px; padding: 10px;
                font-size: 13px; color: #e0e0e0;
            }
        """)
        self.ai_result.setPlaceholderText("点击「分析当前病历」，AI 将基于知识图谱给出可能诊断、用药审查、检查建议与风险预警。")
        self.ai_result.setMinimumHeight(160)
        ai_layout.addWidget(self.ai_result)

        right_layout.addWidget(ai_group)
        splitter.addWidget(right_panel)
        splitter.setSizes([300, 900])

        main_layout.addWidget(splitter)

        # 状态栏
        self.status_bar = QStatusBar()
        self.setStatusBar(self.status_bar)
        self.status_bar.showMessage("就绪 | 请先选择科室和模板")

        # 试用期信息
        self._license_mgr = LicenseManager()
        trial_label = TrialInfoBar.create_label(self._license_mgr)
        if trial_label.text():
            self.status_bar.addPermanentWidget(trial_label)

    def _apply_dark_theme(self):
        app = QApplication.instance()
        app.setStyle("Fusion")
        palette = QPalette()
        palette.setColor(QPalette.Window, QColor(10, 14, 39))
        palette.setColor(QPalette.WindowText, QColor(255, 255, 255))
        palette.setColor(QPalette.Base, QColor(15, 20, 50))
        palette.setColor(QPalette.AlternateBase, QColor(20, 25, 60))
        palette.setColor(QPalette.ToolTipBase, QColor(255, 255, 255))
        palette.setColor(QPalette.ToolTipText, QColor(255, 255, 255))
        palette.setColor(QPalette.Text, QColor(220, 220, 220))
        palette.setColor(QPalette.Button, QColor(30, 40, 80))
        palette.setColor(QPalette.ButtonText, QColor(255, 255, 255))
        palette.setColor(QPalette.BrightText, QColor(255, 0, 0))
        palette.setColor(QPalette.Link, QColor(0, 212, 255))
        palette.setColor(QPalette.Highlight, QColor(0, 212, 255))
        palette.setColor(QPalette.HighlightedText, QColor(10, 14, 39))
        app.setPalette(palette)

    # ══════════════════════════════════════════════════════
    # 科室 / 模板
    # ══════════════════════════════════════════════════════

    def _load_departments(self):
        depts = self.template_engine.get_departments()
        self.dept_combo.clear()
        self.dept_combo.addItems(depts)

    def _on_dept_changed(self, dept):
        self.current_dept = dept
        self.corrector.set_department(dept)
        self.asr.set_hotwords(dept)
        self._refresh_topk_hotwords(silent=True)

        self.template_combo.blockSignals(True)
        self.template_combo.clear()
        templates = self.template_engine.get_templates(dept)
        self._all_template_names = [t["name"] for t in templates]
        for name in self._all_template_names:
            self.template_combo.addItem(name)
        self.template_combo.blockSignals(False)

        self.status_bar.showMessage(f"当前科室：{dept} | 词库已更新")

    def _filter_templates(self, search_text):
        self.template_combo.blockSignals(True)
        self.template_combo.clear()
        if not search_text.strip():
            for name in getattr(self, '_all_template_names', []):
                self.template_combo.addItem(name)
        else:
            keyword = search_text.strip().lower()
            for name in getattr(self, '_all_template_names', []):
                if keyword in name.lower():
                    self.template_combo.addItem(name)
        self.template_combo.blockSignals(False)
        if self.template_combo.count() == 1:
            self.template_combo.setCurrentIndex(0)

    def _on_template_changed(self, template_name):
        if not template_name:
            return
        content = self.template_engine.get_template(self.current_dept, template_name)
        if content:
            current_text = self.text_edit.toPlainText().strip()
            if current_text:
                reply = QMessageBox.question(
                    self, "加载模板",
                    "当前编辑器已有内容，是否用新模板覆盖？",
                    QMessageBox.Yes | QMessageBox.No,
                    QMessageBox.No
                )
                if reply == QMessageBox.No:
                    return
            self.text_edit.setPlainText(content)
            self.status_bar.showMessage(f"已加载模板：{template_name} | 可以开始语音输入")

    def _recommend_template(self, text):
        try:
            if not text or len(text.strip()) < 4:
                return
            best = None
            for dept in self.template_engine.get_departments():
                for t in self.template_engine.get_templates(dept):
                    name = t.get("name", "")
                    if not name:
                        continue
                    score = 0
                    if name in text:
                        score += 3
                    core = name.split("-")[0].replace("【中医】", "")
                    if core and core in text and core != name:
                        score += 2
                    if score > 0 and (best is None or score > best[0]):
                        best = (score, dept, name)
            if best and best[0] >= 2:
                _, dept, name = best
                if self.template_combo.currentText() != name:
                    self.status_bar.showMessage(
                        "\U0001f4a1 推荐模板：%s - %s（可在上方模板下拉选择）" % (dept, name)
                    )
        except Exception as e:
            print(f"[UI] 模板推荐失败: {e}")

    def _autofill_progress_note(self):
        current_text = self.text_edit.toPlainText().strip()
        if not current_text:
            QMessageBox.warning(self, "提示", "当前编辑器无内容，请先填写入院记录。")
            return

        sections = self.parser.parse(current_text)
        if not sections:
            QMessageBox.warning(self, "提示", "未能从当前文本中解析出字段。")
            return

        case_features = []
        basic_info = []
        for field in ['性别', '年龄', '民族']:
            if field in sections and sections[field].strip():
                basic_info.append(f"{field}：{sections[field].strip()}")
        if basic_info:
            case_features.append("，".join(basic_info))

        if '主诉' in sections and sections['主诉'].strip():
            case_features.append(f"主诉：{sections['主诉'].strip()}")
        if '现病史' in sections and sections['现病史'].strip():
            case_features.append(f"现病史：{sections['现病史'].strip()}")
        if '体格检查' in sections and sections['体格检查'].strip():
            case_features.append(f"体格检查：{sections['体格检查'].strip()}")
        if '辅助检查' in sections and sections['辅助检查'].strip():
            case_features.append(f"辅助检查：{sections['辅助检查'].strip()}")

        from datetime import datetime
        now_str = datetime.now().strftime("%Y-%m-%d %H:%M")

        progress_content = f"""日期/时间：{now_str}
病例特点：
{chr(10).join(case_features) if case_features else ''}
诊断依据：
{sections.get('诊断依据', '').strip()}
鉴别诊断：
{sections.get('鉴别诊断', '').strip()}
入院诊断：
{sections.get('初步诊断', sections.get('入院诊断', '')).strip()}
诊疗计划：
{sections.get('诊疗计划', '').strip()}
主治医师意见：
"""

        reply = QMessageBox.question(
            self, "首页→病程",
            f"已从当前文本提取到 {len(sections)} 个字段，是否生成首次病程记录？",
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.Yes
        )
        if reply == QMessageBox.Yes:
            self.text_edit.setPlainText(progress_content)
            self.status_bar.showMessage("已生成首次病程记录")

    def _smart_apply_template(self):
        import re
        current_text = self.text_edit.toPlainText().strip()
        if not current_text:
            QMessageBox.warning(self, "提示", "请先选择常见病模板。")
            return

        x_count = len(re.findall(r'X+', current_text))
        if x_count == 0:
            QMessageBox.information(self, "提示", "当前模板中没有发现占位符 X，无需替换。")
            return

        dialog = QDialog(self)
        dialog.setWindowTitle("一键套用 - 填写患者核心信息")
        dialog.setMinimumWidth(500)
        dialog.setStyleSheet("""
            QDialog { background: #2b2b2b; color: #e0e0e0; }
            QLabel { color: #e0e0e0; font-size: 13px; }
            QLineEdit {
                background: #3c3c3c; color: #e0e0e0; border: 1px solid #555;
                padding: 6px; border-radius: 4px; font-size: 13px;
            }
            QLineEdit:focus { border: 1px solid #4a9eff; }
            QPushButton {
                background: #4a9eff; color: white; border: none;
                padding: 8px 20px; border-radius: 4px; font-size: 13px;
            }
            QPushButton:hover { background: #3a8eef; }
            QPushButton#cancelBtn { background: #555; }
            QPushButton#cancelBtn:hover { background: #666; }
        """)

        layout = QVBoxLayout(dialog)
        info_label = QLabel(f"当前模板中有 {x_count} 处占位符需要填写。")
        info_label.setStyleSheet("color: #4a9eff; font-size: 14px; font-weight: bold;")
        layout.addWidget(info_label)

        hint_label = QLabel("请填写患者核心信息（留空则保留原占位符）：")
        layout.addWidget(hint_label)

        fields = [
            ("姓名", ""), ("性别", ""), ("年龄", ""),
            ("病程时间（如：3天、2周）", ""),
            ("血压（如：160/100）", ""),
            ("体温（如：38.5）", ""),
            ("心率（如：80）", ""),
            ("白细胞（如：12）", ""),
            ("左右侧（左/右）", ""),
            ("部位（如：股骨、胫骨）", ""),
        ]

        line_edits = {}
        for label_text, default in fields:
            row = QHBoxLayout()
            label = QLabel(label_text + "：")
            label.setFixedWidth(180)
            edit = QLineEdit(default)
            edit.setPlaceholderText("留空则保留原占位符")
            row.addWidget(label)
            row.addWidget(edit)
            layout.addLayout(row)
            line_edits[label_text] = edit

        btn_layout = QHBoxLayout()
        ok_btn = QPushButton("✅ 替换")
        cancel_btn = QPushButton("取消")
        cancel_btn.setObjectName("cancelBtn")
        btn_layout.addStretch()
        btn_layout.addWidget(ok_btn)
        btn_layout.addWidget(cancel_btn)
        layout.addLayout(btn_layout)

        ok_btn.clicked.connect(dialog.accept)
        cancel_btn.clicked.connect(dialog.reject)

        if dialog.exec_() != QDialog.Accepted:
            return

        values = {}
        for label_text, edit in line_edits.items():
            val = edit.text().strip()
            if val:
                values[label_text] = val

        if not values:
            QMessageBox.information(self, "提示", "未填写任何信息，模板未修改。")
            return

        result = current_text

        if '姓名' in values:
            result = re.sub(r'姓名：\s*', f'姓名：{values["姓名"]}  ', result)
        if '性别' in values:
            result = re.sub(r'性别：\s*', f'性别：{values["性别"]}  ', result)
        if '年龄' in values:
            result = re.sub(r'年龄：\s*', f'年龄：{values["年龄"]}  ', result)

        time_val = values.get('病程时间', '')
        bp_val = values.get('血压', '')
        temp_val = values.get('体温', '')
        hr_val = values.get('心率', '')
        wbc_val = values.get('白细胞', '')
        side_val = values.get('左右侧', '')
        part_val = values.get('部位', '')

        time_match = re.match(r'(\d+)(\S*)', time_val) if time_val else None
        time_num = time_match.group(1) if time_match else ''
        time_unit = time_match.group(2) if time_match else ''

        if bp_val:
            parts = bp_val.split('/')
            if len(parts) == 2:
                result = re.sub(r'XXX/XXX\s*mmHg', f'{parts[0]}/{parts[1]}mmHg', result)
                result = re.sub(r'XXX/XXX', f'{parts[0]}/{parts[1]}', result)

        if temp_val:
            result = re.sub(r'XX\.X℃', f'{temp_val}℃', result)
            result = re.sub(r'XX-XX℃', f'{temp_val}℃', result)

        if hr_val:
            result = re.sub(r'(?<!\d)XX(?!\.)(?=次/分)', hr_val, result)

        if wbc_val:
            result = re.sub(r'(?<!\d)XX(?!\.)(?=×)', wbc_val, result)

        if side_val:
            result = result.replace('左/右', side_val)
            result = re.sub(r'左/右侧', f'{side_val}侧', result)

        if part_val:
            result = result.replace('侧X', f'侧{part_val}')

        if time_num:
            for unit_suffix in ['年', '月', '天', '小时', '周']:
                result = re.sub(r'(?<=\D)X+' + unit_suffix, f'{time_num}{unit_suffix}', result)

        self.text_edit.setPlainText(result)
        filled_count = sum(1 for v in values.values() if v)
        remaining = len(re.findall(r'X+', result))
        self.status_bar.showMessage(
            f"已替换 {filled_count} 项信息，剩余 {remaining} 处占位符待手动填写"
        )

    def _replace_placeholders_with_voice(self, voice_text, template_text):
        import re
        result = template_text

        cn_map = {'零': '0', '一': '1', '二': '2', '三': '3', '四': '4',
                  '五': '5', '六': '6', '七': '7', '八': '8', '九': '9', '两': '2'}
        def cn_to_arabic(text):
            m = re.match(r'^([一二三四五六七八九])百([一二三四五六七八九])十([一二三四五六七八九])$', text)
            if m:
                return str(int(cn_map[m.group(1)]) * 100 + int(cn_map[m.group(2)]) * 10 + int(cn_map[m.group(3)]))
            m = re.match(r'^([一二三四五六七八九])百([一二三四五六七八九])十$', text)
            if m:
                return str(int(cn_map[m.group(1)]) * 100 + int(cn_map[m.group(2)]) * 10)
            m = re.match(r'^([一二三四五六七八九])百$', text)
            if m:
                return str(int(cn_map[m.group(1)]) * 100)
            m = re.match(r'^([一二三四五六七八九])十([一二三四五六七八九])$', text)
            if m:
                return str(int(cn_map[m.group(1)]) * 10 + int(cn_map[m.group(2)]))
            m = re.match(r'^十([一二三四五六七八九])$', text)
            if m:
                return str(10 + int(cn_map[m.group(1)]))
            m = re.match(r'^([一二三四五六七八九])十$', text)
            if m:
                return str(int(cn_map[m.group(1)]) * 10)
            if text == '十':
                return '10'
            r = ''
            for ch in text:
                if ch in cn_map:
                    r += cn_map[ch]
                else:
                    return text
            return r or text

        sections = self.parser.parse(voice_text)
        inferred = self.classifier.extract_basic_fields(voice_text)

        name = inferred.get('姓名', sections.get('姓名', ''))
        if name:
            result = re.sub(r'姓名：\s*', f'姓名：{name}  ', result)

        gender = inferred.get('性别', sections.get('性别', ''))
        if gender:
            result = re.sub(r'性别：\s*', f'性别：{gender}  ', result)

        age = inferred.get('年龄', sections.get('年龄', ''))
        if age:
            result = re.sub(r'年龄：\s*', f'年龄：{age}  ', result)

        for unit in ['年', '月', '天', '小时', '周']:
            m = re.search(r'(\d+)\s*' + re.escape(unit), voice_text)
            if m:
                result = re.sub(r'(?<=\D)X+' + re.escape(unit), f'{m.group(1)}{unit}', result)
                continue
            m = re.search(r'([零一二两三四五六七八九十百]+)\s*' + re.escape(unit), voice_text)
            if m:
                result = re.sub(r'(?<=\D)X+' + re.escape(unit), f'{cn_to_arabic(m.group(1))}{unit}', result)

        bp_match = re.search(r'(\d{2,3})\s*[/比]\s*(\d{2,3})', voice_text)
        if bp_match:
            result = re.sub(r'XXX\s*/\s*XXX\s*mmHg', f'{bp_match.group(1)}/{bp_match.group(2)}mmHg', result)
            result = re.sub(r'XXX\s*/\s*XXX', f'{bp_match.group(1)}/{bp_match.group(2)}', result)

        temp_match = re.search(r'(\d{2}\.\d)\s*[度℃]', voice_text)
        if not temp_match:
            temp_match = re.search(r'体温\s*(\d{2}\.?\d*)', voice_text)
        if temp_match:
            result = re.sub(r'XX\.X℃', f'{temp_match.group(1)}℃', result)
            result = re.sub(r'XX-XX℃', f'{temp_match.group(1)}℃', result)

        hr_match = re.search(r'(?:心率|脉搏)\s*(\d{2,3})', voice_text)
        if hr_match:
            result = re.sub(r'(?<!\d)XX(?!\.)(?=次/分)', hr_match.group(1), result)

        wbc_match = re.search(r'(?:白细胞|WBC)\s*(\d+\.?\d*)', voice_text)
        if wbc_match:
            result = re.sub(r'(?<!\d)XX(?!\.)(?=×)', wbc_match.group(1), result)

        if re.search(r'左侧|左边', voice_text):
            result = result.replace('左/右', '左')
        elif re.search(r'右侧|右边', voice_text):
            result = result.replace('左/右', '右')

        part_match = re.search(r'[左右]侧(\S{1,4}?)(?:肿痛|疼痛|骨折|肿胀)', voice_text)
        if part_match:
            result = result.replace('侧X', f'侧{part_match.group(1)}')

        return result

    # ══════════════════════════════════════════════════════
    # 录音
    # ══════════════════════════════════════════════════════

    def _toggle_recording(self):
        self.recorder.toggle_recording()

    def moveEvent(self, event):
        super().moveEvent(event)
        panel = getattr(self, 'asr_preview', None)
        if panel is not None and panel.isVisible():
            panel.reposition()

    def resizeEvent(self, event):
        super().resizeEvent(event)
        panel = getattr(self, 'asr_preview', None)
        if panel is not None and panel.isVisible():
            panel.reposition()

    def _start_recording(self):
        self.recorder.start_recording()

    def _stop_recording(self):
        self.recorder.stop_recording()

    def closeEvent(self, event):
        """应用退出时清理资源"""
        if self.is_listening:
            self._stop_recording()
        if hasattr(self, 'asr') and self.asr:
            self.asr.stop_listening()
        super().closeEvent(event)

    def _load_microphones(self):
        self.mic_combo.blockSignals(True)
        self.mic_combo.clear()
        self.mic_combo.addItem("系统默认", None)
        try:
            for dev in get_microphone_list():
                self.mic_combo.addItem(dev["name"], dev["index"])
        except Exception as e:
            print(f"[UI] 加载麦克风列表失败: {e}")
        self.mic_combo.blockSignals(False)

    def _on_mic_changed(self, _index):
        device_index = self.mic_combo.currentData()
        self.asr.set_input_device(device_index)
        self.status_bar.showMessage(f"\U0001f3a4 录音设备：{self.mic_combo.currentText()}")

    def _copy_all_text(self):
        text = self.text_edit.toPlainText()
        if not text.strip():
            self.status_bar.showMessage("没有内容可复制")
            return
        QApplication.clipboard().setText(text)
        self.status_bar.showMessage("\U0001f4cb 已复制全文到剪贴板")
        Toast.show_toast(self, "已复制到剪贴板", "info")

    def _update_record_duration(self):
        if self._record_start_ts is None:
            return
        elapsed = int(time.time() - self._record_start_ts)
        mm, ss = divmod(elapsed, 60)
        self.status_bar.showMessage(f"\U0001f534 录音中... {mm:02d}:{ss:02d}")

    def _poll_audio_level(self):
        level = getattr(self.asr, '_current_level', 0.0)
        self.waveform.add_level(level)

    def _update_field_context(self):
        text = self.text_edit.toPlainText()
        field_keywords = ['主诉', '现病史', '既往史', '个人史', '家族史',
                          '体格检查', '辅助检查', '初步诊断', '诊疗经过',
                          '出院情况', '出院医嘱', '影像表现', '诊断意见']
        last_field = ""
        last_pos = -1
        for kw in field_keywords:
            pos = text.rfind(kw)
            if pos > last_pos:
                last_pos = pos
                last_field = kw
        self.asr.set_field_context(last_field)

    # ══════════════════════════════════════════════════════
    # UI 增强
    # ══════════════════════════════════════════════════════

    def _toggle_left_panel(self):
        self.left_panel.setVisible(not self.left_panel.isVisible())
        state = "已展开" if self.left_panel.isVisible() else "已折叠"
        self.status_bar.showMessage(f"纠错面板{state}")

    def _toggle_ai_panel(self):
        visible = self.ai_result.isVisible()
        self.ai_result.setVisible(not visible)
        self.ai_collapse_btn.setText("▼ 展开" if visible else "▲ 收起")
        if visible:
            self.ai_group.setMaximumHeight(52)
        else:
            self.ai_group.setMaximumHeight(16777215)

    def _toggle_focus_mode(self):
        self._focus_mode = not getattr(self, '_focus_mode', False)
        if self._focus_mode:
            self._pre_focus = {
                'left_visible': self.left_panel.isVisible(),
                'field_visible': self.field_panel.isVisible(),
                'ai_visible': self.ai_group.isVisible(),
                'toolbar_visible': self.findChild(QToolBar).isVisible(),
            }
            self.left_panel.hide()
            self.field_panel.hide()
            self.ai_group.hide()
            self.audio_player_widget.hide()
            self.findChild(QToolBar).hide()
            self.status_bar.showMessage("\U0001f3a7 专注模式：按 F11 退出")
            Toast.show_toast(self, "专注模式 · 按 F11 退出", "info")
        else:
            pf = getattr(self, '_pre_focus', {})
            self.left_panel.setVisible(pf.get('left_visible', True))
            self.field_panel.setVisible(pf.get('field_visible', True))
            self.ai_group.setVisible(pf.get('ai_visible', True))
            self.audio_player_widget.show()
            self.findChild(QToolBar).setVisible(pf.get('toolbar_visible', True))
            self.status_bar.showMessage("已退出专注模式")

    # ══════════════════════════════════════════════════════
    # 拖拽音频文件
    # ══════════════════════════════════════════════════════

    def dragEnterEvent(self, event):
        if event.mimeData().hasUrls():
            for url in event.mimeData().urls():
                path = url.toLocalFile()
                if path.lower().endswith(('.wav', '.mp3', '.m4a', '.flac', '.ogg')):
                    event.acceptProposedAction()
                    return
        event.ignore()

    def dropEvent(self, event):
        if not event.mimeData().hasUrls():
            return
        for url in event.mimeData().urls():
            path = url.toLocalFile()
            if path.lower().endswith(('.wav', '.mp3', '.m4a', '.flac', '.ogg')):
                self._transcribe_audio_file(path)
                event.acceptProposedAction()
                return
        event.ignore()

    def _transcribe_audio_file(self, path):
        if not self.asr.is_ready():
            Toast.show_toast(self, "语音引擎未就绪，无法转写", "warning")
            return
        fname = os.path.basename(path)
        self.status_bar.showMessage(f"\U0001f3a5 正在转写：{fname} ...")
        Toast.show_toast(self, f"开始转写 {fname}", "info", duration=1500)

        def _worker():
            try:
                text = self.asr.transcribe_file(path)
            except Exception as e:
                print(f"[Main] 转写失败: {e}")
                text = ""
            self.file_transcribed.emit(text, fname)
            self._transcribe_semaphore.release()

        # 使用信号量限制并发（最多 2 个转写任务同时进行）
        if not self._transcribe_semaphore.acquire(blocking=False):
            Toast.show_toast(self, "转写任务过多，请稍后再试", "warning")
            return
        threading.Thread(target=_worker, daemon=True).start()

    def _on_file_transcribed(self, text, fname):
        if text:
            cursor = self.text_edit.textCursor()
            cursor.movePosition(QTextCursor.End)
            if self.text_edit.toPlainText().strip():
                cursor.insertText("\n")
            cursor.insertText(text)
            self.text_edit.setTextCursor(cursor)
            self.status_bar.showMessage(f"✅ 转写完成：{fname}（{len(text)} 字）")
            Toast.show_toast(self, f"转写完成（{len(text)} 字）", "success")
        else:
            self.status_bar.showMessage(f"⚠️ 转写无结果：{fname}")
            Toast.show_toast(self, "转写无结果，请检查音频", "warning")

    # ══════════════════════════════════════════════════════
    # ASR 识别结果处理
    # ══════════════════════════════════════════════════════

    def _on_recognized(self, text):
        try:
            print(f"[UI] 收到识别结果，长度: {len(text) if text else 0}")
            self.crash_logger.log_event("ASR识别完成", {"text_length": len(text) if text else 0})
            self.recorder.on_recognized(text)
        except Exception as e:
            print(f"[UI] 处理识别结果时出错: {e}")
            import traceback
            traceback.print_exc()
            self.crash_logger.log_exception(
                type(e), e, e.__traceback__,
                context="处理识别结果"
            )
            self.partial_label.setText("⚠️ 处理识别结果时出错")
            self.status_bar.showMessage(f"错误: {e}")

        self._load_last_audio()

    def _apply_asr_result(self, text):
        template_name = self.template_combo.currentText()
        if template_name and self.current_dept:
            template_content = self.template_engine.get_template(
                self.current_dept, template_name
            )
            if template_content:
                clean_base = getattr(self, '_stream_base_text', '').strip()
                base = clean_base if clean_base else template_content

                import re as _re
                if _re.search(r'X+', base):
                    filled = self._replace_placeholders_with_voice(text, base)
                    self.text_edit.setPlainText(filled)
                    self._last_asr_snapshot = filled
                    remaining = len(_re.findall(r'X+', filled))
                    self.partial_label.setText(f"✓ 已替换占位符，剩余 {remaining} 处")
                    self.status_bar.showMessage(f"套用完成，剩余 {remaining} 处占位符待手动填写")
                    return

                filled = self.classifier.incremental_fill(text, base)
                self.text_edit.setPlainText(filled)
                self._last_asr_snapshot = filled
                self.partial_label.setText("✓ 识别完成")
                self.status_bar.showMessage(f"识别完成，共 {len(filled)} 字")
                return

        current = self.text_edit.toPlainText()
        if current.strip():
            text = current.rstrip() + "\n\n" + text
        self.text_edit.setPlainText(text)
        self._last_asr_snapshot = text
        self.partial_label.setText("✓ 识别完成")
        self.status_bar.showMessage(f"识别完成，共 {len(text)} 字")
        self._recommend_template(self.text_edit.toPlainText())

    def _on_preview_accept(self):
        self.recorder.on_preview_accept()

    def _on_preview_reject(self):
        self.recorder.on_preview_reject()

    def _on_preview_retry(self):
        self.recorder.on_preview_retry()

    def _load_last_audio(self):
        audio_path = getattr(self.asr, 'last_audio_path', None)
        if audio_path and os.path.exists(audio_path):
            self.audio_player_widget.load(audio_path)

    def _on_partial(self, text):
        self.recorder.on_partial(text)

    def _on_asr_stream_error(self, msg):
        """ASR 流式识别恢复状态回调（来自 StreamRecognizer）"""
        if not msg:
            return
        # 更新状态栏
        self.status_bar.showMessage(f"🎤 {msg}")
        # 严重错误弹 Toast
        if "连续失败" in msg or "熔断" in msg:
            Toast.show_toast(self, msg, "warning", duration=4000)
        elif "恢复" in msg:
            Toast.show_toast(self, msg, "success", duration=2000)
        else:
            Toast.show_toast(self, msg, "info", duration=3000)

    # ══════════════════════════════════════════════════════
    # 纠错
    # ══════════════════════════════════════════════════════

    def _run_correction(self):
        text = self.text_edit.toPlainText().strip()
        if not text:
            QMessageBox.information(self, "提示", "请先输入或录制文本")
            return

        self.status_bar.showMessage("正在纠错...")
        QApplication.processEvents()

        self._last_correction_original = text
        self.correct_thread = CorrectThread(self.corrector, text)
        self.correct_thread.correction_done.connect(self._on_correction_done)
        self.correct_thread.start()

    def _on_correction_done(self, corrected, log):
        original_text = getattr(self, '_last_correction_original', '')
        review_accepted = False
        if original_text:
            try:
                dialog = DiffReviewDialog(original_text, corrected, log, self)
                if dialog.exec_() == QDialog.Accepted:
                    corrected = dialog.result_text
                    review_accepted = True
                else:
                    corrected = original_text
            except Exception as e:
                print(f"[DiffReview] error: {e}")
                corrected = original_text if original_text else corrected

        self.text_edit.setPlainText(corrected)
        self._all_logs = log

        try:
            self.feedback.log_corrections(log, source="corrector")
            if review_accepted:
                self.feedback.log_accept_all()
        except Exception as e:
            print(f"[Main] 纠错反馈记录失败: {e}")

        self._apply_filter()

        counts = {"错别字": 0, "逻辑错误": 0, "缺项提醒": 0}
        for item in log:
            cat = item.get("分类", "")
            if cat in counts:
                counts[cat] += 1
        total = sum(counts.values())
        self.stats_label.setText(f"共{total}条 | \U0001f4e2{counts['错别字']} \U0001f9e0{counts['逻辑错误']} ⚠️{counts['缺项提醒']}")

        self.status_bar.showMessage(f"纠错完成，共 {total} 条建议")
        Toast.show_toast(self, f"纠错完成，共 {total} 条建议", "success")

    def _apply_filter(self):
        self.log_list.clear()
        show_typo = self.filter_typo.isChecked()
        show_logic = self.filter_logic.isChecked()
        show_missing = self.filter_missing.isChecked()

        filter_map = {
            "错别字": show_typo,
            "逻辑错误": show_logic,
            "缺项提醒": show_missing,
        }

        self._log_item_map = {}

        for idx, item in enumerate(self._all_logs):
            cat = item.get("分类", "")
            if cat not in filter_map or not filter_map[cat]:
                continue

            if cat == "错别字":
                icon, color = "\U0001f4e2", "#00d4ff"
            elif cat == "逻辑错误":
                icon, color = "\U0001f9e0", "#ff9944"
            elif cat == "缺项提醒":
                icon, color = "⚠️", "#ffdd44"
            else:
                icon, color = "\U0001f4dd", "#b8c5d6"

            level = item.get("级别", "")
            type_name = item.get("type", "")
            orig = item.get("原文", "")
            corr = item.get("修正", "")

            line = f"{icon} <span style='color:{color};'>{type_name}</span>"
            if level:
                line += f" <span style='color:#6b8a9a;font-size:10px;'>[{level}]</span>"

            detail = ""
            if orig and orig != corr:
                detail = f"<span style='color:#ff6b6b;'>{orig}</span> → <span style='color:#51cf66;'>{corr}</span>"
            elif orig:
                detail = f"<span style='color:#e0e0e0;'>{orig}</span>"

            if item.get("相似度"):
                detail += f" <span style='color:#6b8a9a;font-size:10px;'>({item['相似度']})</span>"

            full_text = line
            if detail:
                full_text += f"<br>{detail}"

            list_item = QListWidgetItem(full_text)
            list_item.setData(Qt.UserRole, idx)
            self.log_list.addItem(list_item)

    def _on_log_item_clicked(self, item):
        idx = item.data(Qt.UserRole)
        if idx is None or idx >= len(self._all_logs):
            return
        log_data = self._all_logs[idx]
        orig = log_data.get("原文", "")
        corr = log_data.get("修正", "")

        search_text = corr if corr and corr != orig else orig
        if not search_text:
            return

        self._highlight_text_in_editor(search_text)
        self.log_hint.setText(f"已定位：{orig} → {corr}" if orig != corr else f"已定位：{orig}")
        self._accept_btn.setEnabled(True)
        self._reject_btn.setEnabled(True)

    def _highlight_text_in_editor(self, text):
        doc = self.text_edit.document()
        cursor = self.text_edit.textCursor()
        cursor.select(QTextCursor.Document)
        cursor.clearSelection()
        self.text_edit.setTextCursor(cursor)

        extra = []
        found_any = False
        cursor = doc.find(text, 0)
        while not cursor.isNull():
            found_any = True
            sel = QTextEdit.ExtraSelection()
            sel.cursor = cursor
            sel.format = QTextCharFormat()
            sel.format.setBackground(QColor(255, 200, 50, 100))
            extra.append(sel)
            cursor = doc.find(text, cursor)
        self.text_edit.setExtraSelections(extra)

        if found_any:
            first = doc.find(text, 0)
            if not first.isNull():
                self.text_edit.setTextCursor(first)

    def _accept_correction(self):
        item = self.log_list.currentItem()
        if not item:
            return
        idx = item.data(Qt.UserRole)
        if idx is None or idx >= len(self._all_logs):
            return
        log_data = self._all_logs[idx]
        orig = log_data.get("原文", "")
        corr = log_data.get("修正", "")
        if orig and corr:
            self.corrector.rejections.pop((orig, corr), None)
            try:
                serializable = [
                    "\x00".join(k) for k in self.corrector.rejections.keys()
                    if k[0] and k[1]
                ]
                with open(self.corrector.rejection_path, 'w', encoding='utf-8') as f:
                    json.dump(serializable, f, ensure_ascii=False, indent=2)
            except Exception as e:
                print(f"[Main] 保存 rejection 失败: {e}")
        self.status_bar.showMessage(f"✓ 已接受纠错：{orig} → {corr}")
        try:
            memory = self._get_memory()
            if memory and orig and corr and orig != corr:
                memory.accept_memory_by_values(
                    orig, corr,
                    doctor_id=self.current_user.get("id") if isinstance(self.current_user, dict) else None,
                    dept=getattr(self, "current_dept", "") or ""
                )
        except Exception as e:
            print(f"[Memory] 接受纠错失败: {e}")

    def _accept_all_corrections(self):
        if not self._all_logs:
            Toast.show_toast(self, "没有纠错建议", "info")
            return
        count = sum(1 for item in self._all_logs
                    if not item.get("_rejected") and item.get("原文") and item.get("修正"))
        for item in self._all_logs:
            orig = item.get("原文", "")
            corr = item.get("修正", "")
            if orig and corr:
                self.corrector.rejections.pop((orig, corr), None)
        try:
            serializable = [
                "\x00".join(k) for k in self.corrector.rejections.keys()
                if k[0] and k[1]
            ]
            with open(self.corrector.rejection_path, 'w', encoding='utf-8') as f:
                json.dump(serializable, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f"[Main] 保存 rejection 失败: {e}")
        try:
            self.feedback.log_accept_all()
        except Exception as e:
            print(f"[Main] 批量接受反馈失败: {e}")
        self.status_bar.showMessage(f"✓ 已接受全部 {count} 条纠错")
        Toast.show_toast(self, f"已接受全部 {count} 条纠错", "success")

    def _reject_correction(self):
        item = self.log_list.currentItem()
        if not item:
            return
        idx = item.data(Qt.UserRole)
        if idx is None or idx >= len(self._all_logs):
            return
        log_data = self._all_logs[idx]
        orig = log_data.get("原文", "")
        corr = log_data.get("修正", "")
        if not orig or not corr:
            return

        self.corrector.save_rejection(orig, corr)
        try:
            self.feedback.log_rejection(orig, corr)
        except Exception as e:
            print(f"[Main] 拒绝反馈记录失败: {e}")
        try:
            memory = self._get_memory()
            if memory and orig and corr and orig != corr:
                memory.reject_memory_by_values(
                    orig, corr,
                    doctor_id=self.current_user.get("id") if isinstance(self.current_user, dict) else None,
                    dept=getattr(self, "current_dept", "") or ""
                )
        except Exception as e:
            print(f"[Memory] 拒绝纠错失败: {e}")

        text = self.text_edit.toPlainText()
        if corr in text:
            text = text.replace(corr, orig)
        self.text_edit.setPlainText(text)

        self._all_logs[idx]["_rejected"] = True
        self._apply_filter()

        counts = {"错别字": 0, "逻辑错误": 0, "缺项提醒": 0}
        for item in self._all_logs:
            if item.get("_rejected"):
                continue
            cat = item.get("分类", "")
            if cat in counts:
                counts[cat] += 1
        total = sum(counts.values())
        self.stats_label.setText(f"共{total}条 | \U0001f4e2{counts['错别字']} \U0001f9e0{counts['逻辑错误']} ⚠️{counts['缺项提醒']}")

        self._accept_btn.setEnabled(False)
        self._reject_btn.setEnabled(False)
        self.log_hint.setText(f"已拒绝并记住：{orig}")
        self.status_bar.showMessage(f"✗ 已拒绝纠错并记录偏好：{orig} → {corr}")

    def _get_memory(self):
        if self.memory is None:
            try:
                self.memory = get_memory()
            except Exception as e:
                print(f"[Memory] 初始化失败: {e}")
        return self.memory

    def _get_topk_engine(self):
        if self.topk_engine is None:
            try:
                self.topk_engine = get_topk_engine(memory=self._get_memory())
            except Exception as e:
                print(f"[TopK] 初始化失败: {e}")
        return self.topk_engine

    # ══════════════════════════════════════════════════════
    # 文本操作
    # ══════════════════════════════════════════════════════

    def _clear_text(self):
        self.text_edit.clear()
        self.log_list.clear()
        self._all_logs = []
        self.stats_label.setText("")
        self.partial_label.setText("等待输入...")

    def _save_text(self):
        text = self.text_edit.toPlainText()
        if not text:
            QMessageBox.information(self, "提示", "没有内容可导出")
            return

        menu = QMenu(self)
        txt_action = QAction("\U0001f4c4 导出 .txt", self)
        txt_action.triggered.connect(lambda: self._export_as(text, "txt"))
        menu.addAction(txt_action)

        md_action = QAction("\U0001f4dd 导出 .md", self)
        md_action.triggered.connect(lambda: self._export_as(text, "md"))
        menu.addAction(md_action)

        docx_action = QAction("\U0001f4d8 导出 Word (.docx)", self)
        docx_action.triggered.connect(lambda: self._export_docx(text))
        menu.addAction(docx_action)

        preview_action = QAction("\U0001f5a8 打印预览", self)
        preview_action.triggered.connect(lambda: self._print_preview(text))
        menu.addAction(preview_action)

        sender_btn = self.sender()
        if sender_btn:
            btn_rect = sender_btn.geometry()
            menu.exec_(self.mapToGlobal(btn_rect.bottomLeft()))
        else:
            menu.exec_(self.mapToGlobal(self.record_btn.geometry().bottomLeft()))

    def _export_as(self, text, fmt):
        if fmt == "md":
            md_text = self._convert_to_markdown(text)
            default_name = "病历记录.md"
            filter_str = "Markdown 文件 (*.md);;所有文件 (*)"
        else:
            md_text = text
            default_name = "病历记录.txt"
            filter_str = "文本文件 (*.txt);;所有文件 (*)"

        path, _ = QFileDialog.getSaveFileName(
            self, f"导出 {fmt.upper()}",
            os.path.join(os.path.expanduser("~"), "Desktop", default_name),
            filter_str
        )
        if path:
            with open(path, 'w', encoding='utf-8') as f:
                f.write(md_text)
            self.status_bar.showMessage(f"已导出：{path}")
            try:
                self.feedback.collect_corpus(text)
            except Exception as e:
                print(f"[Main] 语料收集失败: {e}")

    def _convert_to_markdown(self, text):
        lines = text.split('\n')
        result = ["# 病历记录\n"]
        for line in lines:
            line = line.strip()
            if not line:
                continue
            if '：' in line or ':' in line:
                parts = re.split(r'[：:]', line, 1)
                if len(parts) == 2:
                    result.append(f"**{parts[0].strip()}**：{parts[1].strip()}\n")
                else:
                    result.append(f"{line}\n")
            else:
                result.append(f"{line}\n")
        return '\n'.join(result)

    def _export_docx(self, text):
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.oxml.ns import qn
        except ImportError:
            ret = QMessageBox.question(
                self, "缺少依赖",
                "导出 Word 需要 python-docx 库，当前未安装。\n"
                "可在命令行执行：pip install python-docx\n\n"
                "是否改为导出 .txt 文本？",
                QMessageBox.Yes | QMessageBox.No, QMessageBox.Yes
            )
            if ret == QMessageBox.Yes:
                self._export_as(text, "txt")
            return

        path, _ = QFileDialog.getSaveFileName(
            self, "导出 Word",
            os.path.join(os.path.expanduser("~"), "Desktop", "病历记录.docx"),
            "Word 文档 (*.docx)"
        )
        if not path:
            return
        try:
            doc = Document()
            style = doc.styles["Normal"]
            style.font.name = "宋体"
            style.font.size = Pt(11)
            style.element.rcPr.rFonts.set(qn("w:eastAsia"), "宋体")

            title = doc.add_heading("病历记录", level=1)
            for line in text.split("\n"):
                line = line.rstrip()
                if not line:
                    doc.add_paragraph("")
                    continue
                if "：" in line or ":" in line:
                    parts = re.split(r"[：:]", line, 1)
                    if len(parts) == 2:
                        p = doc.add_paragraph()
                        run = p.add_run(parts[0].strip() + "：")
                        run.bold = True
                        p.add_run(parts[1].strip())
                        continue
                doc.add_paragraph(line)
            doc.save(path)
            self.status_bar.showMessage(f"已导出 Word：{path}")
        except Exception as e:
            QMessageBox.warning(self, "导出失败", str(e))

    def _print_preview(self, text):
        from PyQt5.QtPrintSupport import QPrintPreviewDialog, QPrinter
        printer = QPrinter(QPrinter.HighResolution)
        dialog = QPrintPreviewDialog(printer, self)
        dialog.paintRequested.connect(lambda p: self._print_text(p, text))
        dialog.exec_()

    def _print_text(self, printer, text):
        from PyQt5.QtPrintSupport import QPrinter
        doc = QTextDocument()
        html = "<h2>病历记录</h2><hr/>"
        for line in text.split('\n'):
            escaped = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            if '：' in line or ':' in line:
                parts = re.split(r'[：:]', line, 1)
                if len(parts) == 2:
                    html += f"<p><b>{parts[0].strip()}</b>：{parts[1].strip()}</p>"
                else:
                    html += f"<p>{escaped}</p>"
            else:
                html += f"<p>{escaped}</p>"
        doc.setHtml(html)
        doc.print_(printer)

    # ══════════════════════════════════════════════════════
    # 对话框
    # ══════════════════════════════════════════════════════

    def _open_phrase_library(self):
        if self._phrase_dialog is None:
            self._phrase_dialog = PhraseDialog(self.phrase_lib, self)
            self._phrase_dialog.phrase_selected.connect(self._insert_term_at_cursor)
        self._phrase_dialog.show()
        self._phrase_dialog.raise_()
        self._phrase_dialog.activateWindow()

    def _open_backup_menu(self):
        menu = QMenu(self)
        backup_action = QAction("\U0001f4be 立即备份到文件", self)
        backup_action.triggered.connect(self._backup_now)
        menu.addAction(backup_action)

        restore_action = QAction("♻️ 从备份文件恢复", self)
        restore_action.triggered.connect(self._restore_backup)
        menu.addAction(restore_action)

        menu.addSeparator()
        hotword_action = QAction("\U0001f525 从历史病历刷新个人热词", self)
        hotword_action.triggered.connect(lambda: self._refresh_user_hotwords(silent=False))
        menu.addAction(hotword_action)

        sender_btn = self.sender()
        if sender_btn and hasattr(sender_btn, "geometry"):
            menu.exec_(self.mapToGlobal(sender_btn.geometry().bottomLeft()))
        else:
            menu.exec_(self.mapToGlobal(self.record_btn.geometry().bottomLeft()))

    def _backup_now(self):
        if not self.db:
            QMessageBox.warning(self, "提示", "数据库未初始化")
            return
        from datetime import datetime
        default_name = "病历备份_%s.db" % datetime.now().strftime("%Y%m%d_%H%M%S")
        path, _ = QFileDialog.getSaveFileName(
            self, "备份数据库",
            os.path.join(os.path.expanduser("~"), "Desktop", default_name),
            "SQLite 数据库 (*.db)"
        )
        if not path:
            return
        try:
            self.db.backup_to(path)
            QMessageBox.information(self, "备份成功", "已备份到：\n%s" % path)
        except Exception as e:
            QMessageBox.warning(self, "备份失败", str(e))

    def _restore_backup(self):
        if not self.db:
            QMessageBox.warning(self, "提示", "数据库未初始化")
            return
        ret = QMessageBox.warning(
            self, "确认恢复",
            "恢复将用备份数据覆盖当前所有病历，此操作不可撤销！\n建议先做一次备份。确定继续？",
            QMessageBox.Yes | QMessageBox.No, QMessageBox.No
        )
        if ret != QMessageBox.Yes:
            return
        path, _ = QFileDialog.getOpenFileName(
            self, "选择备份文件",
            os.path.expanduser("~"), "SQLite 数据库 (*.db)"
        )
        if not path:
            return
        try:
            self.db.restore_from(path)
            QMessageBox.information(self, "恢复成功", "已从备份恢复。请重启软件以确保数据一致。")
        except Exception as e:
            QMessageBox.warning(self, "恢复失败", str(e))

    def _open_rule_manager(self):
        dialog = RuleManagerDialog(self.rule_engine, self)
        dialog.exec_()
        self.status_bar.showMessage("\U0001f4cf 规则已更新")

    def _save_record(self):
        if not self.db or not self.current_user:
            QMessageBox.warning(self, "提示", "未登录，无法保存病历")
            return
        content = self.text_edit.toPlainText().strip()
        if not content:
            QMessageBox.information(self, "提示", "没有内容可保存")
            return
        patient_name = self._extract_patient_name(content)
        dept = self.current_dept if self.current_dept != "通用" else self.current_user.get("department", "")
        template_name = self.template_combo.currentText() if hasattr(self, "template_combo") else ""
        if self.current_record_id is None:
            self.current_record_id = self.db.create_record(
                self.current_user["id"], patient_name, dept, template_name, content, "草稿"
            )
            self.status_bar.showMessage("\U0001f4be 病历已保存到病历库（新建）")
            Toast.show_toast(self, "病历已保存", "success")
        else:
            self.db.update_record(
                self.current_record_id, patient_name=patient_name,
                department=dept, template_name=template_name, content=content
            )
            self.status_bar.showMessage("\U0001f4be 病历已更新（已记录版本）")
            Toast.show_toast(self, "病历已更新", "success")

        try:
            learned = [w for w in self.corrector.active_words
                       if len(w) >= 2 and content.count(w) >= 1]
            if learned:
                self.asr.update_user_hotwords(learned)
            self._refresh_topk_hotwords(silent=True)
        except Exception as e:
            print(f"[Main] 增量学习热词失败: {e}")

        try:
            self.feedback.collect_corpus(content)
            self.feedback.log_accept_all()
        except Exception as e:
            print(f"[Main] 保存后反馈收集失败: {e}")

        try:
            memory = self._get_memory()
            if memory:
                doctor_id = self.current_user.get("id") if isinstance(self.current_user, dict) else None
                dept = getattr(self, "current_dept", "") or ""
                memory.record_final_text(
                    content, doctor_id=doctor_id, dept=dept,
                    record_id=getattr(self, "current_record_id", None),
                    snapshot=getattr(self, "_last_asr_snapshot", '')
                )
        except Exception as e:
            print(f"[Memory] 记录终稿失败: {e}")

        try:
            self._extract_manual_corrections(content)
        except Exception as e:
            print(f"[Main] 提取手动修正失败: {e}")

    def _extract_patient_name(self, content):
        m = re.search(r'姓名[：:]\s*([^\s　\n]{1,10})', content)
        return m.group(1).strip() if m else ""

    def _extract_manual_corrections(self, final_text):
        snapshot = getattr(self, '_last_asr_snapshot', '')
        if not snapshot or not final_text:
            return
        if snapshot.strip() == final_text.strip():
            return

        sm = difflib.SequenceMatcher(None, snapshot, final_text)
        corrections = []
        for tag, i1, i2, j1, j2 in sm.get_opcodes():
            if tag == 'replace':
                old = snapshot[i1:i2].strip()
                new = final_text[j1:j2].strip()
                if old and new and 2 <= len(old) <= 20 and 2 <= len(new) <= 20:
                    if old.replace(' ', '') != new.replace(' ', ''):
                        corrections.append((old, new))

        if not corrections:
            return

        import datetime
        feedback_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "correction_feedback.jsonl")
        with open(feedback_path, 'a', encoding='utf-8') as f:
            for old, new in corrections:
                record = {
                    "original": old, "corrected": new,
                    "原文": old, "修正": new,
                    "type": "manual_edit", "status": "accepted",
                    "source": "user_manual",
                    "timestamp": datetime.datetime.now().isoformat()
                }
                f.write(json.dumps(record, ensure_ascii=False) + '\n')

        print(f"[Main] 提取手动修正: {len(corrections)} 对 → {corrections[:5]}")
        self._last_asr_snapshot = final_text

    def _open_record_manager(self):
        if not self.db or not self.current_user:
            QMessageBox.warning(self, "提示", "未登录，无法打开病历库")
            return
        dialog = RecordManagerDialog(self.db, self.current_user, self)
        if dialog.exec_() == QDialog.Accepted and dialog.selected_record:
            rec = dialog.selected_record
            self.text_edit.setPlainText(rec["content"])
            self.current_record_id = rec["id"]
            self.status_bar.showMessage(
                "\U0001f4da 已打开病历：%s（%s）" % (rec["patient_name"] or "未命名", rec["updated_at"])
            )

    def _open_user_manager(self):
        if self.current_user.get("role") != "admin":
            QMessageBox.warning(self, "权限不足", "仅管理员可管理用户")
            return
        dialog = UserManagerDialog(self.db, self)
        dialog.exec_()

    # ══════════════════════════════════════════════════════
    # AI 辅助诊断
    # ══════════════════════════════════════════════════════

    def _run_diagnosis(self):
        text = self.text_edit.toPlainText().strip()
        if not text:
            QMessageBox.information(self, "提示", "请先输入或录制病历内容")
            return
        if self.diagnosis_thread is not None and self.diagnosis_thread.isRunning():
            return
        self.ai_analyze_btn.setEnabled(False)
        self.ai_status_label.setText("分析中...")
        self.diagnosis_thread = DiagnosisThread(self.diagnosis_assistant, text)
        self.diagnosis_thread.analysis_done.connect(self._on_diagnosis_done)
        self.diagnosis_thread.start()

    def _on_diagnosis_done(self, result):
        self.ai_analyze_btn.setEnabled(True)
        self.ai_status_label.setText("")
        if result.get("error"):
            self.ai_result.setHtml(f'<div style="color:#ff6b6b;">分析失败：{result["error"]}</div>')
            return
        self.ai_result.setHtml(self._render_diagnosis_html(result))
        self.status_bar.showMessage("\U0001f52c AI 分析完成")

    def _show_qa_dialog(self):
        try:
            from qa_dialog import KnowledgeQADialog
            dialog = KnowledgeQADialog(qa_engine=self.qa_engine, parent=self)
            dialog.exec_()
        except Exception as e:
            import traceback
            self.status_bar.showMessage(f"⚠️ 问答启动失败：{e}")
            traceback.print_exc()

    def _render_diagnosis_html(self, result):
        def section(title, color="#00d4ff"):
            return (f'<div style="color:{color};font-weight:bold;'
                    f'margin:10px 0 4px 0;">{title}</div>')

        html = []

        alerts = result.get("risk_alerts", [])
        if alerts:
            html.append(section("⚠️ 风险预警", "#ff6b6b"))
            html.append('<div style="background:rgba(255,107,107,0.12);'
                        'border:1px solid rgba(255,107,107,0.4);'
                        'border-radius:6px;padding:6px 10px;color:#ff9b9b;">')
            for a in alerts:
                if isinstance(a, (tuple, list)):
                    label, msg = (a[0], a[1]) if len(a) >= 2 else ("", a[0])
                    html.append(f'▸ <b>{label}</b>：{msg}<br>')
                else:
                    html.append(f'▸ {a}<br>')
            html.append('</div>')

        diagnoses = result.get("diagnoses", [])
        if diagnoses:
            html.append(section("\U0001f9ed 可能诊断"))
            html.append('<div style="color:#e0e0e0;">')
            for d in diagnoses:
                if isinstance(d, dict):
                    name = d.get("disease", d.get("name", ""))
                    score = d.get("score")
                    matched = d.get("matched")
                    line = f'• <b>{name}</b>'
                    if score is not None:
                        line += f' <span style="color:#6b8a9a;">(评分 {score})</span>'
                    if matched:
                        line += f' <span style="color:#6b8a9a;">→ {"、".join(matched)}</span>'
                    html.append(line + '<br>')
                else:
                    html.append(f'• {d}<br>')
            html.append('</div>')

        tcm = result.get("tcm_analysis")
        if tcm:
            html.append(section("\U0001f33f 中医辨证", "#4ecdc4"))
            html.append('<div style="color:#e0e0e0;">')
            tcm_dx = tcm.get("tcm_diagnoses", [])
            if tcm_dx:
                html.append(f'• 中医诊断：<b style="color:#4ecdc4;">{"、".join(tcm_dx)}</b><br>')
            syndromes = tcm.get("syndromes", [])
            if syndromes:
                for s in syndromes:
                    if isinstance(s, dict):
                        name = s.get("syndrome", "")
                        score = s.get("score", "")
                        matched = s.get("matched", [])
                        line = f'• 证型：<b style="color:#4ecdc4;">{name}</b>'
                        if score:
                            line += f' <span style="color:#6b8a9a;">(评分 {score})</span>'
                        if matched:
                            line += f' <span style="color:#6b8a9a;">→ {"、".join(matched)}</span>'
                        html.append(line + '<br>')
            treatment = tcm.get("treatment")
            if treatment:
                method = treatment.get("治法", "")
                formula = treatment.get("代表方", "")
                herbs = treatment.get("组成", [])
                html.append(f'• 治法：{method}<br>')
                if formula:
                    html.append(f'• 代表方：<b>{formula}</b><br>')
                if herbs:
                    html.append(f'<span style="color:#6b8a9a;"> 组成：{"、".join(herbs)}</span><br>')
            differential = tcm.get("differential", [])
            if differential:
                html.append('<div style="margin-top:4px;color:#ffa94d;">类证鉴别：</div>')
                for d in differential:
                    if isinstance(d, dict):
                        syn = d.get("syndrome", "")
                        kp = d.get("key_points", "")
                        tf = d.get("治法", "")
                        rx = d.get("代表方", "")
                        html.append(f'  • {syn}：{kp}')
                        if tf:
                            html.append(f' → {tf}')
                        if rx:
                            html.append(f'（{rx}）')
                        html.append('<br>')
            html.append('</div>')

        review = result.get("drug_review") or {}
        if review:
            matched = review.get("matched", [])
            mismatched = review.get("mismatched", [])
            recommended = review.get("recommended", [])
            if matched or mismatched or recommended:
                html.append(section("\U0001f48a 用药审查"))
                html.append('<div style="color:#e0e0e0;">')
                for m in matched:
                    if isinstance(m, dict):
                        drug = m.get("drug", "")
                        forl = m.get("for", [])
                        suffix = f'（适用：{"、".join(forl)}）' if forl else ""
                        html.append(f'✅ {drug}{suffix}<br>')
                    else:
                        html.append(f'✅ {m}<br>')
                for m in mismatched:
                    if isinstance(m, dict):
                        drug = m.get("drug", "")
                        note = m.get("note", "与当前诊断不匹配")
                        html.append(f'<span style="color:#ffa94d;">⚠️ {drug}：{note}</span><br>')
                    else:
                        html.append(f'<span style="color:#ffa94d;">⚠️ {m}</span><br>')
                if recommended:
                    html.append('<span style="color:#6b8a9a;">建议补充用药：'
                                + "、".join(recommended) + '</span><br>')
                html.append('</div>')

        exams = result.get("exam_suggestions", [])
        if exams:
            html.append(section("\U0001f50d 检查建议"))
            html.append('<div style="color:#e0e0e0;">')
            for e in exams:
                if isinstance(e, dict):
                    name = e.get("exam", "")
                    recorded = e.get("recorded")
                    tag = '已记录' if recorded else '<span style="color:#ffa94d;">建议补充</span>'
                    html.append(f'• {name}（{tag}）<br>')
                else:
                    html.append(f'• {e}<br>')
            html.append('</div>')

        disclaimer = result.get("disclaimer", "")
        if disclaimer:
            html.append(f'<div style="color:#6b8a9a;font-size:11px;'
                        f'margin-top:12px;border-top:1px solid rgba(255,255,255,0.1);'
                        f'padding-top:6px;">{disclaimer}</div>')

        if not html:
            return '<div style="color:#6b8a9a;">未提取到可分析的临床信息。</div>'
        return "".join(html)

    def _open_struct_view(self):
        text = self.text_edit.toPlainText().strip()
        if not text:
            QMessageBox.information(self, "提示", "请先输入或录制文本")
            return

        dialog = SectionDialog(self.parser, text, self)
        if dialog.exec_() == QDialog.Accepted:
            structured = dialog.get_result()
            self.text_edit.setPlainText(structured)
            self.status_bar.showMessage("\U0001f4cb 结构化填充完成")

    def _retrain_lm(self):
        import subprocess
        corpus_file = os.path.join(os.path.dirname(os.path.abspath(__file__)), "user_corpus.txt")
        corpus_count = 0
        if os.path.exists(corpus_file):
            with open(corpus_file, "r", encoding="utf-8") as f:
                corpus_count = sum(1 for line in f if line.strip())

        reply = QMessageBox.question(
            self, "重训语言模型",
            f"当前已累积用户语料: {corpus_count} 句\n\n"
            f"重训将合并基础语料 + 用户语料 + 纠错反馈，\n"
            f"生成新的 3-gram 模型（重启后生效）。\n\n开始训练？",
            QMessageBox.Yes | QMessageBox.No
        )
        if reply != QMessageBox.Yes:
            return

        self.status_bar.showMessage("\U0001f9e0 正在重训语言模型...")
        try:
            script = os.path.join(os.path.dirname(os.path.abspath(__file__)), "train_lm.py")
            result = subprocess.run(
                [sys.executable, script],
                capture_output=True, text=True, timeout=120,
                cwd=os.path.dirname(os.path.abspath(__file__))
            )
            if result.returncode == 0:
                output = result.stdout
                Toast.show_toast(self, "语言模型重训完成，重启后生效", "success")
                self.status_bar.showMessage("\U0001f9e0 语言模型已更新，重启程序后生效")
                print(f"[LM] 重训完成:\n{output[-500:]}")
                self._refresh_topk_hotwords(silent=True)
            else:
                QMessageBox.warning(self, "训练失败", result.stderr[-300:] or "未知错误")
                self.status_bar.showMessage("⚠️ 语言模型训练失败")
        except subprocess.TimeoutExpired:
            QMessageBox.warning(self, "超时", "训练超时（>120s），请通过命令行执行")
        except Exception as e:
            QMessageBox.warning(self, "错误", str(e))

    def _view_crash_log(self):
        logs = self.crash_logger.get_recent_logs(lines=200)
        if not logs:
            QMessageBox.information(self, "崩溃日志", "暂无日志记录")
            return

        dialog = QDialog(self)
        dialog.setWindowTitle("崩溃日志")
        dialog.setModal(True)
        dialog.resize(800, 600)

        layout = QVBoxLayout(dialog)

        text_edit = QTextEdit()
        text_edit.setReadOnly(True)
        text_edit.setText(''.join(logs))
        text_edit.setStyleSheet("""
            QTextEdit {
                background: #1a1a2e; color: #e0e0e0;
                font-family: monospace; font-size: 11px;
            }
        """)
        layout.addWidget(text_edit)

        bottom = QHBoxLayout()
        bottom.addStretch()
        clear_btn = QPushButton("清空日志")
        clear_btn.clicked.connect(self._clear_crash_log)
        close_btn = QPushButton("关闭")
        close_btn.clicked.connect(dialog.accept)
        bottom.addWidget(clear_btn)
        bottom.addWidget(close_btn)
        layout.addLayout(bottom)

        dialog.exec_()

    def _clear_crash_log(self):
        log_path = self.crash_logger.get_log_path()
        try:
            with open(log_path, 'w', encoding='utf-8') as f:
                pass
            self.status_bar.showMessage("日志已清空")
        except Exception as e:
            QMessageBox.warning(self, "错误", f"清空日志失败: {e}")

    def _open_template_manager(self):
        dialog = TemplateManagerDialog(
            self.template_engine, self.current_dept, self
        )
        if dialog.exec_() == QDialog.Accepted:
            self._load_departments()
            idx = self.dept_combo.findText(self.current_dept)
            if idx >= 0:
                self.dept_combo.setCurrentIndex(idx)
            self.status_bar.showMessage("\U0001f4dd 模板已更新")


def main():
    """启动入口（thin wrapper，实际逻辑在 gui/__init__.py 的重新导出中）"""
    # 向后兼容：直接在此处保留入口逻辑
    from gui import MedVoiceApp, WebViewApp
    from license_manager import LicenseManager
    from activation_dialog import ActivationDialog, TrialInfoBar
    from database import Database
    from login_dialog import LoginDialog

    app = QApplication(sys.argv)
    app.setStyle('Fusion')

    font = QFont("Microsoft YaHei", 11)
    app.setFont(font)

    license_mgr = LicenseManager()
    status = license_mgr.check_license()

    if status["status"] in ("expired", "tampered"):
        dlg = ActivationDialog(license_mgr, status)
        if dlg.exec_() != QDialog.Accepted:
            sys.exit(0)

    db = Database()
    login = LoginDialog(db)
    if login.exec_() != QDialog.Accepted or not login.current_user:
        sys.exit(0)

    use_webview = '--legacy' not in sys.argv
    if use_webview and _HAS_WEBENGINE:
        try:
            window = WebViewApp(db=db, current_user=login.current_user)
        except Exception as e:
            print(f"[Main] WebView 初始化失败，回退到原生 UI: {e}")
            window = MedVoiceApp(db=db, current_user=login.current_user)
    else:
        window = MedVoiceApp(db=db, current_user=login.current_user)
    window.show()

    sys.exit(app.exec_())


if __name__ == "__main__":
    main()
