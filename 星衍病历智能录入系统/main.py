"""
星衍AI智能病历录入系统 - 主程序
离线使用，基于 Vosk 语音识别 + 医疗词库纠错
"""
import sys
import os

# WebEngine 环境变量（必须在任何 Qt 导入前设置）
import os as _os_pre
_venv_qt = _os_pre.path.join(_os_pre.path.dirname(_os_pre.path.abspath(__file__)),
    'venv', 'lib', 'python3.14', 'site-packages', 'PyQt5', 'Qt5')
_os_pre.environ['QTWEBENGINE_RESOURCES_PATH'] = _os_pre.path.join(_venv_qt, 'lib', 'QtWebEngineCore.framework', 'Resources')
_os_pre.environ['QTWEBENGINE_LOCALES_PATH'] = _os_pre.path.join(_venv_qt, 'lib', 'QtWebEngineCore.framework', 'Resources', 'qtwebengine_locales')
_os_pre.environ.setdefault('QTWEBENGINE_CHROMIUM_FLAGS', '--no-sandbox')
del _os_pre, _venv_qt

import json
import time
import re
import threading

from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QPushButton, QTextEdit, QComboBox, QLabel, QSplitter,
    QListWidget, QListWidgetItem, QStatusBar, QToolBar,
    QMessageBox, QCheckBox, QGroupBox, QFileDialog,
    QDialog, QLineEdit, QTableWidget, QTableWidgetItem,
    QTabWidget, QHeaderView, QTextBrowser, QToolButton,
    QAction, QMenu, QScrollArea, QFrame, QInputDialog
)
from PyQt5.QtCore import Qt, QTimer, pyqtSignal, QThread
from PyQt5.QtGui import QFont, QColor, QPalette, QTextCursor, QTextCharFormat, QTextDocument

# QtWebEngine 必须在 QApplication 创建前导入
try:
    from PyQt5.QtWebEngineWidgets import QWebEngineView  # noqa: F401
    _HAS_WEBENGINE = True
except ImportError:
    _HAS_WEBENGINE = False

from corrector import Corrector
from asr_engine import ASREngine
from template_engine import TemplateEngine
from rule_engine import RuleEngine
from section_parser import SectionParser, SmartDictation
from medical_classifier import MedicalClassifier
from crash_logger import CrashLogger
from database import Database
from login_dialog import LoginDialog, UserManagerDialog
from record_manager_dialog import RecordManagerDialog
from voice_command import VoiceCommandParser
from asr_engine import get_microphone_list
from diagnosis_assistant import DiagnosisAssistant
from knowledge_qa import KnowledgeQA
from license_manager import LicenseManager
from activation_dialog import ActivationDialog, TrialInfoBar
from phrase_library import PhraseLibrary
from phrase_dialog import PhraseDialog
from correction_feedback import CorrectionFeedback
from ux_components import Toast, RecordingIndicator, OnboardingGuide, FieldHighlighter
from audio_widgets import WaveformWidget, AudioPlayerWidget, AsrPreviewPanel
from diff_review_dialog import DiffReviewDialog
from correction_memory import get_memory
from topk_engine import get_topk_engine
from recording_handler import RecordingHandler
from threads import ListenThread, CorrectThread, DiagnosisThread, create_listen_thread, stop_listen_thread


# ==================== 规则管理对话框 ====================
class RuleManagerDialog(QDialog):
    """自定义纠错规则管理界面"""

    def __init__(self, rule_engine, parent=None):
        super().__init__(parent)
        self.rule_engine = rule_engine
        self.setWindowTitle("📏 纠错规则管理")
        self.setModal(True)
        self.resize(700, 500)
        self._init_ui()

    def _init_ui(self):
        layout = QVBoxLayout(self)

        # 标签页：错别字 / 逻辑错误
        tabs = QTabWidget()

        # ---- 错别字标签页 ----
        typo_tab = QWidget()
        typo_layout = QVBoxLayout(typo_tab)

        # 添加错别字规则
        add_typo = QWidget()
        add_typo_layout = QHBoxLayout(add_typo)
        add_typo_layout.setContentsMargins(0, 0, 0, 10)
        self.typo_wrong_input = QLineEdit()
        self.typo_wrong_input.setPlaceholderText("错误写法（如：心电围）")
        self.typo_correct_input = QLineEdit()
        self.typo_correct_input.setPlaceholderText("正确写法（如：心电图）")
        add_typo_btn = QPushButton("➕ 添加规则")
        add_typo_btn.clicked.connect(self._add_typo_rule)
        add_typo_btn.setStyleSheet("""
            QPushButton {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #00d4ff, stop:1 #0066ff);
                color: #0a0e27;
                font-weight: bold;
                padding: 6px 16px;
                border-radius: 15px;
                font-size: 12px;
            }
            QPushButton:hover { padding: 6px 20px; }
        """)
        add_typo_layout.addWidget(QLabel("错误："))
        add_typo_layout.addWidget(self.typo_wrong_input)
        add_typo_layout.addWidget(QLabel("正确："))
        add_typo_layout.addWidget(self.typo_correct_input)
        add_typo_layout.addWidget(add_typo_btn)
        typo_layout.addWidget(add_typo)

        # 错别字规则列表
        self.typo_table = QTableWidget()
        self.typo_table.setColumnCount(3)
        self.typo_table.setHorizontalHeaderLabels(["错误", "正确", "操作"])
        self.typo_table.horizontalHeader().setSectionResizeMode(0, QHeaderView.Stretch)
        self.typo_table.horizontalHeader().setSectionResizeMode(1, QHeaderView.Stretch)
        self.typo_table.horizontalHeader().setSectionResizeMode(2, QHeaderView.Fixed)
        self.typo_table.setColumnWidth(2, 80)
        self.typo_table.setStyleSheet("""
            QTableWidget {
                background: rgba(255,255,255,0.03);
                border: 1px solid rgba(0,212,255,0.1);
                border-radius: 8px;
                color: #e0e0e0;
                font-size: 13px;
            }
            QHeaderView::section {
                background: rgba(0,212,255,0.1);
                color: #00d4ff;
                padding: 6px;
                border: none;
            }
        """)
        typo_layout.addWidget(self.typo_table)
        self._refresh_typo_table()

        tabs.addTab(typo_tab, "🔤 错别字规则")

        # ---- 逻辑错误标签页 ----
        logic_tab = QWidget()
        logic_layout = QVBoxLayout(logic_tab)

        # 添加逻辑错误规则
        add_logic = QWidget()
        add_logic_layout = QHBoxLayout(add_logic)
        add_logic_layout.setContentsMargins(0, 0, 0, 10)
        self.logic_name_input = QLineEdit()
        self.logic_name_input.setPlaceholderText("规则名称（如：疾病与症状不符）")
        self.logic_desc_input = QLineEdit()
        self.logic_desc_input.setPlaceholderText("规则描述")
        add_logic_btn = QPushButton("➕ 添加规则")
        add_logic_btn.clicked.connect(self._add_logic_rule)
        add_logic_btn.setStyleSheet("""
            QPushButton {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #ff9944, stop:1 #ff6600);
                color: white;
                font-weight: bold;
                padding: 6px 16px;
                border-radius: 15px;
                font-size: 12px;
            }
            QPushButton:hover { padding: 6px 20px; }
        """)
        add_logic_layout.addWidget(QLabel("名称："))
        add_logic_layout.addWidget(self.logic_name_input)
        add_logic_layout.addWidget(QLabel("描述："))
        add_logic_layout.addWidget(self.logic_desc_input)
        add_logic_layout.addWidget(add_logic_btn)
        logic_layout.addWidget(add_logic)

        # 逻辑错误规则列表
        self.logic_table = QTableWidget()
        self.logic_table.setColumnCount(3)
        self.logic_table.setHorizontalHeaderLabels(["错误模式", "描述", "操作"])
        self.logic_table.horizontalHeader().setSectionResizeMode(0, QHeaderView.Fixed)
        self.logic_table.setColumnWidth(0, 160)
        self.logic_table.horizontalHeader().setSectionResizeMode(1, QHeaderView.Stretch)
        self.logic_table.horizontalHeader().setSectionResizeMode(2, QHeaderView.Fixed)
        self.logic_table.setColumnWidth(2, 80)
        self.logic_table.setStyleSheet("""
            QTableWidget {
                background: rgba(255,255,255,0.03);
                border: 1px solid rgba(0,212,255,0.1);
                border-radius: 8px;
                color: #e0e0e0;
                font-size: 13px;
            }
            QHeaderView::section {
                background: rgba(255,153,68,0.1);
                color: #ff9944;
                padding: 6px;
                border: none;
            }
        """)
        logic_layout.addWidget(self.logic_table)
        self._refresh_logic_table()

        tabs.addTab(logic_tab, "🧠 逻辑错误规则")

        layout.addWidget(tabs)

        # 底部按钮
        bottom = QHBoxLayout()
        bottom.addStretch()
        close_btn = QPushButton("关闭")
        close_btn.clicked.connect(self.accept)
        close_btn.setStyleSheet("""
            QPushButton {
                background: rgba(255,255,255,0.1);
                color: #b8c5d6;
                padding: 8px 24px;
                border-radius: 15px;
                border: 1px solid rgba(255,255,255,0.1);
            }
            QPushButton:hover { background: rgba(255,255,255,0.15); }
        """)
        bottom.addWidget(close_btn)
        layout.addLayout(bottom)

    def _refresh_typo_table(self):
        """刷新错别字规则列表"""
        rules = self.rule_engine.get_typo_rules()
        self.typo_table.setRowCount(len(rules))
        for i, rule in enumerate(rules):
            self.typo_table.setItem(i, 0, QTableWidgetItem(rule["错误"]))
            self.typo_table.setItem(i, 1, QTableWidgetItem(rule["正确"]))
            # 删除按钮
            del_btn = QPushButton("🗑")
            del_btn.setFixedSize(40, 30)
            del_btn.clicked.connect(lambda _, w=rule["错误"]: self._delete_typo(w))
            del_btn.setStyleSheet("""
                QPushButton {
                    background: rgba(255,80,80,0.1);
                    border: 1px solid rgba(255,80,80,0.2);
                    border-radius: 5px;
                    color: #ff6b6b;
                    font-size: 14px;
                }
                QPushButton:hover { background: rgba(255,80,80,0.2); }
            """)
            self.typo_table.setCellWidget(i, 2, del_btn)

    def _refresh_logic_table(self):
        """刷新逻辑错误规则列表"""
        rules = self.rule_engine.get_logic_rules()
        self.logic_table.setRowCount(len(rules))
        for i, rule in enumerate(rules):
            self.logic_table.setItem(i, 0, QTableWidgetItem(rule.get("错误模式", "")))
            self.logic_table.setItem(i, 1, QTableWidgetItem(rule.get("描述", "")))
            del_btn = QPushButton("🗑")
            del_btn.setFixedSize(40, 30)
            del_btn.clicked.connect(lambda _, n=rule["错误模式"]: self._delete_logic(n))
            del_btn.setStyleSheet("""
                QPushButton {
                    background: rgba(255,80,80,0.1);
                    border: 1px solid rgba(255,80,80,0.2);
                    border-radius: 5px;
                    color: #ff6b6b;
                    font-size: 14px;
                }
                QPushButton:hover { background: rgba(255,80,80,0.2); }
            """)
            self.logic_table.setCellWidget(i, 2, del_btn)

    def _add_typo_rule(self):
        """添加错别字规则"""
        wrong = self.typo_wrong_input.text().strip()
        correct = self.typo_correct_input.text().strip()
        if not wrong or not correct:
            QMessageBox.warning(self, "提示", "请填写错误写法和正确写法")
            return
        if wrong == correct:
            QMessageBox.warning(self, "提示", "错误写法和正确写法不能相同")
            return

        is_new = self.rule_engine.add_typo_rule(wrong, correct)
        self.typo_wrong_input.clear()
        self.typo_correct_input.clear()
        self._refresh_typo_table()
        msg = "规则已添加" if is_new else "规则已更新"
        self.parent().status_bar.showMessage(f"📏 {msg}: {wrong} → {correct}")

    def _add_logic_rule(self):
        """添加逻辑错误规则"""
        name = self.logic_name_input.text().strip()
        desc = self.logic_desc_input.text().strip()
        if not name or not desc:
            QMessageBox.warning(self, "提示", "请填写规则名称和描述")
            return

        self.rule_engine.add_logic_rule(name, desc)
        self.logic_name_input.clear()
        self.logic_desc_input.clear()
        self._refresh_logic_table()
        self.parent().status_bar.showMessage(f"📏 逻辑规则已添加: {name}")

    def _delete_typo(self, wrong):
        """删除错别字规则"""
        confirm = QMessageBox.question(
            self, "确认删除",
            f"确定删除规则「{wrong}」吗？",
            QMessageBox.Yes | QMessageBox.No
        )
        if confirm == QMessageBox.Yes:
            self.rule_engine.delete_typo_rule(wrong)
            self._refresh_typo_table()
            self.parent().status_bar.showMessage(f"📏 规则已删除: {wrong}")

    def _delete_logic(self, name):
        """删除逻辑错误规则"""
        confirm = QMessageBox.question(
            self, "确认删除",
            f"确定删除规则「{name}」吗？",
            QMessageBox.Yes | QMessageBox.No
        )
        if confirm == QMessageBox.Yes:
            self.rule_engine.delete_logic_rule(name)
            self._refresh_logic_table()
            self.parent().status_bar.showMessage(f"📏 规则已删除: {name}")


# ==================== 字段常用词面板 ====================
class FieldWordsPanel(QWidget):
    """字段常用词面板 - 按字段分类展示可点击插入的常用词"""
    term_clicked = pyqtSignal(str)  # 发出 (字段名, 词语) 信号

    # 内置默认常用词（当 field_words.json 不存在时使用）
    DEFAULT_WORDS = {
        "主诉": {
            "label": "主诉",
            "terms": [
                "发热 3 天", "咳嗽 1 周", "胸痛 2 天", "腹痛 1 天",
                "头痛 3 天", "头晕 1 周", "呼吸困难 2 天", "胸闷 1 周",
                "乏力 1 月", "消瘦 2 月", "恶心呕吐 1 天", "腹泻 2 天",
                "便血 1 天", "水肿 1 周", "意识不清 2 小时",
            ]
        },
        "现病史": {
            "label": "现病史",
            "terms": [
                "患者于 X 天前无明显诱因出现",
                "伴发热，体温最高达 38.5℃",
                "无恶心、呕吐、腹泻",
                "自行口服药物后症状无明显缓解",
                "门诊查血常规：白细胞升高",
                "胸片提示右下肺感染",
                "予以抗感染、补液等对症治疗",
                "症状有所缓解，为进一步诊治入院",
            ]
        },
        "既往史": {
            "label": "既往史",
            "terms": [
                "否认高血压、糖尿病、冠心病病史",
                "否认肝炎、结核等传染病史",
                "否认食物及药物过敏史",
                "否认手术史、外伤史、输血史",
                "高血压病史 X 年，口服药物控制可",
                "糖尿病病史 X 年",
                "吸烟史 X 年，约 X 支/日",
                "饮酒史 X 年",
                "父亲患有高血压，母亲患有糖尿病",
            ]
        },
        "体格检查": {
            "label": "体格检查",
            "terms": [
                "T 36.5℃，P 78 次/分，R 20 次/分，BP 128/80 mmHg",
                "神志清楚，精神可，发育正常，营养中等",
                "自主体位，步入病房",
                "全身皮肤黏膜无黄染，未见皮疹、出血点",
                "颈软，无抵抗",
                "胸廓对称，双肺叩诊清音，双肺呼吸音清晰",
                "心前区无隆起，心率 78 次/分，心律齐",
                "腹平坦，腹软，全腹无压痛、反跳痛",
                "肝脾肋下未触及，肠鸣音正常",
                "双下肢无水肿，四肢肌力 V 级",
            ]
        },
        "辅助检查": {
            "label": "辅助检查",
            "terms": [
                "血常规：白细胞 12×10^9/L，中性粒细胞 85%",
                "尿常规：未见明显异常",
                "大便常规：黄色软便，隐血阴性",
                "血生化：谷丙转氨酶 35U/L，肌酐 78μmol/L",
                "凝血功能：PT 12.5s，APTT 30s",
                "血糖：空腹 5.6mmol/L",
                "心电图：窦性心律，心率 75 次/分",
                "胸片：心肺膈未见明显异常",
                "胸部 CT：右肺中叶斑片状高密度影，考虑炎症",
                "头颅 CT：未见出血及占位性病变",
                "腹部 B 超：肝胆胰脾未见明显异常",
            ]
        },
        "初步诊断": {
            "label": "初步诊断",
            "terms": [
                "1. 社区获得性肺炎",
                "2. 高血压病 2 级，很高危",
                "3. 2 型糖尿病",
                "4. 冠状动脉粥样硬化性心脏病",
                "5. 慢性阻塞性肺疾病",
                "6. 急性阑尾炎",
                "7. 急性胆囊炎",
                "8. 脑梗死（急性期）",
                "9. 消化性溃疡伴出血",
                "10. 急性胰腺炎",
            ]
        },
        "诊疗计划": {
            "label": "诊疗计划",
            "terms": [
                "完善血常规、血生化、凝血功能、心电图等检查",
                "抗感染治疗：予头孢曲松钠 2g qd ivgtt",
                "化痰止咳：氨溴索 30mg tid",
                "补液维持水电解质平衡",
                "监测生命体征，定期复查血常规",
                "低盐低脂糖尿病饮食",
                "请示上级医师",
            ]
        },
        "鉴别诊断": {
            "label": "鉴别诊断",
            "terms": [
                "1. 肺结核：需胸片/CT 及结核菌素试验鉴别",
                "2. 支气管肺癌：需 CT 及病理鉴别",
                "3. 支气管扩张：多有反复咳嗽、咳大量脓痰病史",
                "4. 肺脓肿：CT 可见空洞及液平",
            ]
        },
        "专科情况": {
            "label": "专科情况",
            "terms": {
                "呼吸科": [
                    "胸廓对称，双肺叩诊清音",
                    "双肺呼吸音粗，可闻及干湿性啰音",
                    "语音共振无增强或减弱",
                ],
                "心血管内科": [
                    "心前区无隆起，心浊音界正常",
                    "心率 78 次/分，心律齐",
                    "各瓣膜区未闻及病理性杂音",
                    "双下肢无水肿",
                ],
                "神经内科": [
                    "神志清楚，言语流利",
                    "双侧瞳孔等大等圆，对光反射灵敏",
                    "四肢肌力 V 级，肌张力正常",
                    "双侧巴宾斯基征阴性",
                    "颈软，无抵抗",
                ],
                "消化内科": [
                    "腹平坦，未见胃肠型",
                    "腹软，全腹无压痛、反跳痛、肌紧张",
                    "肝脾肋下未触及",
                    "Murphy 征阴性",
                    "肠鸣音正常",
                ],
            }
        },
    }

    def __init__(self, parent=None):
        super().__init__(parent)
        self._words_data = {}
        self._presets_data = {}
        self._current_field = ""
        self._load_words()
        self._load_presets()
        self._init_ui()

    def _load_words(self):
        """加载 field_words.json"""
        # 使用当前文件所在目录，避免硬编码中文目录名
        words_path = os.path.join(os.path.dirname(__file__), "field_words.json")
        try:
            with open(words_path, 'r', encoding='utf-8') as f:
                self._words_data = json.load(f)
        except Exception:
            self._words_data = dict(self.DEFAULT_WORDS)

    def _load_presets(self):
        """加载 field_presets.json（字段常用句）"""
        presets_path = os.path.join(os.path.dirname(__file__), "field_presets.json")
        try:
            with open(presets_path, 'r', encoding='utf-8') as f:
                self._presets_data = json.load(f)
        except Exception:
            self._presets_data = {}

    def _save_presets(self):
        """保存 field_presets.json"""
        presets_path = os.path.join(os.path.dirname(__file__), "field_presets.json")
        with open(presets_path, 'w', encoding='utf-8') as f:
            json.dump(self._presets_data, f, ensure_ascii=False, indent=2)

    def _init_ui(self):
        layout = QVBoxLayout(self)
        layout.setContentsMargins(0, 2, 0, 2)
        layout.setSpacing(2)

        # 字段标签栏
        tabs_bar = QWidget()
        tabs_layout = QHBoxLayout(tabs_bar)
        tabs_layout.setContentsMargins(0, 0, 0, 0)
        tabs_layout.setSpacing(2)

        # 加载顺序（按病历结构顺序）
        field_order = [
            "主诉", "入院方式", "病史陈述者", "现病史", "既往史", "个人史", "婚育史", "家族史",
            "体格检查", "辅助检查", "初步诊断", "鉴别诊断", "诊疗计划",
            "手术记录", "会诊记录", "抢救记录", "死亡病例讨论", "专科情况"
        ]

        self._tab_buttons = {}
        for field in field_order:
            if field in self._words_data:
                btn = QPushButton(field)
                btn.setCheckable(True)
                btn.setFixedHeight(22)
                btn.setToolTip(self._words_data[field].get("description", ""))
                btn.clicked.connect(lambda checked, f=field: self._on_field_selected(f))
                btn.setStyleSheet("""
                    QPushButton {
                        background: rgba(255,255,255,0.05);
                        color: #b8c5d6;
                        border: 1px solid rgba(0,212,255,0.1);
                        border-radius: 10px;
                        padding: 2px 10px;
                        font-size: 11px;
                    }
                    QPushButton:checked {
                        background: rgba(0,212,255,0.2);
                        color: #00d4ff;
                        border-color: rgba(0,212,255,0.4);
                    }
                    QPushButton:hover {
                        background: rgba(0,212,255,0.1);
                    }
                """)
                tabs_layout.addWidget(btn)
                self._tab_buttons[field] = btn

        tabs_layout.addStretch()
        layout.addWidget(tabs_bar)

        # 常用词滚动区
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        scroll.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        scroll.setMaximumHeight(280)
        scroll.setStyleSheet("""
            QScrollArea {
                border: 1px solid rgba(0,212,255,0.08);
                border-radius: 6px;
                background: rgba(255,255,255,0.02);
            }
            QScrollBar:vertical {
                background: rgba(0,0,0,0.2);
                width: 6px;
                border-radius: 3px;
            }
            QScrollBar::handle:vertical {
                background: rgba(0,212,255,0.3);
                border-radius: 3px;
            }
        """)

        self._terms_container = QWidget()
        self._terms_layout = QVBoxLayout(self._terms_container)
        self._terms_layout.setContentsMargins(6, 6, 6, 6)
        self._terms_layout.setSpacing(6)

        scroll.setWidget(self._terms_container)
        layout.addWidget(scroll)

        # 默认选中第一个字段
        if self._tab_buttons:
            first_field = list(self._tab_buttons.keys())[0]
            self._tab_buttons[first_field].setChecked(True)
            self._show_field_terms(first_field)

    def _on_field_selected(self, field):
        """切换字段"""
        for f, btn in self._tab_buttons.items():
            btn.setChecked(f == field)
        self._current_field = field
        self._show_field_terms(field)

    def _show_field_terms(self, field):
        """展示某字段的常用词"""
        # 清空
        while self._terms_layout.count():
            item = self._terms_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()

        field_data = self._words_data.get(field)
        if not field_data:
            return

        terms = field_data.get("terms", {})

        # 如果 terms 是 dict，按子分类展示
        if isinstance(terms, dict):
            for category, word_list in terms.items():
                # 子分类标签
                cat_label = QLabel(category)
                cat_label.setStyleSheet(
                    "color: #00d4ff; font-size: 11px; font-weight: bold; padding: 2px 0;"
                )
                self._terms_layout.addWidget(cat_label)
                # 词语按钮行
                row = self._create_term_row(word_list, field)
                self._terms_layout.addWidget(row)
        else:
            # 简单列表
            row = self._create_term_row(terms, field)
            self._terms_layout.addWidget(row)

        # ─── 常用句区域 ───
        presets = self._presets_data.get(field, [])
        if presets:
            preset_label = QLabel("📝 常用句（点击插入）")
            preset_label.setStyleSheet(
                "color: #ffa500; font-size: 11px; font-weight: bold; padding: 4px 0 2px;"
            )
            self._terms_layout.addWidget(preset_label)
            for sentence in presets:
                btn = QPushButton(sentence[:30] + ("..." if len(sentence) > 30 else ""))
                btn.setToolTip(sentence)
                btn.setFixedHeight(26)
                btn.setCursor(Qt.PointingHandCursor)
                btn.clicked.connect(lambda checked, s=sentence: self.term_clicked.emit(s))
                btn.setStyleSheet("""
                    QPushButton {
                        background: rgba(255,165,0,0.08);
                        color: #e8d5b0;
                        border: 1px solid rgba(255,165,0,0.2);
                        border-radius: 4px;
                        padding: 3px 8px;
                        font-size: 11px;
                        text-align: left;
                    }
                    QPushButton:hover {
                        background: rgba(255,165,0,0.2);
                        color: #ffa500;
                        border-color: rgba(255,165,0,0.5);
                    }
                """)
                self._terms_layout.addWidget(btn)

        # “+ 添加常用句”按钮
        add_btn = QPushButton("➕ 添加常用句")
        add_btn.setFixedHeight(24)
        add_btn.setCursor(Qt.PointingHandCursor)
        add_btn.clicked.connect(lambda: self._add_preset_for_field(field))
        add_btn.setStyleSheet("""
            QPushButton {
                background: transparent;
                color: #788;
                border: 1px dashed rgba(255,255,255,0.15);
                border-radius: 4px;
                padding: 2px 8px;
                font-size: 10px;
            }
            QPushButton:hover { color: #ffa500; border-color: rgba(255,165,0,0.4); }
        """)
        self._terms_layout.addWidget(add_btn)

        self._terms_layout.addStretch()

    def _create_term_row(self, terms, field):
        """创建一行词语按钮"""
        row = QWidget()
        row_layout = QHBoxLayout(row)
        row_layout.setContentsMargins(0, 0, 0, 0)
        row_layout.setSpacing(4)

        for term in terms:
            if not term:
                continue
            btn = QPushButton(term)
            btn.setFixedHeight(24)
            btn.setCursor(Qt.PointingHandCursor)
            btn.clicked.connect(lambda checked, t=term: self.term_clicked.emit(t))
            btn.setStyleSheet("""
                QPushButton {
                    background: rgba(0,212,255,0.08);
                    color: #c8d6e5;
                    border: 1px solid rgba(0,212,255,0.15);
                    border-radius: 12px;
                    padding: 3px 12px;
                    font-size: 11px;
                }
                QPushButton:hover {
                    background: rgba(0,212,255,0.2);
                    color: #00d4ff;
                    border-color: rgba(0,212,255,0.4);
                }
            """)
            row_layout.addWidget(btn)

        row_layout.addStretch()
        return row

    def _add_preset_for_field(self, field):
        """弹出输入框，为当前字段添加常用句"""
        text, ok = QInputDialog.getMultiLineText(
            self, f"添加常用句 — {field}",
            f"请输入「{field}」的常用句：",
            ""
        )
        if ok and text.strip():
            sentence = text.strip()
            if field not in self._presets_data:
                self._presets_data[field] = []
            if sentence not in self._presets_data[field]:
                self._presets_data[field].append(sentence)
                self._save_presets()
                # 刷新显示
                self._show_field_terms(field)

    def set_current_field(self, field):
        """外部设置当前字段"""
        if field in self._tab_buttons:
            self._on_field_selected(field)


# ==================== 主窗口 ====================
class MedVoiceApp(QMainWindow):
    # 音频文件转写完成信号（text, filename）
    file_transcribed = pyqtSignal(str, str)

    def __init__(self, db=None, current_user=None):
        super().__init__()
        self.db = db
        self.current_user = current_user or {}
        # 当前正在编辑的病历 id（None 表示新病历）
        self.current_record_id = None
        uname = self.current_user.get("username", "")
        title = "星衍AI智能病历录入系统 v1.0"
        if uname:
            role_txt = "管理员" if self.current_user.get("role") == "admin" else "医生"
            title += "  |  当前用户：%s（%s）" % (uname, role_txt)
        self.setWindowTitle(title)
        self.setGeometry(100, 100, 1200, 800)

        # 核心引擎
        self.rule_engine = RuleEngine()
        self.corrector = Corrector(rule_engine=self.rule_engine)
        self.template_engine = TemplateEngine()
        self.parser = SectionParser()
        self.smart_dictation = SmartDictation(self.parser)
        self.classifier = MedicalClassifier()

        # AI 辅助诊断（基于知识图谱）
        self.diagnosis_assistant = DiagnosisAssistant()
        self.qa_engine = KnowledgeQA()
        self.diagnosis_thread = None

        # 常用语句库
        self.phrase_lib = PhraseLibrary()
        self._phrase_dialog = None

        # 纠错反馈收集（用于 LM 迭代训练）
        self.feedback = CorrectionFeedback()
        # 统一纠错记忆库（用于 Top-K / prompt 导出 / LM 迭代）
        self.memory = None
        # M3 Top-K 术语引擎
        self.topk_engine = None

        # 崩溃日志
        self.crash_logger = CrashLogger()
        self.crash_logger.log_event("应用启动")

        # 加载语音模型
        model_path = os.path.join(os.path.dirname(__file__), "model")
        self.asr = ASREngine(model_path=model_path)
        self.listen_thread = None

        # 状态
        self.current_dept = "通用"
        self.asr.set_hotwords("通用")
        self.is_listening = False
        self.partial_text = ""
        self._auto_stop_timer = None
        # 语音命令解析器
        self.voice_command = VoiceCommandParser()
        # 录音时长计时
        self._record_start_ts = None
        self._duration_timer = QTimer(self)
        self._duration_timer.setInterval(500)
        self._duration_timer.timeout.connect(self._update_record_duration)

        self._init_ui()
        self._apply_dark_theme()

        # 默认加载通用词库
        self.corrector.set_department("通用")
        self._load_departments()

        # 光标移动时自动检测当前字段
        self.text_edit.cursorPositionChanged.connect(self._on_cursor_moved)

        # 启动时：自动备份数据库 + 从历史病历预热用户热词（移到后台线程，避免阻塞启动）
        self._startup_maintenance_thread = threading.Thread(
            target=self._startup_maintenance,
            daemon=True,
            name="StartupMaintenance"
        )
        self._startup_maintenance_thread.start()

        # 首次使用新手引导
        self._guide = OnboardingGuide.show_if_needed(self)

        # 音频文件转写完成信号
        self.file_transcribed.connect(self._on_file_transcribed)

        # 模板搜索过滤
        self.template_combo.lineEdit().textEdited.connect(self._filter_templates)

    def _startup_maintenance(self):
        """启动维护：数据库自动备份 + 从历史病历提取高频词作为用户热词"""
        if not self.db:
            return
        try:
            path = self.db.auto_backup()
            self.crash_logger.log_event("数据库自动备份: %s" % path)
        except Exception as e:
            print(f"[Main] 自动备份失败: {e}")
        # 从历史病历学习高频医学词 → 用户自适应热词
        try:
            self._refresh_user_hotwords(silent=True)
        except Exception as e:
            print(f"[Main] 用户热词预热失败: {e}")

    def _refresh_user_hotwords(self, silent=False):
        """从当前用户历史病历中提取高频专业词，更新 ASR 用户自适应热词"""
        if not self.db:
            return
        user_id = None if self.current_user.get("role") == "admin" else self.current_user.get("id")
        records = self.db.list_records(user_id=user_id)
        if not records:
            if not silent:
                self.status_bar.showMessage("暂无历史病历，无法提取个人热词")
            return
        from collections import Counter
        counter = Counter()
        # 用分类器/解析器已有的医学词典做候选，统计 2-6 字中文词片段
        for r in records[:200]:
            content = r.get("content", "") or ""
            # 抽取词典中出现过的活跃词
            for w in self.corrector.active_words:
                if len(w) >= 2 and w in content:
                    counter[w] += 1
        # 取高频前 200，出现次数≥2
        common = [w for w, c in counter.most_common(200) if c >= 2]
        if common:
            self.asr.update_user_hotwords(common)
            # 重新按当前科室加载热词以生效
            self.asr.set_hotwords(self.current_dept)
            self._refresh_topk_hotwords(silent=True)
            if not silent:
                self.status_bar.showMessage("已从 %d 份病历提取 %d 个个人高频词并加入热词" % (len(records), len(common)))
        elif not silent:
            self.status_bar.showMessage("未提取到足够的高频词")

    def _refresh_topk_hotwords(self, silent=False):
        """基于记忆库 Top-K 术语刷新 ASR 热词"""
        try:
            topk = self._get_topk_engine()
            if topk:
                topk.refresh_asr_hotwords(
                    self.asr,
                    dept=self.current_dept,
                    doctor_id=self.current_user.get("id") if isinstance(self.current_user, dict) else None,
                    term_budget=300,
                    postprocess_budget=120,
                )
                self.asr.set_hotwords(self.current_dept)
                if not silent:
                    self.status_bar.showMessage("已基于记忆库刷新 Top-K 术语热词")
        except Exception as e:
            print(f"[Main] 刷新 Top-K 热词失败: {e}")

    def _on_cursor_moved(self):
        """光标移动时，检测当前所在的病历字段"""
        cursor = self.text_edit.textCursor()
        pos = cursor.position()
        
        # 优化：只在跨行时重新检测，避免每次光标移动都重新计算
        if not hasattr(self, '_last_cursor_line'):
            self._last_cursor_line = -1
            self._last_detected_field = None
        
        # 获取当前行号
        block = cursor.block()
        current_line = block.blockNumber()
        
        # 如果行号没变，使用缓存结果
        if current_line == self._last_cursor_line:
            field = self._last_detected_field
        else:
            text = self.text_edit.toPlainText()
            field = self._detect_field_at_position(text, pos)
            self._last_cursor_line = current_line
            self._last_detected_field = field
        
        if field and field != self.field_panel._current_field:
            self.field_panel.set_current_field(field)

    def _detect_field_at_position(self, text, pos):
        """检测给定位置属于哪个字段"""
        best_field = None
        best_pos = -1

        for keyword, standard_field in self.parser.keyword_to_field.items():
            # 优化：使用 finditer(text, 0, pos) 避免创建子串
            pattern = re.compile(re.escape(keyword) + r'[：: \t]*')
            for m in pattern.finditer(text, 0, pos):
                if m.end() > best_pos:
                    best_pos = m.end()
                    best_field = standard_field

        return best_field

    def _insert_term_at_cursor(self, term):
        """将选中的常用词插入到编辑器光标位置"""
        cursor = self.text_edit.textCursor()

        # 如果当前在空行或字段开头，直接插入
        # 否则在当前光标位置插入，并加上合适的标点
        cursor.insertText(term)
        self.text_edit.setTextCursor(cursor)
        self.text_edit.setFocus()

    def _init_ui(self):
        """初始化界面"""
        # 工具栏
        toolbar = QToolBar()
        self.addToolBar(toolbar)

        # 科室选择
        toolbar.addWidget(QLabel("  科室："))
        self.dept_combo = QComboBox()
        self.dept_combo.setMinimumWidth(120)
        self.dept_combo.currentTextChanged.connect(self._on_dept_changed)
        toolbar.addWidget(self.dept_combo)

        toolbar.addSeparator()

        # 模板选择（可编辑搜索）
        toolbar.addWidget(QLabel("  模板："))
        self.template_combo = QComboBox()
        self.template_combo.setMinimumWidth(160)
        self.template_combo.setEditable(True)
        self.template_combo.setInsertPolicy(QComboBox.NoInsert)
        self.template_combo.lineEdit().setPlaceholderText("搜索模板...")
        self.template_combo.currentTextChanged.connect(self._on_template_changed)
        toolbar.addWidget(self.template_combo)

        toolbar.addSeparator()

        # 录音按钮
        self.record_btn = QPushButton("🎤 开始录音")
        self.record_btn.setCheckable(True)
        self.record_btn.clicked.connect(self._toggle_recording)
        self.record_btn.setStyleSheet("""
            QPushButton {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #00d4ff, stop:1 #0066ff);
                color: #0a0e27;
                font-weight: bold;
                padding: 8px 20px;
                border-radius: 20px;
                font-size: 14px;
            }
            QPushButton:checked {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #ff4444, stop:1 #cc0000);
                color: white;
            }
            QPushButton:hover {
                padding: 8px 25px;
            }
        """)
        toolbar.addWidget(self.record_btn)

        # 录音状态指示器（脉冲动画）
        self.recording_indicator = RecordingIndicator()
        toolbar.addWidget(self.recording_indicator)

        # 录音模式选择
        mode_label = QLabel("  模式：")
        toolbar.addWidget(mode_label)
        self.record_mode_combo = QComboBox()
        self.record_mode_combo.addItems(["手动停止", "连续录音（60s）", "连续录音（120s）"])
        self.record_mode_combo.setToolTip("手动停止：点击按钮开始/结束\n连续录音：到达时长自动停止")
        self.record_mode_combo.setMinimumWidth(140)
        toolbar.addWidget(self.record_mode_combo)

        # 麦克风设备选择
        mic_label = QLabel("  麦克风：")
        toolbar.addWidget(mic_label)
        self.mic_combo = QComboBox()
        self.mic_combo.setMinimumWidth(150)
        self.mic_combo.setToolTip("选择录音输入设备")
        self._load_microphones()
        self.mic_combo.currentIndexChanged.connect(self._on_mic_changed)
        toolbar.addWidget(self.mic_combo)

        toolbar.addSeparator()

        # 纠错按钮
        correct_btn = QPushButton("✨ 纠错")
        correct_btn.clicked.connect(self._run_correction)
        toolbar.addWidget(correct_btn)

        # 清除按钮
        clear_btn = QPushButton("🗑 清除")
        clear_btn.clicked.connect(self._clear_text)
        toolbar.addWidget(clear_btn)

        # 导出按钮
        save_btn = QPushButton("💾 导出")
        save_btn.clicked.connect(self._save_text)
        toolbar.addWidget(save_btn)

        # 复制全文按钮（Ctrl+Shift+C）
        copy_btn = QPushButton("📋 复制全文")
        copy_btn.setToolTip("复制病历全文到剪贴板，便于粘贴进 HIS 系统（Ctrl+Shift+C）")
        copy_btn.clicked.connect(self._copy_all_text)
        toolbar.addWidget(copy_btn)

        toolbar.addSeparator()

        # 常用语句库按钮（F3）
        phrase_btn = QPushButton("💬 常用语")
        phrase_btn.setToolTip("打开常用语句库，一键插入常用短语（F3）")
        phrase_btn.clicked.connect(self._open_phrase_library)
        toolbar.addWidget(phrase_btn)

        # 首页→病程 自动填充按钮
        autofill_btn = QPushButton("📋 首页→病程")
        autofill_btn.setToolTip("将当前入院记录/首页病程的内容自动填入首次病程记录")
        autofill_btn.clicked.connect(self._autofill_progress_note)
        toolbar.addWidget(autofill_btn)

        # 一键套用按钮
        apply_btn = QPushButton("⚡ 一键套用")
        apply_btn.setToolTip("语音输入核心信息，自动替换模板中的占位符 X")
        apply_btn.clicked.connect(self._smart_apply_template)
        toolbar.addWidget(apply_btn)

        # “更多”溢出菜单（低频功能收纳，减少工具栏拥挤）
        more_btn = QToolButton()
        more_btn.setText("更多 ⌄")
        more_btn.setToolTip("模板管理 / 规则管理 / 结构化 / 崩溃日志")
        more_btn.setPopupMode(QToolButton.InstantPopup)
        more_menu = QMenu(more_btn)
        more_menu.addAction("📝 模板管理", self._open_template_manager)
        more_menu.addAction("📏 规则管理", self._open_rule_manager)
        more_menu.addAction("📋 结构化解析", self._open_struct_view)
        more_menu.addSeparator()
        more_menu.addAction("🧠 重训语言模型", self._retrain_lm)
        more_menu.addAction("📋 崩溃日志", self._view_crash_log)
        more_btn.setMenu(more_menu)
        more_btn.setStyleSheet("""
            QToolButton {
                padding: 6px 12px;
                border-radius: 6px;
            }
            QToolButton:hover { background: rgba(0,212,255,0.1); }
        """)
        toolbar.addWidget(more_btn)

        toolbar.addSeparator()

        # 保存病历按钮（写入数据库，Ctrl+S）
        save_record_btn = QPushButton("💾 保存病历")
        save_record_btn.setToolTip("将当前病历保存到病历库（Ctrl+S）")
        save_record_btn.clicked.connect(self._save_record)
        toolbar.addWidget(save_record_btn)

        # 病历库按钮
        record_lib_btn = QPushButton("📚 病历库")
        record_lib_btn.clicked.connect(self._open_record_manager)
        toolbar.addWidget(record_lib_btn)

        # 数据备份/恢复按钮
        backup_btn = QPushButton("🛡 备份")
        backup_btn.setToolTip("备份/恢复病历数据库，或刷新个人热词")
        backup_btn.clicked.connect(self._open_backup_menu)
        toolbar.addWidget(backup_btn)

        # 用户管理按钮（仅管理员）
        if self.current_user.get("role") == "admin":
            user_mgr_btn = QPushButton("👥 用户管理")
            user_mgr_btn.clicked.connect(self._open_user_manager)
            toolbar.addWidget(user_mgr_btn)

        # 保存病历快捷键 Ctrl+S
        save_shortcut = QAction(self)
        save_shortcut.setShortcut("Ctrl+S")
        save_shortcut.triggered.connect(self._save_record)
        self.addAction(save_shortcut)

        # 复制全文快捷键 Ctrl+Shift+C
        copy_shortcut = QAction(self)
        copy_shortcut.setShortcut("Ctrl+Shift+C")
        copy_shortcut.triggered.connect(self._copy_all_text)
        self.addAction(copy_shortcut)

        # 录音开始/停止快捷键 F2
        record_shortcut = QAction(self)
        record_shortcut.setShortcut("F2")
        record_shortcut.triggered.connect(self._toggle_recording)
        self.addAction(record_shortcut)

        # 录音快捷键 Ctrl+R（更直觉）
        record_shortcut2 = QAction(self)
        record_shortcut2.setShortcut("Ctrl+R")
        record_shortcut2.triggered.connect(self._toggle_recording)
        self.addAction(record_shortcut2)

        # 导出快捷键 Ctrl+E
        export_shortcut = QAction(self)
        export_shortcut.setShortcut("Ctrl+E")
        export_shortcut.triggered.connect(self._save_text)
        self.addAction(export_shortcut)

        # 专注模式快捷键 F11
        focus_shortcut = QAction(self)
        focus_shortcut.setShortcut("F11")
        focus_shortcut.triggered.connect(self._toggle_focus_mode)
        self.addAction(focus_shortcut)

        # 纠错面板显隐快捷键 F9
        panel_shortcut = QAction(self)
        panel_shortcut.setShortcut("F9")
        panel_shortcut.triggered.connect(self._toggle_left_panel)
        self.addAction(panel_shortcut)

        # 常用语句库快捷键 F3
        phrase_shortcut = QAction(self)
        phrase_shortcut.setShortcut("F3")
        phrase_shortcut.triggered.connect(self._open_phrase_library)
        self.addAction(phrase_shortcut)

        # 一键纠错快捷键 F4
        correct_shortcut = QAction(self)
        correct_shortcut.setShortcut("F4")
        correct_shortcut.triggered.connect(self._run_correction)
        self.addAction(correct_shortcut)

        # 中央部件
        central = QWidget()
        self.setCentralWidget(central)
        main_layout = QHBoxLayout(central)
        main_layout.setContentsMargins(10, 10, 10, 10)
        main_layout.setSpacing(10)

        # 分割器
        splitter = QSplitter(Qt.Horizontal)
        self.splitter = splitter

        # 左侧面板 - 纠错日志
        left_panel = QWidget()
        self.left_panel = left_panel
        left_layout = QVBoxLayout(left_panel)
        left_layout.setContentsMargins(0, 0, 0, 0)

        left_layout.addWidget(QLabel("📋 纠错日志"))

        # 筛选按钮组
        filter_bar = QWidget()
        filter_layout = QHBoxLayout(filter_bar)
        filter_layout.setContentsMargins(0, 0, 0, 5)
        filter_layout.setSpacing(5)

        self.filter_typo = QCheckBox("🔤 错别字")
        self.filter_logic = QCheckBox("🧠 逻辑错误")
        self.filter_missing = QCheckBox("⚠️ 缺项提醒")
        self.filter_typo.setChecked(True)
        self.filter_logic.setChecked(True)
        self.filter_missing.setChecked(True)

        # 筛选按钮样式
        for cb in [self.filter_typo, self.filter_logic, self.filter_missing]:
            cb.setStyleSheet("""
                QCheckBox {
                    color: #b8c5d6;
                    font-size: 11px;
                    spacing: 3px;
                }
                QCheckBox::indicator {
                    width: 14px;
                    height: 14px;
                    border: 1px solid rgba(0,212,255,0.3);
                    border-radius: 3px;
                    background: rgba(0,212,255,0.05);
                }
                QCheckBox::indicator:checked {
                    background: rgba(0,212,255,0.3);
                    border-color: #00d4ff;
                }
            """)
            cb.stateChanged.connect(self._apply_filter)
            filter_layout.addWidget(cb)

        filter_layout.addStretch()

        # 统计标签
        self.stats_label = QLabel("")
        self.stats_label.setStyleSheet("color: #6b8a9a; font-size: 10px;")
        filter_layout.addWidget(self.stats_label)

        left_layout.addWidget(filter_bar)

        # 纠错日志列表
        self.log_list = QListWidget()
        self.log_list.setStyleSheet("""
            QListWidget {
                background: rgba(255,255,255,0.03);
                border: 1px solid rgba(0,212,255,0.1);
                border-radius: 8px;
                padding: 5px;
                font-size: 12px;
            }
            QListWidget::item {
                padding: 6px;
                border-bottom: 1px solid rgba(0,212,255,0.05);
            }
        """)
        self.log_list.itemClicked.connect(self._on_log_item_clicked)
        left_layout.addWidget(self.log_list)

        # 接受/拒绝纠错按钮栏
        action_bar = QWidget()
        action_layout = QHBoxLayout(action_bar)
        action_layout.setContentsMargins(0, 5, 0, 0)
        action_layout.setSpacing(6)

        accept_btn = QPushButton("✓ 接受")
        accept_btn.clicked.connect(self._accept_correction)
        accept_btn.setEnabled(False)
        accept_btn.setToolTip("接受当前选中的纠错建议")
        accept_btn.setStyleSheet("""
            QPushButton {
                background: rgba(81, 207, 102, 0.15);
                color: #51cf66;
                padding: 5px 14px;
                border-radius: 12px;
                border: 1px solid rgba(81, 207, 102, 0.3);
                font-size: 11px;
            }
            QPushButton:hover { background: rgba(81, 207, 102, 0.25); }
            QPushButton:disabled { color: #444; border-color: #333; }
        """)

        accept_all_btn = QPushButton("✓✓ 全部接受")
        accept_all_btn.clicked.connect(self._accept_all_corrections)
        accept_all_btn.setToolTip("一键接受所有纠错建议")
        accept_all_btn.setStyleSheet("""
            QPushButton {
                background: rgba(81, 207, 102, 0.1);
                color: #51cf66;
                padding: 5px 10px;
                border-radius: 12px;
                border: 1px solid rgba(81, 207, 102, 0.2);
                font-size: 11px;
            }
            QPushButton:hover { background: rgba(81, 207, 102, 0.2); }
        """)

        reject_btn = QPushButton("✗ 拒绝")
        reject_btn.clicked.connect(self._reject_correction)
        reject_btn.setEnabled(False)
        reject_btn.setToolTip("拒绝当前选中的纠错，系统将记住此偏好")
        reject_btn.setStyleSheet("""
            QPushButton {
                background: rgba(255, 107, 107, 0.15);
                color: #ff6b6b;
                padding: 5px 14px;
                border-radius: 12px;
                border: 1px solid rgba(255, 107, 107, 0.3);
                font-size: 11px;
            }
            QPushButton:hover { background: rgba(255, 107, 107, 0.25); }
            QPushButton:disabled { color: #444; border-color: #333; }
        """)

        self.log_hint = QLabel("点击日志条目查看详情")
        self.log_hint.setStyleSheet("color: #6b8a9a; font-size: 10px;")

        action_layout.addWidget(accept_btn)
        action_layout.addWidget(accept_all_btn)
        action_layout.addWidget(reject_btn)
        action_layout.addStretch()
        action_layout.addWidget(self.log_hint)

        left_layout.addWidget(action_bar)
        self._accept_btn = accept_btn
        self._reject_btn = reject_btn

        # 存储完整纠错数据
        self._all_logs = []

        splitter.addWidget(left_panel)

        # 右侧 - 文本编辑区
        right_panel = QWidget()
        right_layout = QVBoxLayout(right_panel)
        right_layout.setContentsMargins(0, 0, 0, 0)

        # 实时识别显示
        self.partial_label = QLabel("等待输入...")
        self.partial_label.setStyleSheet("""
            color: #00d4ff;
            font-size: 13px;
            padding: 5px;
            background: rgba(0,212,255,0.05);
            border-radius: 5px;
        """)
        self.partial_label.setWordWrap(True)
        self.partial_label.setMaximumHeight(120)
        right_layout.addWidget(self.partial_label)

        # 悬浮识别预览面板（录音中实时显示 + 识别后接受/拒绝/重听确认）
        self.asr_preview = AsrPreviewPanel(self)
        self.asr_preview.accepted.connect(self._on_preview_accept)
        self.asr_preview.rejected.connect(self._on_preview_reject)
        self.asr_preview.retried.connect(self._on_preview_retry)
        self._pending_asr_text = ''  # 待确认的识别结果

        # 录音事件处理器（录音/识别/预览确认逻辑，已从 main.py 拆分）
        self.recorder = RecordingHandler(self)

        # 实时录音波形图（录音时显示滚动音量柱状图）
        self.waveform = WaveformWidget()
        self.waveform.setVisible(False)
        right_layout.addWidget(self.waveform)

        # 波形轮询定时器（~30fps 从 ASR 引擎读取实时音量电平）
        self._wave_timer = QTimer(self)
        self._wave_timer.setInterval(33)
        self._wave_timer.timeout.connect(self._poll_audio_level)

        # 字段常用词面板
        self.field_panel = FieldWordsPanel()
        self.field_panel.setMaximumHeight(340)
        self.field_panel.term_clicked.connect(self._insert_term_at_cursor)
        right_layout.addWidget(self.field_panel)

        # 文本编辑区
        self.text_edit = QTextEdit()
        self.text_edit.setPlaceholderText(
            "选择模板开始，或直接输入病历内容。\n"
            "点击「开始录音」用语音输入，系统会自动纠错。"
        )
        self.text_edit.setStyleSheet("""
            QTextEdit {
                background: rgba(255,255,255,0.03);
                border: 1px solid rgba(0,212,255,0.15);
                border-radius: 10px;
                padding: 15px;
                font-size: 14px;
                line-height: 1.8;
                color: #e0e0e0;
            }
            QTextEdit:focus {
                border: 1px solid rgba(0,212,255,0.4);
            }
        """)
        right_layout.addWidget(self.text_edit)

        # 音文对照播放器（录音结束后自动加载，可回放校对）
        self.audio_player_widget = AudioPlayerWidget()
        right_layout.addWidget(self.audio_player_widget)

        # 字段名语法高亮
        self.field_highlighter = FieldHighlighter(self.text_edit.document())

        # 支持拖拽音频文件转写
        self.setAcceptDrops(True)

        # ==================== AI 辅助诊断面板（可折叠） ====================
        ai_group = QGroupBox("🔬 AI 辅助诊断")
        self.ai_group = ai_group
        ai_group.setStyleSheet("""
            QGroupBox {
                color: #00d4ff;
                font-size: 13px;
                font-weight: bold;
                border: 1px solid rgba(0,212,255,0.2);
                border-radius: 10px;
                margin-top: 10px;
                padding-top: 8px;
            }
            QGroupBox::title {
                subcontrol-origin: margin;
                left: 12px;
                padding: 0 5px;
            }
        """)
        ai_layout = QVBoxLayout(ai_group)
        ai_layout.setContentsMargins(10, 8, 10, 10)

        ai_header = QHBoxLayout()
        self.ai_analyze_btn = QPushButton("🔬 分析当前病历")
        self.ai_analyze_btn.clicked.connect(self._run_diagnosis)
        self.ai_analyze_btn.setStyleSheet("""
            QPushButton {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #00d4ff, stop:1 #0088cc);
                color: #0a0e27;
                padding: 6px 16px;
                border-radius: 12px;
                font-size: 12px;
                font-weight: bold;
            }
            QPushButton:hover { background: #00d4ff; }
            QPushButton:disabled { background: #333; color: #666; }
        """)
        qa_btn = QPushButton("💡 知识问答")
        qa_btn.clicked.connect(self._show_qa_dialog)
        qa_btn.setStyleSheet("""
            QPushButton {
                background: qlineargradient(x1:0,y1:0,x2:1,y2:0, stop:0 #ff9a5b, stop:1 #f64236);
                color: #fff;
                padding: 6px 16px;
                border-radius: 12px;
                font-size: 12px;
                font-weight: bold;
            }
            QPushButton:hover { background: #ffaf7b; }
        """)
        self.ai_status_label = QLabel("")
        self.ai_status_label.setStyleSheet("color: #6b8a9a; font-size: 10px;")
        ai_header.addWidget(self.ai_analyze_btn)
        ai_header.addWidget(qa_btn)
        ai_header.addWidget(self.ai_status_label)
        ai_header.addStretch()

        # AI 面板折叠按钮
        self.ai_collapse_btn = QPushButton("▲ 收起")
        self.ai_collapse_btn.setFixedWidth(64)
        self.ai_collapse_btn.setToolTip("折叠/展开 AI 诊断结果区域")
        self.ai_collapse_btn.setStyleSheet("""
            QPushButton {
                background: transparent;
                color: #6b8a9a;
                border: 1px solid rgba(0,212,255,0.2);
                border-radius: 8px;
                padding: 3px 8px;
                font-size: 11px;
            }
            QPushButton:hover { color: #00d4ff; border-color: rgba(0,212,255,0.5); }
        """)
        self.ai_collapse_btn.clicked.connect(self._toggle_ai_panel)
        ai_header.addWidget(self.ai_collapse_btn)
        ai_layout.addLayout(ai_header)

        self.ai_result = QTextBrowser()
        self.ai_result.setOpenExternalLinks(False)
        self.ai_result.setStyleSheet("""
            QTextBrowser {
                background: rgba(255,255,255,0.03);
                border: 1px solid rgba(0,212,255,0.15);
                border-radius: 8px;
                padding: 10px;
                font-size: 13px;
                color: #e0e0e0;
            }
        """)
        self.ai_result.setPlaceholderText("点击「分析当前病历」，AI 将基于知识图谱给出可能诊断、用药审查、检查建议与风险预警。")
        self.ai_result.setMinimumHeight(160)
        ai_layout.addWidget(self.ai_result)

        right_layout.addWidget(ai_group)

        splitter.addWidget(right_panel)
        splitter.setSizes([300, 900])

        main_layout.addWidget(splitter)

        # 状态栏
        self.status_bar = QStatusBar()
        self.setStatusBar(self.status_bar)
        self.status_bar.showMessage("就绪 | 请先选择科室和模板")

        # 试用期信息
        self._license_mgr = LicenseManager()
        trial_label = TrialInfoBar.create_label(self._license_mgr)
        if trial_label.text():
            self.status_bar.addPermanentWidget(trial_label)

    def _apply_dark_theme(self):
        """深色主题"""
        app = QApplication.instance()
        app.setStyle("Fusion")
        palette = QPalette()
        palette.setColor(QPalette.Window, QColor(10, 14, 39))
        palette.setColor(QPalette.WindowText, QColor(255, 255, 255))
        palette.setColor(QPalette.Base, QColor(15, 20, 50))
        palette.setColor(QPalette.AlternateBase, QColor(20, 25, 60))
        palette.setColor(QPalette.ToolTipBase, QColor(255, 255, 255))
        palette.setColor(QPalette.ToolTipText, QColor(255, 255, 255))
        palette.setColor(QPalette.Text, QColor(220, 220, 220))
        palette.setColor(QPalette.Button, QColor(30, 40, 80))
        palette.setColor(QPalette.ButtonText, QColor(255, 255, 255))
        palette.setColor(QPalette.BrightText, QColor(255, 0, 0))
        palette.setColor(QPalette.Link, QColor(0, 212, 255))
        palette.setColor(QPalette.Highlight, QColor(0, 212, 255))
        palette.setColor(QPalette.HighlightedText, QColor(10, 14, 39))
        app.setPalette(palette)

    def _load_departments(self):
        """加载科室列表"""
        depts = self.template_engine.get_departments()
        self.dept_combo.clear()
        self.dept_combo.addItems(depts)

    def _on_dept_changed(self, dept):
        """科室切换"""
        self.current_dept = dept
        self.corrector.set_department(dept)
        self.asr.set_hotwords(dept)
        self._refresh_topk_hotwords(silent=True)

        # 更新模板列表
        self.template_combo.blockSignals(True)
        self.template_combo.clear()
        templates = self.template_engine.get_templates(dept)
        self._all_template_names = [t["name"] for t in templates]
        for name in self._all_template_names:
            self.template_combo.addItem(name)
        self.template_combo.blockSignals(False)

        self.status_bar.showMessage(f"当前科室：{dept} | 词库已更新")

    def _filter_templates(self, search_text):
        """模板搜索过滤（输入时实时筛选）"""
        self.template_combo.blockSignals(True)
        self.template_combo.clear()
        if not search_text.strip():
            # 空搜索→显示全部
            for name in getattr(self, '_all_template_names', []):
                self.template_combo.addItem(name)
        else:
            keyword = search_text.strip().lower()
            for name in getattr(self, '_all_template_names', []):
                if keyword in name.lower():
                    self.template_combo.addItem(name)
        self.template_combo.blockSignals(False)
        # 如果只剩一个结果，自动选中
        if self.template_combo.count() == 1:
            self.template_combo.setCurrentIndex(0)

    def _on_template_changed(self, template_name):
        """模板选择"""
        if not template_name:
            return
        content = self.template_engine.get_template(
            self.current_dept, template_name
        )
        if content:
            # 修复：如果编辑器已有内容，提示用户是否覆盖，避免数据丢失
            current_text = self.text_edit.toPlainText().strip()
            if current_text:
                reply = QMessageBox.question(
                    self, "加载模板",
                    "当前编辑器已有内容，是否用新模板覆盖？",
                    QMessageBox.Yes | QMessageBox.No,
                    QMessageBox.No
                )
                if reply == QMessageBox.No:
                    return
            self.text_edit.setPlainText(content)
            self.status_bar.showMessage(
                f"已加载模板：{template_name} | 可以开始语音输入"
            )

    def _recommend_template(self, text):
        """根据当前文本内容，在所有科室模板中推荐最匹配的一个（仅状态栏提示，不自动切换）"""
        try:
            if not text or len(text.strip()) < 4:
                return
            best = None  # (score, dept, name)
            for dept in self.template_engine.get_departments():
                for t in self.template_engine.get_templates(dept):
                    name = t.get("name", "")
                    if not name:
                        continue
                    # 模板名（如疾病名）直接出现在文本中 → 强信号
                    score = 0
                    if name in text:
                        score += 3
                    # 拆分名称关键字（去掉“-中医”等后缀）逐字匹配
                    core = name.split("-")[0].replace("【中医】", "")
                    if core and core in text and core != name:
                        score += 2
                    if score > 0 and (best is None or score > best[0]):
                        best = (score, dept, name)
            if best and best[0] >= 2:
                _, dept, name = best
                # 已经选中该模板则不提示
                if self.template_combo.currentText() != name:
                    self.status_bar.showMessage(
                        "💡 推荐模板：%s - %s（可在上方模板下拉选择）" % (dept, name)
                    )
        except Exception as e:
            print(f"[UI] 模板推荐失败: {e}")

    def _autofill_progress_note(self):
        """将当前入院记录/首页病程的内容自动填入首次病程记录"""
        current_text = self.text_edit.toPlainText().strip()
        if not current_text:
            QMessageBox.warning(self, "提示", "当前编辑器无内容，请先填写入院记录/首页病程。")
            return

        # 解析当前文本中的字段
        sections = self.parser.parse(current_text)
        if not sections:
            QMessageBox.warning(self, "提示", "未能从当前文本中解析出字段。")
            return

        # 构建病例特点（从主诉+现病史+体格检查+辅助检查提取）
        case_features = []
        # 基本信息
        basic_info = []
        for field in ['性别', '年龄', '民族']:
            if field in sections and sections[field].strip():
                basic_info.append(f"{field}：{sections[field].strip()}")
        if basic_info:
            case_features.append("，".join(basic_info))

        if '主诉' in sections and sections['主诉'].strip():
            case_features.append(f"主诉：{sections['主诉'].strip()}")
        if '现病史' in sections and sections['现病史'].strip():
            case_features.append(f"现病史：{sections['现病史'].strip()}")
        if '体格检查' in sections and sections['体格检查'].strip():
            case_features.append(f"体格检查：{sections['体格检查'].strip()}")
        if '辅助检查' in sections and sections['辅助检查'].strip():
            case_features.append(f"辅助检查：{sections['辅助检查'].strip()}")

        # 构建首次病程记录
        from datetime import datetime
        now_str = datetime.now().strftime("%Y-%m-%d %H:%M")

        progress_content = f"""日期/时间：{now_str}
病例特点：
{chr(10).join(case_features) if case_features else ''}
诊断依据：
{sections.get('诊断依据', '').strip()}
鉴别诊断：
{sections.get('鉴别诊断', '').strip()}
入院诊断：
{sections.get('初步诊断', sections.get('入院诊断', '')).strip()}
诊疗计划：
{sections.get('诊疗计划', '').strip()}
主治医师意见：
"""

        # 确认是否生成
        reply = QMessageBox.question(
            self, "首页→病程",
            f"已从当前文本提取到 {len(sections)} 个字段，是否生成首次病程记录？",
            QMessageBox.Yes | QMessageBox.No,
            QMessageBox.Yes
        )
        if reply == QMessageBox.Yes:
            self.text_edit.setPlainText(progress_content)
            self.status_bar.showMessage("已生成首次病程记录（从首页病程自动填充）")

    def _smart_apply_template(self):
        """智能套用模板：语音输入核心信息，自动替换模板中的X占位符"""
        import re
        current_text = self.text_edit.toPlainText().strip()
        if not current_text:
            QMessageBox.warning(self, "提示", "请先选择常见病模板。")
            return

        # 检查是否有X占位符
        x_count = len(re.findall(r'X+', current_text))
        if x_count == 0:
            QMessageBox.information(self, "提示", "当前模板中没有发现占位符 X，无需替换。")
            return

        # 创建输入对话框
        dialog = QDialog(self)
        dialog.setWindowTitle("一键套用 - 填写患者核心信息")
        dialog.setMinimumWidth(500)
        dialog.setStyleSheet("""
            QDialog { background: #2b2b2b; color: #e0e0e0; }
            QLabel { color: #e0e0e0; font-size: 13px; }
            QLineEdit { 
                background: #3c3c3c; color: #e0e0e0; border: 1px solid #555;
                padding: 6px; border-radius: 4px; font-size: 13px;
            }
            QLineEdit:focus { border: 1px solid #4a9eff; }
            QPushButton { 
                background: #4a9eff; color: white; border: none;
                padding: 8px 20px; border-radius: 4px; font-size: 13px;
            }
            QPushButton:hover { background: #3a8eef; }
            QPushButton#cancelBtn { background: #555; }
            QPushButton#cancelBtn:hover { background: #666; }
        """)

        layout = QVBoxLayout(dialog)

        # 提示信息
        info_label = QLabel(f"当前模板中有 {x_count} 处占位符需要填写。")
        info_label.setStyleSheet("color: #4a9eff; font-size: 14px; font-weight: bold;")
        layout.addWidget(info_label)

        hint_label = QLabel("请填写患者核心信息（留空则保留原占位符）：")
        layout.addWidget(hint_label)

        # 定义常用字段
        fields = [
            ("姓名", ""), ("性别", ""), ("年龄", ""),
            ("病程时间（如：3天、2周）", ""),
            ("血压（如：160/100）", ""),
            ("体温（如：38.5）", ""),
            ("心率（如：80）", ""),
            ("白细胞（如：12）", ""),
            ("左右侧（左/右）", ""),
            ("部位（如：股骨、胫骨）", ""),
        ]

        line_edits = {}
        for label_text, default in fields:
            row = QHBoxLayout()
            label = QLabel(label_text + "：")
            label.setFixedWidth(180)
            edit = QLineEdit(default)
            edit.setPlaceholderText("留空则保留原占位符")
            row.addWidget(label)
            row.addWidget(edit)
            layout.addLayout(row)
            line_edits[label_text] = edit

        # 按钮
        btn_layout = QHBoxLayout()
        ok_btn = QPushButton("✅ 替换")
        cancel_btn = QPushButton("取消")
        cancel_btn.setObjectName("cancelBtn")
        btn_layout.addStretch()
        btn_layout.addWidget(ok_btn)
        btn_layout.addWidget(cancel_btn)
        layout.addLayout(btn_layout)

        ok_btn.clicked.connect(dialog.accept)
        cancel_btn.clicked.connect(dialog.reject)

        if dialog.exec_() != QDialog.Accepted:
            return

        # 收集用户输入
        values = {}
        for label_text, edit in line_edits.items():
            val = edit.text().strip()
            if val:
                values[label_text] = val

        if not values:
            QMessageBox.information(self, "提示", "未填写任何信息，模板未修改。")
            return

        # 执行替换
        result = current_text

        # 1. 基本信息替换
        if '姓名' in values:
            result = re.sub(r'姓名：\s*', f'姓名：{values["姓名"]}  ', result)
        if '性别' in values:
            result = re.sub(r'性别：\s*', f'性别：{values["性别"]}  ', result)
        if '年龄' in values:
            result = re.sub(r'年龄：\s*', f'年龄：{values["年龄"]}  ', result)

        # 2. 数值类替换
        time_val = values.get('病程时间', '')
        bp_val = values.get('血压', '')
        temp_val = values.get('体温', '')
        hr_val = values.get('心率', '')
        wbc_val = values.get('白细胞', '')
        side_val = values.get('左右侧', '')
        part_val = values.get('部位', '')

        # 提取时间数字和单位
        time_match = re.match(r'(\d+)(\S*)', time_val) if time_val else None
        time_num = time_match.group(1) if time_match else ''
        time_unit = time_match.group(2) if time_match else ''

        # 替换血压 XXX/XXX
        if bp_val:
            parts = bp_val.split('/')
            if len(parts) == 2:
                result = re.sub(r'XXX/XXX\s*mmHg', f'{parts[0]}/{parts[1]}mmHg', result)
                result = re.sub(r'XXX/XXX', f'{parts[0]}/{parts[1]}', result)

        # 替换体温 XX.X
        if temp_val:
            result = re.sub(r'XX\.X℃', f'{temp_val}℃', result)
            result = re.sub(r'XX-XX℃', f'{temp_val}℃', result)

        # 替换心率
        if hr_val:
            result = re.sub(r'(?<!\d)XX(?!\.)(?=次/分)', hr_val, result)

        # 替换白细胞
        if wbc_val:
            result = re.sub(r'(?<!\d)XX(?!\.)(?=×)', wbc_val, result)

        # 替换左右侧
        if side_val:
            result = result.replace('左/右', side_val)
            result = re.sub(r'左/右侧', f'{side_val}侧', result)

        # 替换部位
        if part_val:
            result = result.replace('侧X', f'侧{part_val}')

        # 替换通用 X 占位符
        if time_num:
            # X年/月/天/小时 → 数字+单位
            result = re.sub(r'(?<=\D)X年(?!\()', f'{time_num}年', result)
            result = re.sub(r'(?<=\D)X月(?!\()', f'{time_num}月', result)
            result = re.sub(r'(?<=\D)X天', f'{time_num}天', result)
            result = re.sub(r'(?<=\D)X小时', f'{time_num}小时', result)
            result = re.sub(r'(?<=\D)X周', f'{time_num}周', result)

        # 替换剩余的单个 X（数字占位）
        # 保留未填写的 X 占位符

        self.text_edit.setPlainText(result)
        filled_count = sum(1 for v in values.values() if v)
        remaining = len(re.findall(r'X+', result))
        self.status_bar.showMessage(
            f"已替换 {filled_count} 项信息，剩余 {remaining} 处占位符待手动填写"
        )

    def _replace_placeholders_with_voice(self, voice_text, template_text):
        """将语音输入的核心信息自动替换模板中的X占位符"""
        import re
        result = template_text

        # 0. 中文数字转阿拉伯数字
        cn_map = {'零': '0', '一': '1', '二': '2', '三': '3', '四': '4',
                  '五': '5', '六': '6', '七': '7', '八': '8', '九': '9', '两': '2'}
        def cn_to_arabic(text):
            """简单中文数字转阿拉伯数字（支持个位、十位、百位）"""
            # X百X十X 模式（如三百六十五 = 365）
            m = re.match(r'^([一二三四五六七八九])百([一二三四五六七八九])十([一二三四五六七八九])$', text)
            if m:
                return str(int(cn_map[m.group(1)]) * 100 + int(cn_map[m.group(2)]) * 10 + int(cn_map[m.group(3)]))
            # X百X十 模式（如一百二十 = 120）
            m = re.match(r'^([一二三四五六七八九])百([一二三四五六七八九])十$', text)
            if m:
                return str(int(cn_map[m.group(1)]) * 100 + int(cn_map[m.group(2)]) * 10)
            # X百 模式（如一百 = 100）
            m = re.match(r'^([一二三四五六七八九])百$', text)
            if m:
                return str(int(cn_map[m.group(1)]) * 100)
            # X十X 模式（如六十五 = 65）
            m = re.match(r'^([一二三四五六七八九])十([一二三四五六七八九])$', text)
            if m:
                return str(int(cn_map[m.group(1)]) * 10 + int(cn_map[m.group(2)]))
            # 十X 模式（如十二 = 12）
            m = re.match(r'^十([一二三四五六七八九])$', text)
            if m:
                return str(10 + int(cn_map[m.group(1)]))
            # X十 模式（如三十 = 30）
            m = re.match(r'^([一二三四五六七八九])十$', text)
            if m:
                return str(int(cn_map[m.group(1)]) * 10)
            # 单独的十 = 10
            if text == '十':
                return '10'
            # 纯个位
            r = ''
            for ch in text:
                if ch in cn_map:
                    r += cn_map[ch]
                else:
                    return text
            return r or text

        # 1. 用 SectionParser 解析语音文本中的字段
        sections = self.parser.parse(voice_text)

        # 2. 用 extract_basic_fields 提取基本信息
        inferred = self.classifier.extract_basic_fields(voice_text)

        # 3. 替换基本信息占位符
        name = inferred.get('姓名', sections.get('姓名', ''))
        if name:
            result = re.sub(r'姓名：\s*', f'姓名：{name}  ', result)

        gender = inferred.get('性别', sections.get('性别', ''))
        if gender:
            result = re.sub(r'性别：\s*', f'性别：{gender}  ', result)

        age = inferred.get('年龄', sections.get('年龄', ''))
        if age:
            result = re.sub(r'年龄：\s*', f'年龄：{age}  ', result)

        # 4. 提取时间信息（支持阿拉伯数字和中文数字，替换所有单位）
        for unit in ['年', '月', '天', '小时', '周']:
            # 阿拉伯数字
            m = re.search(r'(\d+)\s*' + re.escape(unit), voice_text)
            if m:
                time_val = m.group(1)
                result = re.sub(r'(?<=\D)X+' + re.escape(unit), f'{time_val}{unit}', result)
                continue
            # 中文数字
            m = re.search(r'([零一二两三四五六七八九十百]+)\s*' + re.escape(unit), voice_text)
            if m:
                time_val = cn_to_arabic(m.group(1))
                result = re.sub(r'(?<=\D)X+' + re.escape(unit), f'{time_val}{unit}', result)

        # 5. 提取血压（XXX/XXX）
        bp_match = re.search(r'(\d{2,3})\s*[/\u6bd4]\s*(\d{2,3})', voice_text)
        if bp_match:
            systolic, diastolic = bp_match.group(1), bp_match.group(2)
            result = re.sub(r'XXX\s*/\s*XXX\s*mmHg', f'{systolic}/{diastolic}mmHg', result)
            result = re.sub(r'XXX\s*/\s*XXX', f'{systolic}/{diastolic}', result)

        # 6. 提取体温
        temp_match = re.search(r'(\d{2}\.\d)\s*[度℃]', voice_text)
        if not temp_match:
            temp_match = re.search(r'体温\s*(\d{2}\.?\d*)', voice_text)
        if temp_match:
            temp = temp_match.group(1)
            result = re.sub(r'XX\.X℃', f'{temp}℃', result)
            result = re.sub(r'XX-XX℃', f'{temp}℃', result)

        # 7. 提取心率
        hr_match = re.search(r'(?:心率|脉搏)\s*(\d{2,3})', voice_text)
        if hr_match:
            hr = hr_match.group(1)
            result = re.sub(r'(?<!\d)XX(?!\.)(?=次/分)', hr, result)

        # 8. 提取白细胞
        wbc_match = re.search(r'(?:白细胞|WBC)\s*(\d+\.?\d*)', voice_text)
        if wbc_match:
            wbc = wbc_match.group(1)
            result = re.sub(r'(?<!\d)XX(?!\.)(?=×)', wbc, result)

        # 9. 提取左右侧
        if re.search(r'左侧|左边', voice_text):
            result = result.replace('左/右', '左')
        elif re.search(r'右侧|右边', voice_text):
            result = result.replace('左/右', '右')

        # 10. 提取部位
        part_match = re.search(r'[左右]侧(\S{1,4}?)(?:肿痛|疼痛|骨折|肿胀)', voice_text)
        if part_match:
            part = part_match.group(1)
            result = result.replace('侧X', f'侧{part}')

        return result

    def _toggle_recording(self):
        """开始/停止录音（委托给 RecordingHandler）"""
        self.recorder.toggle_recording()

    def moveEvent(self, event):
        """主窗口移动时，悬浮预览面板跟随定位"""
        super().moveEvent(event)
        panel = getattr(self, 'asr_preview', None)
        if panel is not None and panel.isVisible():
            panel.reposition()

    def resizeEvent(self, event):
        """主窗口缩放时，悬浮预览面板跟随定位"""
        super().resizeEvent(event)
        panel = getattr(self, 'asr_preview', None)
        if panel is not None and panel.isVisible():
            panel.reposition()

    def _start_recording(self):
        """开始录音（委托给 RecordingHandler）"""
        self.recorder.start_recording()

    def _stop_recording(self):
        """停止录音（委托给 RecordingHandler）"""
        self.recorder.stop_recording()

    def _load_microphones(self):
        """加载麦克风设备列表到下拉框"""
        self.mic_combo.blockSignals(True)
        self.mic_combo.clear()
        self.mic_combo.addItem("系统默认", None)
        try:
            for dev in get_microphone_list():
                self.mic_combo.addItem(dev["name"], dev["index"])
        except Exception as e:
            print(f"[UI] 加载麦克风列表失败: {e}")
        self.mic_combo.blockSignals(False)

    def _on_mic_changed(self, _index):
        """切换录音设备"""
        device_index = self.mic_combo.currentData()
        self.asr.set_input_device(device_index)
        self.status_bar.showMessage(f"🎤 录音设备：{self.mic_combo.currentText()}")

    def _copy_all_text(self):
        """复制病历全文到剪贴板"""
        text = self.text_edit.toPlainText()
        if not text.strip():
            self.status_bar.showMessage("没有内容可复制")
            return
        QApplication.clipboard().setText(text)
        self.status_bar.showMessage("📋 已复制全文到剪贴板")
        Toast.show_toast(self, "已复制到剪贴板", "info")

    def _update_record_duration(self):
        """状态栏实时显示录音时长"""
        if self._record_start_ts is None:
            return
        elapsed = int(time.time() - self._record_start_ts)
        mm, ss = divmod(elapsed, 60)
        self.status_bar.showMessage(f"🔴 录音中... {mm:02d}:{ss:02d}")

    def _poll_audio_level(self):
        """从 ASR 引擎轮询实时音量电平，喜入波形图"""
        level = getattr(self.asr, '_current_level', 0.0)
        self.waveform.add_level(level)

    def _update_field_context(self):
        """检测编辑器末尾的字段名，设置 ASR 字段上下文（用于 LM 偏置）"""
        text = self.text_edit.toPlainText()
        # 从末尾向前找最近的字段名
        field_keywords = ['主诉', '现病史', '既往史', '个人史', '家族史',
                          '体格检查', '辅助检查', '初步诊断', '诊疗经过',
                          '出院情况', '出院医嘱', '影像表现', '诊断意见']
        last_field = ""
        last_pos = -1
        for kw in field_keywords:
            pos = text.rfind(kw)
            if pos > last_pos:
                last_pos = pos
                last_field = kw
        self.asr.set_field_context(last_field)

    # ==================== UI 增强：面板折叠 / 专注模式 / 拖拽 ====================

    def _toggle_left_panel(self):
        """折叠/展开左侧纠错日志面板（F9）"""
        self.left_panel.setVisible(not self.left_panel.isVisible())
        state = "已展开" if self.left_panel.isVisible() else "已折叠"
        self.status_bar.showMessage(f"纠错面板{state}")

    def _toggle_ai_panel(self):
        """折叠/展开 AI 诊断结果区域"""
        visible = self.ai_result.isVisible()
        self.ai_result.setVisible(not visible)
        self.ai_collapse_btn.setText("▼ 展开" if visible else "▲ 收起")
        # 收起时缩小面板高度
        if visible:
            self.ai_group.setMaximumHeight(52)
        else:
            self.ai_group.setMaximumHeight(16777215)  # 解除限制

    def _toggle_focus_mode(self):
        """专注录音模式（F11）：隐藏所有干扰面板，只留编辑器"""
        self._focus_mode = not getattr(self, '_focus_mode', False)
        if self._focus_mode:
            # 记住原始状态
            self._pre_focus = {
                'left_visible': self.left_panel.isVisible(),
                'field_visible': self.field_panel.isVisible(),
                'ai_visible': self.ai_group.isVisible(),
                'toolbar_visible': self.findChild(QToolBar).isVisible(),
            }
            self.left_panel.hide()
            self.field_panel.hide()
            self.ai_group.hide()
            self.audio_player_widget.hide()
            self.findChild(QToolBar).hide()
            self.status_bar.showMessage("🎧 专注模式：按 F11 退出")
            Toast.show_toast(self, "专注模式 · 按 F11 退出", "info")
        else:
            pf = getattr(self, '_pre_focus', {})
            self.left_panel.setVisible(pf.get('left_visible', True))
            self.field_panel.setVisible(pf.get('field_visible', True))
            self.ai_group.setVisible(pf.get('ai_visible', True))
            self.audio_player_widget.show()
            self.findChild(QToolBar).setVisible(pf.get('toolbar_visible', True))
            self.status_bar.showMessage("已退出专注模式")

    def dragEnterEvent(self, event):
        """拖拽进入：接受音频文件"""
        if event.mimeData().hasUrls():
            for url in event.mimeData().urls():
                path = url.toLocalFile()
                if path.lower().endswith(('.wav', '.mp3', '.m4a', '.flac', '.ogg')):
                    event.acceptProposedAction()
                    return
        event.ignore()

    def dropEvent(self, event):
        """拖拽释放：转写音频文件"""
        if not event.mimeData().hasUrls():
            return
        for url in event.mimeData().urls():
            path = url.toLocalFile()
            if path.lower().endswith(('.wav', '.mp3', '.m4a', '.flac', '.ogg')):
                self._transcribe_audio_file(path)
                event.acceptProposedAction()
                return
        event.ignore()

    def _transcribe_audio_file(self, path):
        """后台转写音频文件并填入编辑器"""
        if not self.asr.is_ready():
            Toast.show_toast(self, "语音引擎未就绪，无法转写", "warning")
            return
        fname = os.path.basename(path)
        self.status_bar.showMessage(f"🎧 正在转写：{fname} ...")
        Toast.show_toast(self, f"开始转写 {fname}", "info", duration=1500)

        def _run():
            try:
                text = self.asr.transcribe_file(path)
                return text
            except Exception as e:
                print(f"[Main] 转写失败: {e}")
                return ""

        import threading

        def _worker():
            text = _run()
            # 通过信号回到主线程更新 UI
            self.file_transcribed.emit(text, fname)

        threading.Thread(target=_worker, daemon=True).start()

    def _on_file_transcribed(self, text, fname):
        """音频转写完成回调（主线程）"""
        if text:
            cursor = self.text_edit.textCursor()
            cursor.movePosition(QTextCursor.End)
            if self.text_edit.toPlainText().strip():
                cursor.insertText("\n")
            cursor.insertText(text)
            self.text_edit.setTextCursor(cursor)
            self.status_bar.showMessage(f"✅ 转写完成：{fname}（{len(text)} 字）")
            Toast.show_toast(self, f"转写完成（{len(text)} 字）", "success")
        else:
            self.status_bar.showMessage(f"⚠️ 转写无结果：{fname}")
            Toast.show_toast(self, "转写无结果，请检查音频", "warning")

    def _handle_voice_command(self, text):
        """拦截语音命令。命中则执行对应动作并返回 True（不再作为病历文本填充）"""
        command, arg = self.voice_command.parse(text)
        if command is None:
            return False
        if command == "clear":
            self._clear_text()
            self.status_bar.showMessage("🗣 语音命令：已清除内容")
        elif command == "export":
            self._save_text()
        elif command == "correct":
            self._run_correction()
            self.status_bar.showMessage("🗣 语音命令：开始纠错")
        elif command == "save":
            self._save_record()
        elif command == "copy":
            self._copy_all_text()
        elif command == "open_library":
            self._open_record_manager()
        elif command == "stop_record":
            if self.is_listening:
                self._stop_recording()
        elif command == "start_record":
            if not self.is_listening:
                self._start_recording()
        elif command == "switch_template":
            self._voice_switch_template(arg)
        elif command == "switch_dept":
            self._voice_switch_dept(arg)
        else:
            return False
        return True

    def _voice_switch_template(self, name):
        """按语音命令切换模板（模糊匹配下拉框选项）"""
        for i in range(self.template_combo.count()):
            item = self.template_combo.itemText(i)
            if name in item or item in name:
                self.template_combo.setCurrentIndex(i)
                self.status_bar.showMessage(f"🗣 语音命令：已切换模板→{item}")
                return
        self.status_bar.showMessage(f"🗣 未找到模板：{name}")

    def _voice_switch_dept(self, name):
        """按语音命令切换科室"""
        for i in range(self.dept_combo.count()):
            item = self.dept_combo.itemText(i)
            if name in item or item in name:
                self.dept_combo.setCurrentIndex(i)
                self.status_bar.showMessage(f"🗣 语音命令：已切换科室→{item}")
                return
        self.status_bar.showMessage(f"🗣 未找到科室：{name}")

    def _on_recognized(self, text):
        """识别到文本：显示悬浮预览，等待用户确认（委托给 RecordingHandler）"""
        try:
            print(f"[UI] 收到识别结果，长度: {len(text) if text else 0}")
            self.crash_logger.log_event("ASR识别完成", {"text_length": len(text) if text else 0})
            self.recorder.on_recognized(text)
        except Exception as e:
            print(f"[UI] 处理识别结果时出错: {e}")
            import traceback
            traceback.print_exc()
            self.crash_logger.log_exception(
                type(e), e, e.__traceback__,
                context="处理识别结果"
            )
            self.partial_label.setText("⚠️ 处理识别结果时出错")
            self.status_bar.showMessage(f"错误: {e}")

        # 录音结束后自动加载音频到播放器（供回放校对）
        self._load_last_audio()

    def _apply_asr_result(self, text):
        """把已确认的识别文本按病历格式结构化填充（模板增量填充 / X占位符 / 追加模式）"""
        # 如果选了模板，用增量填充（只填空字段，不覆盖已有内容）
        template_name = self.template_combo.currentText()
        if template_name and self.current_dept:
            template_content = self.template_engine.get_template(
                self.current_dept, template_name
            )
            if template_content:
                # 使用录音前的干净内容作为 base（避免流式识别追加的垃圾文本干扰）
                clean_base = getattr(self, '_stream_base_text', '').strip()
                base = clean_base if clean_base else template_content

                # 检查是否是常见病模板（含X占位符）
                import re as _re
                if _re.search(r'X+', base):
                    # 常见病模板模式：用语音输入替换X占位符
                    filled = self._replace_placeholders_with_voice(text, base)
                    self.text_edit.setPlainText(filled)
                    self._last_asr_snapshot = filled  # 记录ASR填充后快照
                    remaining = len(_re.findall(r'X+', filled))
                    self.partial_label.setText(f"✓ 已替换占位符，剩余 {remaining} 处")
                    self.status_bar.showMessage(f"套用完成，剩余 {remaining} 处占位符待手动填写")
                    print(f"[UI] 常见病模板占位符替换完成")
                    return

                filled = self.classifier.incremental_fill(text, base)
                self.text_edit.setPlainText(filled)
                self._last_asr_snapshot = filled  # 记录ASR填充后快照
                self.partial_label.setText("✓ 识别完成")
                self.status_bar.showMessage(f"识别完成，共 {len(filled)} 字")
                print(f"[UI] 增量填充完成")
                return

        # 将识别文本插入到编辑器（追加模式，不覆盖已有内容）
        current = self.text_edit.toPlainText()
        if current.strip():
            # 有内容时，在末尾追加识别结果
            text = current.rstrip() + "\n\n" + text
        self.text_edit.setPlainText(text)
        self._last_asr_snapshot = text  # 记录ASR填充后快照
        self.partial_label.setText("✓ 识别完成")
        self.status_bar.showMessage(f"识别完成，共 {len(text)} 字")
        print(f"[UI] 文本已插入编辑器")
        # 智能推荐模板（非侵入，仅状态栏提示）
        self._recommend_template(self.text_edit.toPlainText())

    def _on_preview_accept(self):
        """接受预览：把识别结果按病历格式填充到编辑器（委托给 RecordingHandler）"""
        self.recorder.on_preview_accept()

    def _on_preview_reject(self):
        """拒绝预览：丢弃本次识别结果（委托给 RecordingHandler）"""
        self.recorder.on_preview_reject()

    def _on_preview_retry(self):
        """重听：重新开始录音（委托给 RecordingHandler）"""
        self.recorder.on_preview_retry()

    def _load_last_audio(self):
        """把最近一次录音加载到音文对照播放器"""
        audio_path = getattr(self.asr, 'last_audio_path', None)
        if audio_path and os.path.exists(audio_path):
            self.audio_player_widget.load(audio_path)

    def _on_partial(self, text):
        """流式识别中间结果（委托给 RecordingHandler）"""
        self.recorder.on_partial(text)

    def _run_correction(self):
        """执行纠错"""
        text = self.text_edit.toPlainText().strip()
        if not text:
            QMessageBox.information(self, "提示", "请先输入或录制文本")
            return

        self.status_bar.showMessage("正在纠错...")
        QApplication.processEvents()

        self._last_correction_original = text
        self.correct_thread = CorrectThread(self.corrector, text)
        self.correct_thread.correction_done.connect(self._on_correction_done)
        self.correct_thread.start()

    def _on_correction_done(self, corrected, log):
        """纠错完成"""
        original_text = getattr(self, '_last_correction_original', '')
        review_accepted = False
        if original_text:
            try:
                dialog = DiffReviewDialog(original_text, corrected, log, self)
                if dialog.exec_() == QDialog.Accepted:
                    corrected = dialog.result_text
                    review_accepted = True
                else:
                    corrected = original_text
            except Exception as e:
                print(f"[DiffReview] error: {e}")
                corrected = original_text if original_text else corrected

        self.text_edit.setPlainText(corrected)

        # 保存完整日志
        self._all_logs = log

        # 记录纠错反馈（用于 LM 迭代训练）
        try:
            self.feedback.log_corrections(log, source="corrector")
            if review_accepted:
                self.feedback.log_accept_all()
        except Exception:
            pass

        # 应用筛选
        self._apply_filter()

        # 统计
        counts = {"错别字": 0, "逻辑错误": 0, "缺项提醒": 0}
        for item in log:
            cat = item.get("分类", "")
            if cat in counts:
                counts[cat] += 1
        total = sum(counts.values())
        self.stats_label.setText(f"共{total}条 | 🔤{counts['错别字']} 🧠{counts['逻辑错误']} ⚠️{counts['缺项提醒']}")

        self.status_bar.showMessage(f"纠错完成，共 {total} 条建议")
        Toast.show_toast(self, f"纠错完成，共 {total} 条建议", "success")

    def _apply_filter(self):
        """根据筛选按钮过滤日志"""
        self.log_list.clear()

        show_typo = self.filter_typo.isChecked()
        show_logic = self.filter_logic.isChecked()
        show_missing = self.filter_missing.isChecked()

        filter_map = {
            "错别字": show_typo,
            "逻辑错误": show_logic,
            "缺项提醒": show_missing,
        }

        self._log_item_map = {}  # list item → log data index

        for idx, item in enumerate(self._all_logs):
            cat = item.get("分类", "")
            if cat not in filter_map or not filter_map[cat]:
                continue

            # 分类图标和颜色
            if cat == "错别字":
                icon = "🔤"
                color = "#00d4ff"
            elif cat == "逻辑错误":
                icon = "🧠"
                color = "#ff9944"
            elif cat == "缺项提醒":
                icon = "⚠️"
                color = "#ffdd44"
            else:
                icon = "📝"
                color = "#b8c5d6"

            level = item.get("级别", "")
            type_name = item.get("type", "")
            orig = item.get("原文", "")
            corr = item.get("修正", "")

            # 类型名称
            line = f"{icon} <span style='color:{color};'>{type_name}</span>"
            if level:
                line += f" <span style='color:#6b8a9a;font-size:10px;'>[{level}]</span>"

            # 原文 → 修正
            detail = ""
            if orig and orig != corr:
                detail = f"<span style='color:#ff6b6b;'>{orig}</span> → <span style='color:#51cf66;'>{corr}</span>"
            elif orig:
                detail = f"<span style='color:#e0e0e0;'>{orig}</span>"

            # 相似度
            if item.get("相似度"):
                detail += f" <span style='color:#6b8a9a;font-size:10px;'>({item['相似度']})</span>"

            # 合并为一条 item
            full_text = line
            if detail:
                full_text += f"<br>{detail}"

            list_item = QListWidgetItem(full_text)
            list_item.setData(Qt.UserRole, idx)  # 映射到 _all_logs 索引
            self.log_list.addItem(list_item)

    def _on_log_item_clicked(self, item):
        """点击日志条目 → 高亮编辑器对应文本"""
        idx = item.data(Qt.UserRole)
        if idx is None or idx >= len(self._all_logs):
            return
        log_data = self._all_logs[idx]
        orig = log_data.get("原文", "")
        corr = log_data.get("修正", "")

        # 优先高亮修正后的文本，找不到再高亮原文
        search_text = corr if corr and corr != orig else orig
        if not search_text:
            return

        self._highlight_text_in_editor(search_text)
        self.log_hint.setText(f"已定位：{orig} → {corr}" if orig != corr else f"已定位：{orig}")
        self._accept_btn.setEnabled(True)
        self._reject_btn.setEnabled(True)

    def _highlight_text_in_editor(self, text):
        """在编辑器中高亮所有匹配文本"""
        doc = self.text_edit.document()
        cursor = self.text_edit.textCursor()
        cursor.select(QTextCursor.Document)
        cursor.clearSelection()
        self.text_edit.setTextCursor(cursor)

        extra = []
        found_any = False
        cursor = doc.find(text, 0)
        while not cursor.isNull():
            found_any = True
            sel = QTextEdit.ExtraSelection()
            sel.cursor = cursor
            sel.format = QTextCharFormat()
            sel.format.setBackground(QColor(255, 200, 50, 100))
            extra.append(sel)
            cursor = doc.find(text, cursor)
        self.text_edit.setExtraSelections(extra)

        if found_any:
            # 跳到第一个匹配位置
            first = doc.find(text, 0)
            if not first.isNull():
                self.text_edit.setTextCursor(first)

    def _accept_correction(self):
        """接受当前纠错（默认已应用，此操作为确认，从拒绝列表中移除）"""
        item = self.log_list.currentItem()
        if not item:
            return
        idx = item.data(Qt.UserRole)
        if idx is None or idx >= len(self._all_logs):
            return
        log_data = self._all_logs[idx]
        orig = log_data.get("原文", "")
        corr = log_data.get("修正", "")
        if orig and corr:
            self.corrector.rejections.pop((orig, corr), None)
            self.corrector.save_rejection("", "")  # 刷新文件
            # 重建 rejection 文件（去掉空键）
            try:
                serializable = [
                    "\x00".join(k) for k in self.corrector.rejections.keys()
                    if k[0] and k[1]
                ]
                with open(self.corrector.rejection_path, 'w', encoding='utf-8') as f:
                    json.dump(serializable, f, ensure_ascii=False, indent=2)
            except Exception:
                pass
        self.status_bar.showMessage(f"✓ 已接受纠错：{orig} → {corr}")
        # 记忆库：接受当前纠错
        try:
            memory = self._get_memory()
            if memory and orig and corr and orig != corr:
                memory.accept_memory_by_values(orig, corr, doctor_id=self.current_user.get("id") if isinstance(self.current_user, dict) else None, dept=getattr(self, "current_dept", "") or "")
        except Exception as e:
            print(f"[Memory] 接受纠错失败: {e}")


    def _accept_all_corrections(self):
        """一键接受所有纠错建议"""
        if not self._all_logs:
            Toast.show_toast(self, "没有纠错建议", "info")
            return
        count = sum(1 for item in self._all_logs
                    if not item.get("_rejected") and item.get("原文") and item.get("修正"))
        # 标记所有为已接受（清除拒绝列表中的对应项）
        for item in self._all_logs:
            orig = item.get("原文", "")
            corr = item.get("修正", "")
            if orig and corr:
                self.corrector.rejections.pop((orig, corr), None)
        # 刷新 rejection 文件
        try:
            serializable = [
                "\x00".join(k) for k in self.corrector.rejections.keys()
                if k[0] and k[1]
            ]
            with open(self.corrector.rejection_path, 'w', encoding='utf-8') as f:
                json.dump(serializable, f, ensure_ascii=False, indent=2)
        except Exception:
            pass
        # 反馈收集：标记所有 pending 为 accepted
        try:
            self.feedback.log_accept_all()
        except Exception:
            pass
        self.status_bar.showMessage(f"✓ 已接受全部 {count} 条纠错")
        Toast.show_toast(self, f"已接受全部 {count} 条纠错", "success")

    def _reject_correction(self):
        """拒绝当前纠错，写入个人偏好"""
        item = self.log_list.currentItem()
        if not item:
            return
        idx = item.data(Qt.UserRole)
        if idx is None or idx >= len(self._all_logs):
            return
        log_data = self._all_logs[idx]
        orig = log_data.get("原文", "")
        corr = log_data.get("修正", "")
        if not orig or not corr:
            return

        # 记录拒绝规则
        self.corrector.save_rejection(orig, corr)

        # 记录拒绝反馈（用于 LM 迭代训练）
        try:
            self.feedback.log_rejection(orig, corr)
        except Exception:
            pass

        # 记忆库：拒绝当前纠错
        try:
            memory = self._get_memory()
            if memory and orig and corr and orig != corr:
                memory.reject_memory_by_values(orig, corr, doctor_id=self.current_user.get("id") if isinstance(self.current_user, dict) else None, dept=getattr(self, "current_dept", "") or "")
        except Exception as e:
            print(f"[Memory] 拒绝纠错失败: {e}")

        # 恢复原文（从编辑器中撤销这条纠错）
        # 修复：只替换第一个匹配，避免 replace 替换所有相同文本导致误伤
        text = self.text_edit.toPlainText()
        idx_in_text = text.find(corr)
        if idx_in_text >= 0:
            text = text[:idx_in_text] + orig + text[idx_in_text + len(corr):]
        self.text_edit.setPlainText(text)

        # 从日志中移除
        self._all_logs[idx]["_rejected"] = True
        self._apply_filter()

        # 更新统计
        counts = {"错别字": 0, "逻辑错误": 0, "缺项提醒": 0}
        for item in self._all_logs:
            if item.get("_rejected"):
                continue
            cat = item.get("分类", "")
            if cat in counts:
                counts[cat] += 1
        total = sum(counts.values())
        self.stats_label.setText(f"共{total}条 | 🔤{counts['错别字']} 🧠{counts['逻辑错误']} ⚠️{counts['缺项提醒']}")

        self._accept_btn.setEnabled(False)
        self._reject_btn.setEnabled(False)
        self.log_hint.setText(f"已拒绝并记住：{orig}")
        self.status_bar.showMessage(f"✗ 已拒绝纠错并记录偏好：{orig} → {corr}")

    def _get_memory(self):
        if self.memory is None:
            try:
                self.memory = get_memory()
            except Exception as e:
                print(f"[Memory] 初始化失败: {e}")
        return self.memory

    def _get_topk_engine(self):
        if self.topk_engine is None:
            try:
                self.topk_engine = get_topk_engine(memory=self._get_memory())
            except Exception as e:
                print(f"[TopK] 初始化失败: {e}")
        return self.topk_engine


    def _clear_text(self):
        """清除文本"""
        self.text_edit.clear()
        self.log_list.clear()
        self._all_logs = []
        self.stats_label.setText("")
        self.partial_label.setText("等待输入...")

    def _save_text(self):
        """导出文本（.txt / .md / 打印预览）"""
        text = self.text_edit.toPlainText()
        if not text:
            QMessageBox.information(self, "提示", "没有内容可导出")
            return

        # 弹出格式选择菜单
        menu = QMenu(self)
        txt_action = QAction("📄 导出 .txt", self)
        txt_action.triggered.connect(lambda: self._export_as(text, "txt"))
        menu.addAction(txt_action)

        md_action = QAction("📝 导出 .md", self)
        md_action.triggered.connect(lambda: self._export_as(text, "md"))
        menu.addAction(md_action)

        docx_action = QAction("📘 导出 Word (.docx)", self)
        docx_action.triggered.connect(lambda: self._export_docx(text))
        menu.addAction(docx_action)

        preview_action = QAction("🖨 打印预览", self)
        preview_action.triggered.connect(lambda: self._print_preview(text))
        menu.addAction(preview_action)

        # 在导出按钮位置显示菜单
        sender_btn = self.sender()
        if sender_btn:
            btn_rect = sender_btn.geometry()
            menu.exec_(self.mapToGlobal(btn_rect.bottomLeft()))
        else:
            menu.exec_(self.mapToGlobal(self.record_btn.geometry().bottomLeft()))

    def _export_as(self, text, fmt):
        """按指定格式导出"""
        if fmt == "md":
            # 将病历文本转为 Markdown 格式
            md_text = self._convert_to_markdown(text)
            default_name = "病历记录.md"
            filter_str = "Markdown 文件 (*.md);;所有文件 (*)"
        else:
            md_text = text
            default_name = "病历记录.txt"
            filter_str = "文本文件 (*.txt);;所有文件 (*)"

        path, _ = QFileDialog.getSaveFileName(
            self, f"导出 {fmt.upper()}",
            os.path.join(os.path.expanduser("~"), "Desktop", default_name),
            filter_str
        )
        if path:
            with open(path, 'w', encoding='utf-8') as f:
                f.write(md_text)
            self.status_bar.showMessage(f"已导出：{path}")
            # 收集导出文本作为高质量语料（用于 LM 迭代训练）
            try:
                self.feedback.collect_corpus(text)
            except Exception:
                pass

    def _convert_to_markdown(self, text):
        """将纯文本病历转为 Markdown 格式"""
        lines = text.split('\n')
        result = ["# 病历记录\n"]
        for line in lines:
            line = line.strip()
            if not line:
                continue
            # 字段名加粗
            if '：' in line or ':' in line:
                parts = re.split(r'[：:]', line, 1)
                if len(parts) == 2:
                    field, content = parts[0].strip(), parts[1].strip()
                    result.append(f"**{field}**：{content}\n")
                else:
                    result.append(f"{line}\n")
            else:
                result.append(f"{line}\n")
        return '\n'.join(result)

    def _export_docx(self, text):
        """导出为 Word (.docx)。需 python-docx，未安装时提示并降级。"""
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.oxml.ns import qn
        except ImportError:
            ret = QMessageBox.question(
                self, "缺少依赖",
                "导出 Word 需要 python-docx 库，当前未安装。\n"
                "可在命令行执行：pip install python-docx\n\n"
                "是否改为导出 .txt 文本？",
                QMessageBox.Yes | QMessageBox.No, QMessageBox.Yes
            )
            if ret == QMessageBox.Yes:
                self._export_as(text, "txt")
            return

        path, _ = QFileDialog.getSaveFileName(
            self, "导出 Word",
            os.path.join(os.path.expanduser("~"), "Desktop", "病历记录.docx"),
            "Word 文档 (*.docx)"
        )
        if not path:
            return
        try:
            doc = Document()
            # 设置中文字体
            style = doc.styles["Normal"]
            style.font.name = "宋体"
            style.font.size = Pt(11)
            style.element.rcPr.rFonts.set(qn("w:eastAsia"), "宋体")

            title = doc.add_heading("病历记录", level=1)
            for line in text.split("\n"):
                line = line.rstrip()
                if not line:
                    doc.add_paragraph("")
                    continue
                # 字段名加粗
                if "：" in line or ":" in line:
                    parts = re.split(r"[：:]", line, 1)
                    if len(parts) == 2:
                        p = doc.add_paragraph()
                        run = p.add_run(parts[0].strip() + "：")
                        run.bold = True
                        p.add_run(parts[1].strip())
                        continue
                doc.add_paragraph(line)
            doc.save(path)
            self.status_bar.showMessage(f"已导出 Word：{path}")
        except Exception as e:
            QMessageBox.warning(self, "导出失败", str(e))

    def _print_preview(self, text):
        """打印预览"""
        from PyQt5.QtPrintSupport import QPrintPreviewDialog, QPrinter
        printer = QPrinter(QPrinter.HighResolution)
        dialog = QPrintPreviewDialog(printer, self)
        dialog.paintRequested.connect(lambda p: self._print_text(p, text))
        dialog.exec_()

    def _print_text(self, printer, text):
        """实际打印/预览渲染"""
        from PyQt5.QtPrintSupport import QPrinter
        doc = QTextDocument()
        # 简单 HTML 格式
        html = "<h2>病历记录</h2><hr/>"
        for line in text.split('\n'):
            escaped = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
            if '：' in line or ':' in line:
                parts = re.split(r'[：:]', line, 1)
                if len(parts) == 2:
                    html += f"<p><b>{parts[0].strip()}</b>：{parts[1].strip()}</p>"
                else:
                    html += f"<p>{escaped}</p>"
            else:
                html += f"<p>{escaped}</p>"
        doc.setHtml(html)
        doc.print_(printer)

    def _open_phrase_library(self):
        """打开常用语句库（非模态，可连续插入多条）"""
        if self._phrase_dialog is None:
            self._phrase_dialog = PhraseDialog(self.phrase_lib, self)
            self._phrase_dialog.phrase_selected.connect(self._insert_term_at_cursor)
        self._phrase_dialog.show()
        self._phrase_dialog.raise_()
        self._phrase_dialog.activateWindow()

    def _open_backup_menu(self):
        """备份/恢复/刷新热词菜单"""
        menu = QMenu(self)
        backup_action = QAction("💾 立即备份到文件", self)
        backup_action.triggered.connect(self._backup_now)
        menu.addAction(backup_action)

        restore_action = QAction("♻ 从备份文件恢复", self)
        restore_action.triggered.connect(self._restore_backup)
        menu.addAction(restore_action)

        menu.addSeparator()
        hotword_action = QAction("🔥 从历史病历刷新个人热词", self)
        hotword_action.triggered.connect(lambda: self._refresh_user_hotwords(silent=False))
        menu.addAction(hotword_action)

        sender_btn = self.sender()
        if sender_btn and hasattr(sender_btn, "geometry"):
            menu.exec_(self.mapToGlobal(sender_btn.geometry().bottomLeft()))
        else:
            menu.exec_(self.mapToGlobal(self.record_btn.geometry().bottomLeft()))

    def _backup_now(self):
        """手动备份数据库到用户选定位置"""
        if not self.db:
            QMessageBox.warning(self, "提示", "数据库未初始化")
            return
        from datetime import datetime
        default_name = "病历备份_%s.db" % datetime.now().strftime("%Y%m%d_%H%M%S")
        path, _ = QFileDialog.getSaveFileName(
            self, "备份数据库",
            os.path.join(os.path.expanduser("~"), "Desktop", default_name),
            "SQLite 数据库 (*.db)"
        )
        if not path:
            return
        try:
            self.db.backup_to(path)
            QMessageBox.information(self, "备份成功", "已备份到：\n%s" % path)
        except Exception as e:
            QMessageBox.warning(self, "备份失败", str(e))

    def _restore_backup(self):
        """从备份文件恢复数据库"""
        if not self.db:
            QMessageBox.warning(self, "提示", "数据库未初始化")
            return
        ret = QMessageBox.warning(
            self, "确认恢复",
            "恢复将用备份数据覆盖当前所有病历，此操作不可撤销！\n建议先做一次备份。确定继续？",
            QMessageBox.Yes | QMessageBox.No, QMessageBox.No
        )
        if ret != QMessageBox.Yes:
            return
        path, _ = QFileDialog.getOpenFileName(
            self, "选择备份文件",
            os.path.expanduser("~"), "SQLite 数据库 (*.db)"
        )
        if not path:
            return
        try:
            self.db.restore_from(path)
            QMessageBox.information(
                self, "恢复成功",
                "已从备份恢复。请重启软件以确保数据一致。"
            )
        except Exception as e:
            QMessageBox.warning(self, "恢复失败", str(e))

    def _open_rule_manager(self):
        """打开规则管理对话框"""
        dialog = RuleManagerDialog(self.rule_engine, self)
        dialog.exec_()
        # 刷新后重新加载规则
        self.status_bar.showMessage("📏 规则已更新")

    def _save_record(self):
        """保存当前病历到数据库（新建或更新并记录版本）"""
        if not self.db or not self.current_user:
            QMessageBox.warning(self, "提示", "未登录，无法保存病历")
            return
        content = self.text_edit.toPlainText().strip()
        if not content:
            QMessageBox.information(self, "提示", "没有内容可保存")
            return
        # 从编辑器提取患者姓名（尽量）
        patient_name = self._extract_patient_name(content)
        dept = self.current_dept if self.current_dept != "通用" else self.current_user.get("department", "")
        template_name = self.template_combo.currentText() if hasattr(self, "template_combo") else ""
        if self.current_record_id is None:
            self.current_record_id = self.db.create_record(
                self.current_user["id"], patient_name, dept, template_name, content, "草稿"
            )
            self.status_bar.showMessage("💾 病历已保存到病历库（新建）")
            Toast.show_toast(self, "病历已保存", "success")
        else:
            self.db.update_record(
                self.current_record_id, patient_name=patient_name,
                department=dept, template_name=template_name, content=content
            )
            self.status_bar.showMessage("💾 病历已更新（已记录版本）")
            Toast.show_toast(self, "病历已更新", "success")

        # 从本次保存的内容增量学习高频专业词 → 用户自适应热词
        try:
            learned = [w for w in self.corrector.active_words
                       if len(w) >= 2 and content.count(w) >= 1]
            if learned:
                self.asr.update_user_hotwords(learned)
            self._refresh_topk_hotwords(silent=True)
        except Exception as e:
            print(f"[Main] 增量学习热词失败: {e}")

        # 收集高质量语料 + 标记纠错为已接受（用于 LM 迭代训练）
        try:
            self.feedback.collect_corpus(content)
            self.feedback.log_accept_all()
        except Exception:
            pass

        # 记忆库：保存时记录医生确认后的最终稿与手动修正
        try:
            memory = self._get_memory()
            if memory:
                doctor_id = self.current_user.get("id") if isinstance(self.current_user, dict) else None
                dept = getattr(self, "current_dept", "") or ""
                memory.record_final_text(content, doctor_id=doctor_id, dept=dept, record_id=getattr(self, "current_record_id", None), snapshot=getattr(self, "_last_asr_snapshot", ''))
        except Exception as e:
            print(f"[Memory] 记录终稿失败: {e}")

        # 提取用户手动修正（对比 ASR 快照 vs 终稿）→ 写入混淆对候选
        try:
            self._extract_manual_corrections(content)
        except Exception as e:
            print(f"[Main] 提取手动修正失败: {e}")

    def _extract_patient_name(self, content):
        """从病历文本中提取“姓名：XXX”字段，提不到返回空字符串"""
        m = re.search(r'姓名[：:]\s*([^\s　\n]{1,10})', content)
        return m.group(1).strip() if m else ""

    def _extract_manual_corrections(self, final_text):
        """对比 ASR 快照与终稿，提取用户手动修正的词语对，写入混淆对候选"""
        import difflib
        snapshot = getattr(self, '_last_asr_snapshot', '')
        if not snapshot or not final_text:
            return
        # 如果终稿和快照完全相同，没有手动修改
        if snapshot.strip() == final_text.strip():
            return

        # 用 difflib 找出替换块
        sm = difflib.SequenceMatcher(None, snapshot, final_text)
        corrections = []
        for tag, i1, i2, j1, j2 in sm.get_opcodes():
            if tag == 'replace':
                old = snapshot[i1:i2].strip()
                new = final_text[j1:j2].strip()
                # 过滤：太长的不算单词级修正（可能是整段重写）
                # 太短的不算（单字可能是误触）
                if old and new and 2 <= len(old) <= 20 and 2 <= len(new) <= 20:
                    # 过滤纯标点/空白变化
                    if old.replace(' ', '') != new.replace(' ', ''):
                        corrections.append((old, new))

        if not corrections:
            return

        # 写入 correction_feedback.jsonl（标记为 manual_edit 来源）
        import json, datetime
        feedback_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "correction_feedback.jsonl")
        with open(feedback_path, 'a', encoding='utf-8') as f:
            for old, new in corrections:
                record = {
                    "original": old,
                    "corrected": new,
                    "原文": old,
                    "修正": new,
                    "type": "manual_edit",
                    "status": "accepted",
                    "source": "user_manual",
                    "timestamp": datetime.datetime.now().isoformat()
                }
                f.write(json.dumps(record, ensure_ascii=False) + '\n')

        print(f"[Main] 提取手动修正: {len(corrections)} 对 → {corrections[:5]}")
        # 清除快照（避免下次保存重复提取）
        self._last_asr_snapshot = final_text

    def _open_record_manager(self):
        """打开病历库，选中后回填到编辑器"""
        if not self.db or not self.current_user:
            QMessageBox.warning(self, "提示", "未登录，无法打开病历库")
            return
        dialog = RecordManagerDialog(self.db, self.current_user, self)
        if dialog.exec_() == QDialog.Accepted and dialog.selected_record:
            rec = dialog.selected_record
            self.text_edit.setPlainText(rec["content"])
            self.current_record_id = rec["id"]
            self.status_bar.showMessage(
                "📚 已打开病历：%s（%s）" % (rec["patient_name"] or "未命名", rec["updated_at"])
            )

    def _open_user_manager(self):
        """打开用户管理（仅管理员）"""
        if self.current_user.get("role") != "admin":
            QMessageBox.warning(self, "权限不足", "仅管理员可管理用户")
            return
        dialog = UserManagerDialog(self.db, self)
        dialog.exec_()

    # ==================== AI 辅助诊断 ====================
    def _run_diagnosis(self):
        """启动后台线程分析当前病历"""
        text = self.text_edit.toPlainText().strip()
        if not text:
            QMessageBox.information(self, "提示", "请先输入或录制病历内容")
            return
        if self.diagnosis_thread is not None and self.diagnosis_thread.isRunning():
            return
        self.ai_analyze_btn.setEnabled(False)
        self.ai_status_label.setText("分析中...")
        self.diagnosis_thread = DiagnosisThread(self.diagnosis_assistant, text)
        self.diagnosis_thread.analysis_done.connect(self._on_diagnosis_done)
        self.diagnosis_thread.start()

    def _on_diagnosis_done(self, result):
        """接收分析结果并分组展示"""
        self.ai_analyze_btn.setEnabled(True)
        self.ai_status_label.setText("")
        if result.get("error"):
            self.ai_result.setHtml(
                f'<div style="color:#ff6b6b;">分析失败：{result["error"]}</div>'
            )
            return
        self.ai_result.setHtml(self._render_diagnosis_html(result))
        self.status_bar.showMessage("🔬 AI 分析完成")

    def _show_qa_dialog(self):
        """打开知识问答对话框"""
        try:
            from qa_dialog import KnowledgeQADialog
            dialog = KnowledgeQADialog(qa_engine=self.qa_engine, parent=self)
            dialog.exec_()
        except Exception as e:
            import traceback
            self.status_bar.showMessage(f"⚠ 问答启动失败：{e}")
            traceback.print_exc()

    def _render_diagnosis_html(self, result):
        """将分析结果渲染为 HTML"""
        def section(title, color="#00d4ff"):
            return (f'<div style="color:{color};font-weight:bold;'
                    f'margin:10px 0 4px 0;">{title}</div>')

        html = []

        # 风险预警（置顶、红色高亮）
        alerts = result.get("risk_alerts", [])
        if alerts:
            html.append(section("⚠️ 风险预警", "#ff6b6b"))
            html.append('<div style="background:rgba(255,107,107,0.12);'
                        'border:1px solid rgba(255,107,107,0.4);'
                        'border-radius:6px;padding:6px 10px;color:#ff9b9b;">')
            for a in alerts:
                if isinstance(a, (tuple, list)):
                    label, msg = (a[0], a[1]) if len(a) >= 2 else ("", a[0])
                    html.append(f'▸ <b>{label}</b>：{msg}<br>')
                else:
                    html.append(f'▸ {a}<br>')
            html.append('</div>')

        # 可能诊断
        diagnoses = result.get("diagnoses", [])
        if diagnoses:
            html.append(section("🧭 可能诊断"))
            html.append('<div style="color:#e0e0e0;">')
            for d in diagnoses:
                if isinstance(d, dict):
                    name = d.get("disease", d.get("name", ""))
                    score = d.get("score")
                    matched = d.get("matched")
                    line = f'• <b>{name}</b>'
                    if score is not None:
                        line += f' <span style="color:#6b8a9a;">(评分 {score})</span>'
                    if matched:
                        line += f' <span style="color:#6b8a9a;">→ {"、".join(matched)}</span>'
                    html.append(line + '<br>')
                else:
                    html.append(f'• {d}<br>')
            html.append('</div>')

        # 中医辨证
        tcm = result.get("tcm_analysis")
        if tcm:
            html.append(section("🌿 中医辨证", "#4ecdc4"))
            html.append('<div style="color:#e0e0e0;">')
            # 中医诊断
            tcm_dx = tcm.get("tcm_diagnoses", [])
            if tcm_dx:
                html.append(f'• 中医诊断：<b style="color:#4ecdc4;">'
                            f'{"、".join(tcm_dx)}</b><br>')
            # 证型
            syndromes = tcm.get("syndromes", [])
            if syndromes:
                for s in syndromes:
                    if isinstance(s, dict):
                        name = s.get("syndrome", "")
                        score = s.get("score", "")
                        matched = s.get("matched", [])
                        line = f'• 证型：<b style="color:#4ecdc4;">{name}</b>'
                        if score:
                            line += f' <span style="color:#6b8a9a;">(评分 {score})</span>'
                        if matched:
                            line += f' <span style="color:#6b8a9a;">→ {"、".join(matched)}</span>'
                        html.append(line + '<br>')
            # 治法方药
            treatment = tcm.get("treatment")
            if treatment:
                method = treatment.get("治法", "")
                formula = treatment.get("代表方", "")
                herbs = treatment.get("组成", [])
                html.append(f'• 治法：{method}<br>')
                if formula:
                    html.append(f'• 代表方：<b>{formula}</b><br>')
                if herbs:
                    html.append(f'<span style="color:#6b8a9a;"> 组成：{"、".join(herbs)}</span><br>')
            # 类证鉴别
            differential = tcm.get("differential", [])
            if differential:
                html.append('<div style="margin-top:4px;color:#ffa94d;">类证鉴别：</div>')
                for d in differential:
                    if isinstance(d, dict):
                        syn = d.get("syndrome", "")
                        kp = d.get("key_points", "")
                        tf = d.get("治法", "")
                        rx = d.get("代表方", "")
                        html.append(f'  • {syn}：{kp}')
                        if tf:
                            html.append(f' → {tf}')
                        if rx:
                            html.append(f'（{rx}）')
                        html.append('<br>')
            html.append('</div>')

        # 用药审查
        review = result.get("drug_review") or {}
        if review:
            matched = review.get("matched", [])
            mismatched = review.get("mismatched", [])
            recommended = review.get("recommended", [])
            if matched or mismatched or recommended:
                html.append(section("💊 用药审查"))
                html.append('<div style="color:#e0e0e0;">')
                for m in matched:
                    if isinstance(m, dict):
                        drug = m.get("drug", "")
                        forl = m.get("for", [])
                        suffix = f'（适用：{"、".join(forl)}）' if forl else ""
                        html.append(f'✅ {drug}{suffix}<br>')
                    else:
                        html.append(f'✅ {m}<br>')
                for m in mismatched:
                    if isinstance(m, dict):
                        drug = m.get("drug", "")
                        note = m.get("note", "与当前诊断不匹配")
                        html.append(f'<span style="color:#ffa94d;">'
                                    f'⚠️ {drug}：{note}</span><br>')
                    else:
                        html.append(f'<span style="color:#ffa94d;">⚠️ {m}</span><br>')
                if recommended:
                    html.append('<span style="color:#6b8a9a;">建议补充用药：'
                                + "、".join(recommended) + '</span><br>')
                html.append('</div>')

        # 检查建议
        exams = result.get("exam_suggestions", [])
        if exams:
            html.append(section("🔍 检查建议"))
            html.append('<div style="color:#e0e0e0;">')
            for e in exams:
                if isinstance(e, dict):
                    name = e.get("exam", "")
                    recorded = e.get("recorded")
                    tag = '已记录' if recorded else '<span style="color:#ffa94d;">建议补充</span>'
                    html.append(f'• {name}（{tag}）<br>')
                else:
                    html.append(f'• {e}<br>')
            html.append('</div>')

        # 免责声明
        disclaimer = result.get("disclaimer", "")
        if disclaimer:
            html.append(f'<div style="color:#6b8a9a;font-size:11px;'
                        f'margin-top:12px;border-top:1px solid rgba(255,255,255,0.1);'
                        f'padding-top:6px;">{disclaimer}</div>')

        if not html:
            return '<div style="color:#6b8a9a;">未提取到可分析的临床信息。</div>'
        return "".join(html)

    def _open_struct_view(self):
        """打开结构化解析视图"""
        text = self.text_edit.toPlainText().strip()
        if not text:
            QMessageBox.information(self, "提示", "请先输入或录制文本")
            return

        dialog = SectionDialog(self.parser, text, self)
        if dialog.exec_() == QDialog.Accepted:
            # 将结构化结果填充回编辑器
            structured = dialog.get_result()
            self.text_edit.setPlainText(structured)
            self.status_bar.showMessage("📋 结构化填充完成")

    def _retrain_lm(self):
        """一键重训 3-gram 语言模型（从累积的用户语料 + 基础语料）"""
        import subprocess
        corpus_file = os.path.join(os.path.dirname(os.path.abspath(__file__)), "user_corpus.txt")
        corpus_count = 0
        if os.path.exists(corpus_file):
            with open(corpus_file, "r", encoding="utf-8") as f:
                corpus_count = sum(1 for line in f if line.strip())

        reply = QMessageBox.question(
            self, "重训语言模型",
            f"当前已累积用户语料: {corpus_count} 句\n\n"
            f"重训将合并基础语料 + 用户语料 + 纠错反馈，\n"
            f"生成新的 3-gram 模型（重启后生效）。\n\n开始训练？",
            QMessageBox.Yes | QMessageBox.No
        )
        if reply != QMessageBox.Yes:
            return

        self.status_bar.showMessage("🧠 正在重训语言模型...")
        try:
            script = os.path.join(os.path.dirname(os.path.abspath(__file__)), "train_lm.py")
            result = subprocess.run(
                [sys.executable, script],
                capture_output=True, text=True, timeout=120,
                cwd=os.path.dirname(os.path.abspath(__file__))
            )
            if result.returncode == 0:
                # 提取关键信息
                output = result.stdout
                Toast.show_toast(self, "语言模型重训完成，重启后生效", "success")
                self.status_bar.showMessage("🧠 语言模型已更新，重启程序后生效")
                print(f"[LM] 重训完成:\n{output[-500:]}")
                self._refresh_topk_hotwords(silent=True)
            else:
                QMessageBox.warning(self, "训练失败", result.stderr[-300:] or "未知错误")
                self.status_bar.showMessage("⚠️ 语言模型训练失败")
        except subprocess.TimeoutExpired:
            QMessageBox.warning(self, "超时", "训练超时（>120s），请通过命令行执行")
        except Exception as e:
            QMessageBox.warning(self, "错误", str(e))

    def _view_crash_log(self):
        """查看崩溃日志"""
        logs = self.crash_logger.get_recent_logs(lines=200)
        if not logs:
            QMessageBox.information(self, "崩溃日志", "暂无日志记录")
            return

        dialog = QDialog(self)
        dialog.setWindowTitle("崩溃日志")
        dialog.setModal(True)
        dialog.resize(800, 600)

        layout = QVBoxLayout(dialog)

        text_edit = QTextEdit()
        text_edit.setReadOnly(True)
        text_edit.setText(''.join(logs))
        text_edit.setStyleSheet("""
            QTextEdit {
                background: #1a1a2e;
                color: #e0e0e0;
                font-family: monospace;
                font-size: 11px;
            }
        """)
        layout.addWidget(text_edit)

        bottom = QHBoxLayout()
        bottom.addStretch()
        clear_btn = QPushButton("清空日志")
        clear_btn.clicked.connect(self._clear_crash_log)
        close_btn = QPushButton("关闭")
        close_btn.clicked.connect(dialog.accept)
        bottom.addWidget(clear_btn)
        bottom.addWidget(close_btn)
        layout.addLayout(bottom)

        dialog.exec_()

    def _clear_crash_log(self):
        """清空崩溃日志"""
        log_path = self.crash_logger.get_log_path()
        try:
            with open(log_path, 'w', encoding='utf-8') as f:
                pass
            self.status_bar.showMessage("日志已清空")
        except Exception as e:
            QMessageBox.warning(self, "错误", f"清空日志失败: {e}")

    def _open_template_manager(self):
        """打开模板管理对话框"""
        dialog = TemplateManagerDialog(
            self.template_engine, self.current_dept, self
        )
        if dialog.exec_() == QDialog.Accepted:
            # 刷新科室和模板列表
            self._load_departments()
            # 重新选当前科室
            idx = self.dept_combo.findText(self.current_dept)
            if idx >= 0:
                self.dept_combo.setCurrentIndex(idx)
            self.status_bar.showMessage("📝 模板已更新")


# ==================== 模板管理对话框 ====================
class TemplateManagerDialog(QDialog):
    """模板增删改查管理界面"""

    def __init__(self, template_engine, current_dept, parent=None):
        super().__init__(parent)
        self.template_engine = template_engine
        self.current_dept = current_dept
        self._editing_index = -1  # -1 = 新增模式，>=0 = 编辑模式
        self.setWindowTitle("📝 模板管理")
        self.setModal(True)
        self.resize(700, 550)
        self._init_ui()
        self._refresh_dept_list()
        self._refresh_template_list()

    def _init_ui(self):
        layout = QVBoxLayout(self)

        # 上部：科室选择 + 模板列表
        top_split = QWidget()
        top_layout = QHBoxLayout(top_split)
        top_layout.setContentsMargins(0, 0, 0, 0)

        # 科室列表
        dept_widget = QWidget()
        dept_layout = QVBoxLayout(dept_widget)
        dept_layout.setContentsMargins(0, 0, 0, 0)
        dept_layout.addWidget(QLabel("科室："))
        self.dept_list = QListWidget()
        self.dept_list.setMaximumWidth(120)
        self.dept_list.itemClicked.connect(self._on_dept_selected)
        dept_layout.addWidget(self.dept_list)
        top_layout.addWidget(dept_widget)

        # 模板列表
        tpl_widget = QWidget()
        tpl_layout = QVBoxLayout(tpl_widget)
        tpl_layout.setContentsMargins(0, 0, 0, 0)
        tpl_layout.addWidget(QLabel("模板："))
        self.tpl_list = QListWidget()
        self.tpl_list.itemClicked.connect(self._on_template_selected)
        tpl_layout.addWidget(self.tpl_list)

        # 模板操作按钮
        tpl_btn_bar = QWidget()
        tpl_btn_layout = QHBoxLayout(tpl_btn_bar)
        tpl_btn_layout.setContentsMargins(0, 5, 0, 0)
        add_tpl_btn = QPushButton("➕ 新建")
        add_tpl_btn.clicked.connect(self._add_template)
        del_tpl_btn = QPushButton("🗑 删除")
        del_tpl_btn.clicked.connect(self._delete_template)
        tpl_btn_layout.addWidget(add_tpl_btn)
        tpl_btn_layout.addWidget(del_tpl_btn)
        tpl_layout.addWidget(tpl_btn_bar)

        top_layout.addWidget(tpl_widget)
        layout.addWidget(top_split)

        # 中部：模板编辑区
        edit_group = QGroupBox("模板内容编辑")
        edit_layout = QVBoxLayout(edit_group)

        name_layout = QHBoxLayout()
        name_layout.addWidget(QLabel("模板名称："))
        self.name_edit = QLineEdit()
        name_layout.addWidget(self.name_edit)
        edit_layout.addLayout(name_layout)

        self.content_edit = QTextEdit()
        self.content_edit.setPlaceholderText("在此编辑模板内容...\n用「字段名：」格式定义病历字段")
        self.content_edit.setMinimumHeight(200)
        edit_layout.addWidget(self.content_edit)

        layout.addWidget(edit_group)

        # 底部按钮
        bottom = QHBoxLayout()
        bottom.addStretch()
        cancel_btn = QPushButton("取消")
        cancel_btn.clicked.connect(self.reject)
        save_btn = QPushButton("✓ 保存")
        save_btn.clicked.connect(self._save_template)
        save_btn.setDefault(True)
        bottom.addWidget(cancel_btn)
        bottom.addWidget(save_btn)
        layout.addLayout(bottom)

    def _refresh_dept_list(self):
        """刷新科室列表"""
        self.dept_list.clear()
        depts = self.template_engine.get_departments()
        for dept in depts:
            self.dept_list.addItem(dept)
        # 选中当前科室
        for i in range(self.dept_list.count()):
            if self.dept_list.item(i).text() == self.current_dept:
                self.dept_list.setCurrentRow(i)
                break

    def _refresh_template_list(self):
        """刷新当前科室的模板列表"""
        self.tpl_list.clear()
        dept = self._get_selected_dept()
        if not dept:
            return
        templates = self.template_engine.get_templates(dept)
        for t in templates:
            self.tpl_list.addItem(t["name"])

    def _get_selected_dept(self):
        item = self.dept_list.currentItem()
        return item.text() if item else ""

    def _on_dept_selected(self, item):
        """切换科室"""
        self._editing_index = -1
        self._refresh_template_list()
        self.name_edit.clear()
        self.content_edit.clear()

    def _on_template_selected(self, item):
        """选中模板，加载内容"""
        dept = self._get_selected_dept()
        if not dept:
            return
        tpl_name = item.text()
        content = self.template_engine.get_template(dept, tpl_name)
        if content is not None:
            self.name_edit.setText(tpl_name)
            self.content_edit.setPlainText(content)
            self._editing_index = self.tpl_list.currentRow()

    def _add_template(self):
        """新增模板"""
        dept = self._get_selected_dept()
        if not dept:
            QMessageBox.information(self, "提示", "请先选择科室")
            return
        self._editing_index = -1
        self.name_edit.clear()
        self.content_edit.clear()
        self.name_edit.setFocus()

    def _delete_template(self):
        """删除模板"""
        dept = self._get_selected_dept()
        if not dept:
            return
        row = self.tpl_list.currentRow()
        if row < 0:
            QMessageBox.information(self, "提示", "请先选择要删除的模板")
            return
        tpl_name = self.tpl_list.item(row).text()
        confirm = QMessageBox.question(
            self, "确认删除",
            f"确定删除模板「{tpl_name}」吗？",
            QMessageBox.Yes | QMessageBox.No
        )
        if confirm == QMessageBox.Yes:
            templates = self.template_engine.get_templates(dept)
            if row < len(templates):
                templates.pop(row)
                self._save_dept_templates(dept, templates)
                self._refresh_template_list()
                self.name_edit.clear()
                self.content_edit.clear()
                self._editing_index = -1
                self.status_bar.showMessage(f"📝 模板已删除：{tpl_name}")

    def _save_template(self):
        """保存模板"""
        dept = self._get_selected_dept()
        if not dept:
            QMessageBox.warning(self, "提示", "请选择科室")
            return
        name = self.name_edit.text().strip()
        content = self.content_edit.toPlainText().strip()
        if not name:
            QMessageBox.warning(self, "提示", "请输入模板名称")
            return
        if not content:
            QMessageBox.warning(self, "提示", "请输入模板内容")
            return

        templates = self.template_engine.get_templates(dept)
        if self._editing_index >= 0 and self._editing_index < len(templates):
            # 更新已有模板
            templates[self._editing_index] = {"name": name, "content": content}
        else:
            # 新增模板
            templates.append({"name": name, "content": content})

        self._save_dept_templates(dept, templates)
        self._refresh_template_list()
        # 选中刚保存的模板
        for i in range(self.tpl_list.count()):
            if self.tpl_list.item(i).text() == name:
                self.tpl_list.setCurrentRow(i)
                break
        self._editing_index = self.tpl_list.currentRow()
        self.accept()

    def _save_dept_templates(self, dept, templates):
        """保存科室模板到文件"""
        import json
        import os
        templates_dir = os.path.join(os.path.dirname(__file__), "templates")
        path = os.path.join(templates_dir, f"{dept}.json")
        data = {"templates": templates}
        with open(path, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        # 同时更新内存中的模板
        self.template_engine.load_templates()


# ==================== 病历结构化对话框 ====================
class SectionDialog(QDialog):
    """病历结构化编辑对话框"""

    # 标准病历字段（按常见顺序排列）
    STANDARD_FIELDS = [
        "主诉", "现病史", "既往史", "体格检查", "辅助检查",
        "初步诊断", "诊疗经过", "出院情况", "出院医嘱",
        "术前诊断", "手术名称", "术中情况", "术后诊断", "术后医嘱",
        # 影像科字段
        "检查项目", "检查部位", "检查方法", "影像表现", "诊断意见",
        "增强特征", "血管描述", "超声所见", "超声提示",
        "建议",
        "急救措施", "用药情况", "效果评估",
        "日期", "患者情况", "处理意见"
    ]

    def __init__(self, parser, raw_text, parent=None):
        super().__init__(parent)
        self.parser = parser
        self.raw_text = raw_text
        self.field_edits = {}
        self.setWindowTitle("📋 病历结构化")
        self.setModal(True)
        self.resize(750, 600)
        self._init_ui()

    def _init_ui(self):
        # 先解析文本
        sections = self.parser.parse(self.raw_text)

        layout = QVBoxLayout(self)

        # 顶部说明
        info = QLabel(f"从语音文本中识别出 {len(sections)} 个病历部分，可逐项编辑")
        info.setStyleSheet("color: #b8c5d6; font-size: 12px; padding: 5px;")
        layout.addWidget(info)

        # 滚动区域
        from PyQt5.QtWidgets import QScrollArea
        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        scroll.setStyleSheet("""
            QScrollArea {
                border: 1px solid rgba(0,212,255,0.1);
                border-radius: 8px;
                background: rgba(255,255,255,0.02);
            }
        """)

        container = QWidget()
        form_layout = QVBoxLayout(container)
        form_layout.setSpacing(8)

        # 为每个标准字段创建输入框
        for field in self.STANDARD_FIELDS:
            row = QWidget()
            row_layout = QHBoxLayout(row)
            row_layout.setContentsMargins(0, 0, 0, 0)
            row_layout.setSpacing(8)

            # 字段标签
            label = QLabel(f"{field}：")
            label.setFixedWidth(70)
            label.setStyleSheet("color: #00d4ff; font-size: 13px; font-weight: bold;")
            row_layout.addWidget(label)

            # 内容输入框
            edit = QLineEdit()
            content = sections.get(field, "")
            edit.setText(content)
            edit.setPlaceholderText(f"请输入{field}...")
            edit.setStyleSheet("""
                QLineEdit {
                    background: rgba(255,255,255,0.05);
                    border: 1px solid rgba(0,212,255,0.15);
                    border-radius: 6px;
                    padding: 6px 10px;
                    color: #e0e0e0;
                    font-size: 13px;
                }
                QLineEdit:focus {
                    border: 1px solid rgba(0,212,255,0.4);
                }
            """)
            row_layout.addWidget(edit)
            self.field_edits[field] = edit

            # 如果有内容，高亮标签
            if content:
                label.setStyleSheet("color: #51cf66; font-size: 13px; font-weight: bold;")

            form_layout.addWidget(row)

        scroll.setWidget(container)
        layout.addWidget(scroll)

        # 底部按钮
        bottom = QHBoxLayout()
        bottom.addStretch()

        clear_btn = QPushButton("🗑 清空")
        clear_btn.clicked.connect(self._clear_all)
        clear_btn.setStyleSheet("""
            QPushButton {
                background: rgba(255,80,80,0.1);
                color: #ff6b6b;
                padding: 8px 20px;
                border-radius: 15px;
                border: 1px solid rgba(255,80,80,0.2);
            }
            QPushButton:hover { background: rgba(255,80,80,0.2); }
        """)
        bottom.addWidget(clear_btn)

        cancel_btn = QPushButton("取消")
        cancel_btn.clicked.connect(self.reject)
        cancel_btn.setStyleSheet("""
            QPushButton {
                background: rgba(255,255,255,0.1);
                color: #b8c5d6;
                padding: 8px 20px;
                border-radius: 15px;
                border: 1px solid rgba(255,255,255,0.1);
            }
            QPushButton:hover { background: rgba(255,255,255,0.15); }
        """)
        bottom.addWidget(cancel_btn)

        confirm_btn = QPushButton("✓ 确认填充")
        confirm_btn.clicked.connect(self._confirm)
        confirm_btn.setStyleSheet("""
            QPushButton {
                background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                    stop:0 #00d4ff, stop:1 #0066ff);
                color: #0a0e27;
                font-weight: bold;
                padding: 8px 24px;
                border-radius: 15px;
            }
            QPushButton:hover { padding: 8px 28px; }
        """)
        bottom.addWidget(confirm_btn)

        layout.addLayout(bottom)

    def _clear_all(self):
        """清空所有字段"""
        for edit in self.field_edits.values():
            edit.clear()

    def _confirm(self):
        """确认，生成结构化文本"""
        self.accept()

    def get_result(self):
        """获取结构化后的病历文本"""
        lines = []
        for field in self.STANDARD_FIELDS:
            edit = self.field_edits.get(field)
            if edit:
                content = edit.text().strip()
                if content:
                    lines.append(f"{field}：{content}")
        return "\n\n".join(lines)


# ==================== WebView 主窗口 ====================
class WebViewApp(QMainWindow):
    """基于 QWebEngineView 的新版主界面"""

    def __init__(self, db=None, current_user=None):
        super().__init__()
        self.db = db
        self.current_user = current_user or {}
        self.current_record_id = None
        self.current_dept = "通用"
        self._current_field = "主诉"
        self.is_listening = False
        self._record_start_ts = None
        self._auto_stop_timer = None

        uname = self.current_user.get("username", "")
        title = "星衍AI · 智能病历录入"
        if uname:
            title += f"  |  {uname}"
        self.setWindowTitle(title)
        self.setGeometry(80, 60, 1400, 900)

        # 核心引擎
        self.rule_engine = RuleEngine()
        self.corrector = Corrector(rule_engine=self.rule_engine)
        self.template_engine = TemplateEngine()
        self.parser = SectionParser()
        self.classifier = MedicalClassifier()
        self.qa_engine = KnowledgeQA()
        self.feedback = CorrectionFeedback()
        self.crash_logger = CrashLogger()
        self.crash_logger.log_event("应用启动(WebView)")
        self.voice_command = VoiceCommandParser()

        # ASR
        model_path = os.path.join(os.path.dirname(__file__), "model")
        self.asr = ASREngine(model_path=model_path)
        self.listen_thread = None

        # 常用句
        self._presets_path = os.path.join(os.path.dirname(__file__), "field_presets.json")
        self._presets_data = {}
        self._load_presets()
        self._asr_preview_timer = None

        # WebView
        from webview_bridge import WebViewMain
        self.webview = WebViewMain(self)
        self.setCentralWidget(self.webview)

        # 连接桥接信号
        br = self.webview.bridge
        br.sig_rec_toggle.connect(self._toggle_recording)
        br.sig_save.connect(self._save_record)
        br.sig_qa.connect(self._show_qa)
        br.sig_qa_ask.connect(self._on_qa_ask)
        br.sig_qa_close.connect(self.webview.js_close_qa)
        br.sig_template_mgr.connect(self._open_template_manager)
        br.sig_retrain.connect(self._retrain_lm)
        br.sig_dept_changed.connect(self._on_dept_changed)
        br.sig_template_changed.connect(self._on_template_changed)
        br.sig_field_changed.connect(self._on_field_changed)
        br.sig_editor_changed.connect(self._on_editor_changed)
        br.sig_chip_click.connect(self._on_chip_click)
        br.sig_preset_click.connect(self._on_preset_click)
        br.sig_add_preset.connect(self._on_add_preset)
        br.sig_asr_accept.connect(self._on_asr_accept)
        br.sig_asr_reject.connect(self._on_asr_reject)
        br.sig_asr_retry.connect(self._on_asr_retry)

        # 录音计时
        self._duration_timer = QTimer(self)
        self._duration_timer.setInterval(1000)
        self._duration_timer.timeout.connect(self._update_rec_time)
        self._rec_seconds = 0

        # 延迟初始化 UI 数据（等页面加载完成）
        self.webview.set_on_ready(self._init_webview_data)
        self._recent_history_opened = False

    def _load_presets(self):
        try:
            with open(self._presets_path, 'r', encoding='utf-8') as f:
                self._presets_data = json.load(f)
        except Exception:
            self._presets_data = {}

    def _save_presets(self):
        with open(self._presets_path, 'w', encoding='utf-8') as f:
            json.dump(self._presets_data, f, ensure_ascii=False, indent=2)

    def _init_webview_data(self):
        """WebView 加载完成后推送初始数据"""
        # 科室列表（内科优先，与默认模板兜底一致）
        depts = ["内科", "外科", "妇产科", "儿科", "全科"]
        self.webview.js_set_depts(depts)
        # 模板列表
        self._refresh_templates()
        # 默认加载入院记录骨架，让用户启动即看到结构化字段（语音录入可直接归类）
        self._on_template_changed("入院记录")
        # 字段导航
        fields = ["主诉", "现病史", "既往史", "个人史", "婚育史", "家族史",
                  "体格检查", "辅助检查", "初步诊断", "诊疗计划"]
        self.webview.js_set_fields(fields, "主诉")
        # 状态栏统计
        hw = len(self.asr._current_hotwords.split()) if self.asr._current_hotwords else 0
        kg_count = len(self.qa_engine.kg.entities) if self.qa_engine.kg else 0
        drug_count = len(self.qa_engine.kg.drug_inserts) if self.qa_engine.kg else 0
        self.webview.js_set_stats(str(hw), str(kg_count), str(drug_count))
        # 初始加载最近病历
        try:
            self._show_history()
        except Exception:
            pass
        # 默认科室热词
        self.corrector.set_department("通用")
        self.asr.set_hotwords("通用")
        # 默认上下文面板
        self._update_context_panel("主诉")
        # 启动检查
        self._show_startup_checks()

    def _dept_for_templates(self):
        """模板查找用科室：通用/全科无独立模板，兜底到内科"""
        dept = self.current_dept
        if dept in ("通用", "全科"):
            return "内科"
        return dept

    def _refresh_templates(self):
        dept = self._dept_for_templates()
        tpls = self.template_engine.get_templates(dept)
        tpl_names = [t["name"] for t in tpls]
        self.webview.js_set_templates(tpl_names)

    def _update_context_panel(self, field):
        """更新右侧上下文面板"""
        # 从 field_words.json 获取常用词
        import os as _os
        fw_path = _os.path.join(_os.path.dirname(__file__), "field_words.json")
        sections = []
        try:
            with open(fw_path, 'r', encoding='utf-8') as f:
                fw = json.load(f)
            field_data = fw.get(field, {})
            terms = field_data.get("terms", {})
            if isinstance(terms, dict):
                for cat, words in terms.items():
                    sections.append({"title": cat, "words": words[:12]})
            elif isinstance(terms, list):
                sections.append({"title": "常用词", "words": terms[:15]})
        except Exception:
            pass
        # 常用句
        presets = self._presets_data.get(field, [])
        self.webview.js_set_context_panel(f"常用词 · {field}", sections, presets)

    # ─── 桥接信号处理 ───

    def _toggle_recording(self):
        if self.is_listening:
            self._stop_recording()
        else:
            self._start_recording()

    def _start_recording(self):
        self.is_listening = True
        self._rec_seconds = 0
        self._record_start_ts = time.time()
        self._duration_timer.start()
        self.webview.js_set_recording(True, "主诉", "00:00")
        # 启动 ASR（复用 threads.create_listen_thread 统一线程管理）
        self.listen_thread = create_listen_thread(
            self.asr,
            on_text_ready=self._on_recognized,
            on_partial=self._on_partial,
        )
        self.listen_thread.start()

    def _stop_recording(self):
        self.is_listening = False
        self._duration_timer.stop()
        self.webview.js_set_recording(False)
        # 复用 threads.stop_listen_thread 统一停止逻辑
        stop_listen_thread(self.listen_thread, timeout_ms=3000)
        # 兜底：3 秒后若预览面板仍残留（无最终结果时）则清空
        try:
            QTimer.singleShot(3000, lambda: self.webview.js_set_asr_preview(""))
        except Exception:
            pass

    def _update_rec_time(self):
        self._rec_seconds += 1
        m, s = divmod(self._rec_seconds, 60)
        self.webview.js_set_recording(True, "", f"{m:02d}:{s:02d}")

    def _on_recognized(self, text):
        if not text:
            return
        self._last_asr_preview_text = text
        try:
            self.webview.js_set_asr_preview(text)
            self.webview.js_set_asr_actions(True)
        except Exception as e:
            print(f"[ASR] 预览更新失败: {e}")
        # 语音命令拦截
        cmd, arg = self.voice_command.parse(text)
        if cmd == "stop_record":
            self._stop_recording()
            return
        if cmd == "save":
            self._save_record()
            return
        # 结构化填充：将识别文本智能分配到模板字段
        self._fill_and_update(text)
        # 保留最终预览一会再清空，避免一闪而过
        try:
            if self._asr_preview_timer is not None:
                self._asr_preview_timer.stop()
            self._asr_preview_timer = QTimer(self)
            self._asr_preview_timer.setSingleShot(True)
            self._asr_preview_timer.timeout.connect(lambda: self.webview.js_set_asr_preview(""))
            self._asr_preview_timer.timeout.connect(lambda: self.webview.js_set_asr_actions(False))
            self._asr_preview_timer.start(2500)
        except Exception as e:
            print(f"[ASR] 预览延时清空失败: {e}")

    def _fill_and_update(self, asr_text):
        """调用 MedicalClassifier.incremental_fill 做结构化填充"""
        try:
            base = getattr(self, '_last_editor_text', '') or ''
            if not base:
                # 尝试从模板获取
                dept = self._dept_for_templates()
                tpls = self.template_engine.get_templates(dept)
                if tpls:
                    base = tpls[0].get('content', '')
            if not base:
                # 无模板，直接插入
                print(f"[Fill] 无模板兜底，直接插入: {asr_text[:50]}")
                self.webview.js_insert_text(asr_text)
                return
            inferred = self.classifier.extract_basic_fields(asr_text)
            print(f"[Fill] 识别文本: {asr_text[:80]}")
            print(f"[Fill] base首行: {base.splitlines()[0] if base else '(空)'!r} (共{len(base)}字)")
            print(f"[Fill] 推断字段: {inferred}")
            filled = self.classifier.incremental_fill(asr_text, base)
            if filled != base:
                html = self._text_to_editor_html(filled)
                self.webview.js_set_content(html)
                self._last_editor_text = filled
                print("[Fill] ✓ 填充完成，已更新编辑器")
            else:
                # 填充无变化，降级为直接插入
                print("[Fill] ⚠ 填充无变化，降级为直接插入")
                self.webview.js_insert_text(asr_text)
        except Exception as e:
            import traceback
            print(f"[Fill] error: {e}")
            traceback.print_exc()
            self.webview.js_insert_text(asr_text)

    def _on_partial(self, text):
        if not text:
            return
        # 流式结果连续刷新预览；若正在显示最终结果，先取消延时清空
        try:
            if self._asr_preview_timer is not None:
                self._asr_preview_timer.stop()
                self._asr_preview_timer = None
            self.webview.js_set_asr_preview(text)
        except Exception as e:
            print(f"[ASR] 预览更新失败: {e}")

    def _on_editor_changed(self, text):
        """编辑器内容变化（用于保存时获取）"""
        self._last_editor_text = text
        try:
            alerts = self.rule_engine.realtime_checks(text, dept=self.current_dept, field=getattr(self, '_current_field', ''))
            self.webview.js_set_qc_status(len(alerts))
        except Exception:
            pass
        print(f"[Editor] 内容同步: {len(text)}字 首行={text.splitlines()[0][:30] if text else '(空)'!r}")

    def _save_record(self):
        text = getattr(self, '_last_editor_text', '')
        if not text:
            self.webview.get_editor_text(self._do_save)
        else:
            self._do_save(text)

    def _do_save(self, text):
        if not text or not text.strip():
            self.webview.js_show_toast("没有内容可保存")
            return
        if not self.db or not self.current_user:
            self.webview.js_show_toast("未登录")
            return
        content = text.strip()
        patient_name = ""
        m = re.search(r'姓名[：:]\s*([^\s　\n]{1,10})', content)
        if m:
            patient_name = m.group(1).strip()
        dept = self.current_dept if self.current_dept != "通用" else ""
        if self.current_record_id is None:
            self.current_record_id = self.db.create_record(
                self.current_user["id"], patient_name, dept, "", content, "草稿"
            )
            self.webview.js_show_toast("💾 病历已保存")
        else:
            self.db.update_record(self.current_record_id, content=content)
            self.webview.js_show_toast("💾 病历已更新")
        # 收集语料
        try:
            self.feedback.collect_corpus(content)
        except Exception:
            pass

    def _show_qa(self):
        self._show_qa_for_selection()

    def _on_qa_ask(self, question):
        """用户在问答面板输入问题：交给知识图谱引擎回答"""
        try:
            question = (question or '').strip()
            if not question:
                return
            result = self.qa_engine.answer(question)
            answer_text = result.get('text', '') if isinstance(result, dict) else str(result)
            import html as _html
            html_text = _html.escape(answer_text or '未找到相关知识，请换个问法试试。')
            html_text = html_text.replace('\n', '<br>')
            self.webview.js_set_qa(f'<b>问：{_html.escape(question)}</b><br><br>{html_text}')
            print(f"[QA] 提问: {question[:40]} -> 回答{len(answer_text)}字")
        except Exception as e:
            print(f"[QA] 提问处理失败: {e}")

    def _show_qa_dialog(self):
        try:
            from qa_dialog import KnowledgeQADialog
            dlg = KnowledgeQADialog(self.qa_engine, self)
            dlg.exec_()
        except Exception as e:
            QMessageBox.warning(self, "问答错误", f"知识问答模块加载失败：\n{e}")

    def _show_history(self):
        try:
            if not self.db:
                self.webview.js_set_history('数据库未初始化')
                return
            user_id = self.current_user.get("id") if isinstance(self.current_user, dict) else None
            records = self.db.list_records(user_id=user_id, limit=20)
            if not records:
                self.webview.js_set_history('暂无最近病历')
                return
            lines = []
            for rec in records[:8]:
                title = rec.get('patient_name') or ('病历#' + str(rec.get('record_id')))
                snippet = (rec.get('content') or '').strip().splitlines()[0]
                lines.append(f"{title}：{snippet}")
            self.webview.js_set_history('<br>'.join(lines))
        except Exception as e:
            print(f"[History] 失败: {e}")

    def _show_startup_checks(self):
        try:
            checks = []
            if not self.asr.is_ready():
                checks.append("ASR 模型未就绪，请检查模型目录")
            hw = len(self.asr._current_hotwords.split()) if self.asr._current_hotwords else 0
            if hw < 10:
                checks.append("当前热词较少，建议先做 Top-K 刷新")
            if self.rule_engine.get_stats().get("错别字规则数", 0) < 5:
                checks.append("纠错规则较少，建议补充 postprocess 规则")
            if checks:
                self.webview.js_show_toast("；".join(checks[:3]))
        except Exception as e:
            print(f"[Check] 失败: {e}")
        try:
            from qa_dialog import KnowledgeQADialog
            dlg = KnowledgeQADialog(self.qa_engine, self)
            dlg.exec_()
        except Exception as e:
            QMessageBox.warning(self, "问答错误", f"知识问答模块加载失败：\n{e}")

    def _open_template_manager(self):
        from main import TemplateManagerDialog
        dlg = TemplateManagerDialog(self.template_engine, self.current_dept, self)
        if dlg.exec_() == QDialog.Accepted:
            self._refresh_templates()

    def _retrain_lm(self):
        import subprocess
        reply = QMessageBox.question(self, "重训语言模型",
            "将合并用户语料+纠错反馈重训 3-gram 模型。\n开始？",
            QMessageBox.Yes | QMessageBox.No)
        if reply != QMessageBox.Yes:
            return
        script = os.path.join(os.path.dirname(os.path.abspath(__file__)), "train_lm.py")
        try:
            subprocess.run([sys.executable, script], capture_output=True, text=True, timeout=120,
                          cwd=os.path.dirname(os.path.abspath(__file__)))
            self.webview.js_show_toast("🧠 语言模型重训完成，重启后生效")
        except Exception as e:
            self.webview.js_show_toast(f"重训失败: {e}")

    def _on_dept_changed(self, dept):
        self.current_dept = dept
        self.corrector.set_department(dept)
        self.asr.set_hotwords(dept)
        self._refresh_templates()

    def _on_template_changed(self, tpl_name):
        dept = self._dept_for_templates()
        content = self.template_engine.get_template(dept, tpl_name)
        if content:
            # 将模板文本转为 HTML（高亮字段名）
            html = self._text_to_editor_html(content)
            self.webview.js_set_content(html)
            self._last_editor_text = content

    def _on_field_changed(self, field):
        self._current_field = field
        self._update_context_panel(field)
        # 设置 ASR 字段上下文
        self.asr.set_field_context(field)
        # 刷新字段级 Top-K 热词
        self._refresh_field_hotwords(field)

    def _refresh_field_hotwords(self, field):
        try:
            topk = self._get_topk_engine()
            if not topk or not field:
                return
            prompt_pack = topk.build_field_prompt_pack(field=field, dept=self.current_dept, doctor_id=self.current_user.get("id") if isinstance(self.current_user, dict) else None, top_k=220)
            self.asr.set_prompt_pack(prompt_pack)
            self.asr.apply_prompt_pack()
            self.asr.set_hotwords(self.current_dept)
        except Exception as e:
            print(f"[Main] 刷新字段级热词失败: {e}")

    def _on_chip_click(self, word):
        self.webview.js_insert_text(word)

    def _on_preset_click(self, sentence):
        self.webview.js_insert_text(sentence)

    def _on_add_preset(self):
        text, ok = QInputDialog.getMultiLineText(self, "添加常用句", "请输入常用句：", "")
        if ok and text.strip():
            field = self._current_field
            if field not in self._presets_data:
                self._presets_data[field] = []
            if text.strip() not in self._presets_data[field]:
                self._presets_data[field].append(text.strip())
                self._save_presets()
            self._update_context_panel(field)

    @staticmethod
    def _text_to_editor_html(text):
        """将纯文本模板转为带字段高亮的 HTML"""
        fields = ['姓名', '性别', '年龄', '民族', '婚姻状况', '出生地', '职业',
                  '入院时间', '入院方式', '病史陈述者', '可靠程度', '主诉', '现病史',
                  '既往史', '个人史', '婚育史', '家族史', '体格检查', '辅助检查',
                  '初步诊断', '鉴别诊断', '诊疗计划', '诊疗经过', '出院情况', '出院医嘱',
                  '术前诊断', '手术名称', '术中情况', '术后诊断', '术后医嘱']
        import html as _html
        escaped = _html.escape(text)
        for f in fields:
            pattern = re.escape(f) + r'[：:]'
            escaped = re.sub(pattern, f'<span class="fl">{f}：</span>', escaped)
        escaped = escaped.replace('\n', '<br>')
        return escaped

    def _on_asr_accept(self):
        try:
            text = getattr(self, '_last_asr_preview_text', '') or ''
            if text:
                self.webview.js_insert_text(text)
            self.webview.js_set_asr_preview('')
            self.webview.js_set_asr_actions(False)
            self._last_asr_preview_text = ''
        except Exception as e:
            print(f"[ASR] 接受预览失败: {e}")

    def _on_asr_reject(self):
        try:
            self.webview.js_set_asr_preview('')
            self.webview.js_set_asr_actions(False)
            self._last_asr_preview_text = ''
        except Exception as e:
            print(f"[ASR] 拒绝预览失败: {e}")

    def _on_asr_retry(self):
        try:
            self._stop_recording()
            self.webview.js_set_asr_preview('')
            self.webview.js_set_asr_actions(False)
            self._last_asr_preview_text = ''
            self._toggle_recording()
        except Exception as e:
            print(f"[ASR] 重试失败: {e}")

    def _show_qa_for_selection(self):
        try:
            text = getattr(self, '_last_editor_text', '') or ''
            if not text:
                # 兜底：异步从编辑器取文本（get_editor_text 是回调式异步）
                self.webview.get_editor_text(lambda t: self._do_show_qa(t or ''))
                return
            self._do_show_qa(text)
        except Exception as e:
            print(f"[QA] 失败: {e}")

    def _do_show_qa(self, text):
        try:
            suggestions = self._build_qa_suggestions(text)
            print(f"[QA] 文本{len(text)}字 -> 提示{len(suggestions)}字")
            self.webview.js_set_qa(suggestions)
        except Exception as e:
            print(f"[QA] 构建提示失败: {e}")

    def _show_history(self):
        try:
            if not self.db:
                self.webview.js_set_history('数据库未初始化')
                return
            user_id = self.current_user.get("id") if isinstance(self.current_user, dict) else None
            records = self.db.list_records(user_id=user_id, limit=20)
            if not records:
                self.webview.js_set_history('暂无最近病历')
                return
            lines = []
            for rec in records[:8]:
                title = rec.get('patient_name') or ('病历#' + str(rec.get('record_id')))
                snippet = (rec.get('content') or '').strip().splitlines()[0]
                lines.append(f"{title}：{snippet}")
            self.webview.js_set_history('<br>'.join(lines))
        except Exception as e:
            print(f"[History] 失败: {e}")

    def _build_qa_suggestions(self, text):
        # 先判断病历是否有实质内容（剥离字段名与标点）
        stripped = text or ''
        for f in self.classifier.STANDARD_FIELDS:
            stripped = stripped.replace(f, '')
        stripped = re.sub(r'[：:\s，,。、；;（）()]+', '', stripped)
        if len(stripped) < 4:
            return ('当前病历还没有实质内容。<br>'
                    '语音或打字填入内容后，这里会自动提示<b>可能疾病、用药、检查、质控提醒</b>。'
                    '<br><br>例：录入"患者男性，56岁，主诉胸痛2小时，'
                    '初步诊断冠心病"后再按 ⌘/ 试试。')
        items = []
        try:
            diseases = self.qa_engine.extract_diseases(text)[:5]
            if diseases:
                items.append('<b>可能疾病</b>：' + '、'.join(diseases))
        except Exception:
            pass
        try:
            drugs = self.qa_engine.extract_drugs(text)[:8]
            if drugs:
                items.append('<b>用药</b>：' + '、'.join(drugs))
        except Exception:
            pass
        try:
            exams = self.qa_engine.extract_exams(text)[:8]
            if exams:
                items.append('<b>检查</b>：' + '、'.join(exams))
        except Exception:
            pass
        try:
            alerts = self.rule_engine.realtime_checks(text, dept=self.current_dept, field=getattr(self, '_current_field', ''))[:8]
            if alerts:
                items.append('<b>质控提醒</b>：<br>' + '<br>'.join([a.get('message','') for a in alerts]))
        except Exception:
            pass
        return '<br><br>'.join(items) if items else '已扫描病历文本，未命中知识图谱中的疾病/药品/检查实体。可尝试补充诊断、用药或检查描述。'

    def closeEvent(self, event):
        if self.is_listening:
            self._stop_recording()
        super().closeEvent(event)


# ==================== 启动入口 ====================
def main():
    app = QApplication(sys.argv)
    app.setStyle('Fusion')

    # 设置全局字体
    font = QFont("Microsoft YaHei", 11)
    app.setFont(font)

    # ─── 授权检查 ───
    license_mgr = LicenseManager()
    status = license_mgr.check_license()

    if status["status"] in ("expired", "tampered"):
        # 试用到期或篡改：必须输入激活码才能继续
        dlg = ActivationDialog(license_mgr, status)
        if dlg.exec_() != QDialog.Accepted:
            sys.exit(0)
    # "trial" 或 "activated" → 直接进入

    # 登录（首次运行自动创建管理员）
    db = Database()
    login = LoginDialog(db)
    if login.exec_() != QDialog.Accepted or not login.current_user:
        sys.exit(0)

    # 默认启用新版 WebView UI，--legacy 回退原生界面
    use_webview = '--legacy' not in sys.argv
    if use_webview and _HAS_WEBENGINE:
        try:
            window = WebViewApp(db=db, current_user=login.current_user)
        except Exception as e:
            print(f"[Main] WebView 初始化失败，回退到原生 UI: {e}")
            window = MedVoiceApp(db=db, current_user=login.current_user)
    else:
        window = MedVoiceApp(db=db, current_user=login.current_user)
    window.show()

    sys.exit(app.exec_())


if __name__ == "__main__":
    main()
